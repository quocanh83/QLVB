import re
from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from rest_framework.decorators import action
from django.db import transaction, models
from django.utils import timezone
from django.http import HttpResponse
from .models import ComparisonProject, DraftVersion, ComparisonNode, ComparisonMapping, ComparisonAIResult, StandaloneReview
from .serializers import (
    ComparisonProjectSerializer, DraftVersionSerializer, 
    ComparisonNodeSerializer, ComparisonMappingSerializer,
    ComparisonAIResultSerializer, StandaloneReviewSerializer
)
from .utils.parser_engine import ComparisonParser
from .utils.legislative_diff import legislative_diff
from .utils.automap_service import automap_nodes
from .utils.export_service import export_comparison_table
from .utils.ai_service import UnifiedAIService
from .utils.reference_export import export_reference_excel, export_reference_word

class ComparisonProjectViewSet(viewsets.ModelViewSet):
    queryset = ComparisonProject.objects.all().order_by('-created_at')
    serializer_class = ComparisonProjectSerializer
    permission_classes = [permissions.IsAuthenticated]

    def perform_create(self, serializer):
        project = serializer.save(uploaded_by=self.request.user)
        # Tự động bóc tách bản gốc ngay khi tạo
        self._parse_base_document(project)

    def _parse_base_document(self, project):
        parser = ComparisonParser(project.base_file.path)
        structure = parser.parse()
        
        def save_nodes_recursive(node_list, parent=None):
            for node_data in node_list:
                children = node_data.pop('children', [])
                node = ComparisonNode.objects.create(
                    project=project,
                    parent=parent,
                    **node_data
                )
                if children:
                    save_nodes_recursive(children, parent=node)
        
        save_nodes_recursive(structure)
        
    @action(detail=True, methods=['post'])
    def replace_base_file(self, request, pk=None):
        """Thay thế văn bản gốc và bóc tách lại"""
        project = self.get_object()
        file_obj = request.FILES.get('base_file')
        if not file_obj:
            return Response({"error": "Vui lòng chọn file .docx mới."}, status=400)
            
        with transaction.atomic():
            # Xóa các node gốc cũ
            ComparisonNode.objects.filter(project=project, version__isnull=True).delete()
            # Cập nhật file
            project.base_file = file_obj
            project.save()
            # Bóc tách lại (Bản gốc không cần đánh số lại)
            self._parse_base_document(project)
            
        return Response(ComparisonProjectSerializer(project).data)

    @action(detail=True, methods=['get'])
    def base_nodes(self, request, pk=None):
        """Lấy tất cả node gốc của dự án này"""
        project = self.get_object()
        nodes = ComparisonNode.objects.filter(project=project).order_by('order_index')
        serializer = ComparisonNodeSerializer(nodes, many=True)
        return Response(serializer.data)

class DraftVersionViewSet(viewsets.ModelViewSet):
    queryset = DraftVersion.objects.all().order_by('-created_at')
    serializer_class = DraftVersionSerializer
    permission_classes = [permissions.IsAuthenticated]
    
    def get_permissions(self):
        if self.action == 'export_word':
            return [permissions.AllowAny()]
        return [permissions.IsAuthenticated()]

    def create(self, request, *args, **kwargs):
        project_id = request.data.get('project')
        user_note = request.data.get('user_note', '')
        file_obj = request.FILES.get('file_path')

        if not project_id or not file_obj:
            return Response({"error": "Thiếu project_id hoặc file dự thảo."}, status=400)

        # Tạo nhãn phiên bản: [Ngày giờ] - {user_note}
        from django.utils import timezone
        now = timezone.now().strftime('%Y-%m-%d %H:%M')
        version_label = f"[{now}] {user_note}".strip()

        with transaction.atomic():
            version = DraftVersion.objects.create(
                project_id=project_id,
                file_path=file_obj,
                user_note=user_note,
                version_label=version_label
            )
            # Bóc tách cấu trúc cho phiên bản này
            self._parse_draft_document(version)
            
            # Tự động ánh xạ từ phiên bản trước (nếu có) hoặc từ bản gốc
            self._initial_automap(version)

            # 1. Tự động kế thừa URL Google Sheet từ phiên bản cũ nhất gần đây
            prev_version = DraftVersion.objects.filter(project_id=project_id).exclude(id=version.id).order_by('-created_at').first()
            if prev_version and prev_version.explanation_sheet_url:
                version.explanation_sheet_url = prev_version.explanation_sheet_url
                version.save()

        # 2. Tự động đồng bộ thuyết minh nếu có URL
        if version.explanation_sheet_url:
            try:
                self._perform_gsheet_sync(version, version.explanation_sheet_url)
            except Exception as e:
                # Không chặn quá trình tạo version nếu đồng bộ lỗi
                print(f"Auto-sync explanation failed: {str(e)}")

        return Response(DraftVersionSerializer(version).data, status=201)

    def _perform_gsheet_sync(self, version, sheet_url):
        """Tiện ích thực hiện đồng bộ thuyết minh thực sự với thuật toán đối soát đa lớp"""
        from .utils.gsheet_sync import sync_explanation_from_gsheet, extract_norm_label, get_content_fingerprint
        exp_dict = sync_explanation_from_gsheet(sheet_url)
        
        count = 0
        # Đồng bộ thuyết minh cho các mục Điều, Phụ lục, Chương, Mục
        nodes = ComparisonNode.objects.filter(version=version, node_type__in=['Điều', 'Phụ lục', 'Chương', 'Mục'])
        for node in nodes:
            # 1. Khớp theo ID node_{id}
            node_id_key = f"node_{node.id}"
            # 2. Khớp theo Nhãn chuẩn hóa
            label_key = extract_norm_label(node.node_label)
            # 3. Khớp theo Fingerprint nội dung
            fp_key = f"fp_{get_content_fingerprint(node.content)}"

            # Ưu tiên khớp ID -> Label -> Fingerprint
            match = exp_dict.get(node_id_key)
            if not match: match = exp_dict.get(label_key)
            if not match: match = exp_dict.get(fp_key)

            if match:
                node.explanation = match.get('exp', '')
                node.save()
                count += 1
        return count

    def _parse_draft_document(self, version):
        parser = ComparisonParser(version.file_path.path)
        # Bóc tách cấu trúc chuẩn cho bản dự thảo
        structure = parser.parse()
        
        def save_nodes_recursive(node_list, parent=None):
            for node_data in node_list:
                children = node_data.pop('children', [])
                node = ComparisonNode.objects.create(
                    version=version,
                    parent=parent,
                    **node_data
                )
                if children:
                    save_nodes_recursive(children, parent=node)
        
        save_nodes_recursive(structure)

    def _initial_automap(self, version):
        project = version.project
        # 1. Tìm phiên bản trước đó gần nhất (không phải phiên bản hiện tại)
        prev_version = DraftVersion.objects.filter(project=project).exclude(id=version.id).order_by('-created_at').first()
        
        inherited_mapping = {}
        if prev_version:
            # Lấy ánh xạ của phiên bản trước
            prev_mappings = ComparisonMapping.objects.filter(version=prev_version).select_related('draft_node')
            # Tạo bản đồ: {id_node_gốc: nhãn_node_dự_thảo_cũ}
            # Chúng ta dùng nhãn (node_label) để làm "cầu nối" giữa 2 phiên bản dự thảo (V1 -> V2)
            prev_map_by_label = {m.base_node_id: m.draft_node.node_label for m in prev_mappings}
            
            # Lấy danh sách node của phiên bản hiện tại, đánh chỉ mục theo nhãn
            current_draft_nodes_by_label = {
                n.node_label: n.id 
                for n in ComparisonNode.objects.filter(version=version)
            }
            
            # Xây dựng ánh xạ kế thừa: Nếu Điều 1 gốc v1 khớp nhãn "Điều 5" dự thảo v1, 
            # thì ở v2 ta vẫn cố gắng khớp Điều 1 gốc với cái gì có nhãn "Điều 5" ở dự thảo v2.
            for b_id, label in prev_map_by_label.items():
                if label in current_draft_nodes_by_label:
                    inherited_mapping[b_id] = current_draft_nodes_by_label[label]

        # 2. Thực hiện automap tiêu chuẩn (cho các node chưa có trong ánh xạ kế thừa)
        base_nodes = ComparisonNode.objects.filter(project=project, node_type__in=['Điều', 'Phụ lục'], version__isnull=True)
        draft_nodes = ComparisonNode.objects.filter(version=version, node_type__in=['Điều', 'Phụ lục'])
        
        standard_mapping = automap_nodes(base_nodes, draft_nodes)
        
        # 3. Gộp kết quả: Ánh xạ kế thừa (Inherited) được ưu tiên vì nó chứa các tùy chỉnh thủ công của người dùng
        # Dictionary merge in Python 3.9+: final_mapping = standard_mapping | inherited_mapping
        final_mapping = {**standard_mapping, **inherited_mapping}
        
        # 4. Lưu vào DB
        new_mappings = [
            ComparisonMapping(
                project=project,
                version=version,
                base_node_id=b_id,
                draft_node_id=d_id
            ) for b_id, d_id in final_mapping.items()
        ]
        
        if new_mappings:
            ComparisonMapping.objects.bulk_create(new_mappings)

    def _get_full_content(self, node, nodes_by_parent=None):
        """Gộp nội dung tiêu đề và tất cả các con (Khoản, Điểm) của một node. Sử dụng map bộ nhớ để tránh N+1."""
        full_text = node.content if node.content else ""
        
        if node.node_type == 'Chương':
            return full_text.strip()
            
        # Ưu tiên lấy từ map nếu có để tránh truy vấn DB
        if nodes_by_parent is not None:
            children = nodes_by_parent.get(node.id, [])
        else:
            children = ComparisonNode.objects.filter(parent=node).order_by('order_index')
            
        for child in children:
            child_text = self._get_full_content(child, nodes_by_parent)
            if child_text:
                full_text += f"\n{child.node_label} {child_text}"
        return full_text.strip()

    def _get_full_explanation(self, node, nodes_by_parent=None):
        """Gộp thuyết minh của node và tất cả các mục con. Sử dụng map bộ nhớ để tránh N+1."""
        full_exp = node.explanation if node.explanation else ""
        
        if node.node_type == 'Chương':
            return full_exp.strip()
            
        if nodes_by_parent is not None:
            children = nodes_by_parent.get(node.id, [])
        else:
            children = ComparisonNode.objects.filter(parent=node).order_by('order_index')
            
        for child in children:
            child_exp = self._get_full_explanation(child, nodes_by_parent)
            if child_exp:
                full_exp += f"\n{child.node_label} {child_exp}"
        return full_exp.strip()

    def _get_interleaved_rows(self, version):
        """Logic trộn hàng tối ưu hóa: Sử dụng bulk loading và in-memory processing"""
        project = version.project
        
        # 1. Bulk load TOÀN BỘ nodes của dự án này để xử lý trong bộ nhớ
        # Bao gồm cả nodes của v bản gốc (project=project, version is null) và nodes của version hiện tại
        all_relevant_nodes = list(ComparisonNode.objects.filter(
            models.Q(project=project, version__isnull=True) | models.Q(version=version)
        ).select_related('parent').order_by('order_index'))
        
        # 2. Xây dựng Parent-Child Map & Node Map (O(N))
        from collections import defaultdict
        nodes_by_parent = defaultdict(list)
        base_nodes = []
        draft_nodes = []
        
        for node in all_relevant_nodes:
            if node.parent_id:
                nodes_by_parent[node.parent_id].append(node)
            
            # Chỉ lấy các node cấp độ cao cho danh sách hàng
            if node.node_type in ['Chương', 'Điều', 'Phụ lục', 'Vấn đề khác']:
                if node.version_id is None:
                    base_nodes.append(node)
                elif node.version_id == version.id:
                    draft_nodes.append(node)
        
        # 3. Lấy Mappings tối ưu
        mappings = ComparisonMapping.objects.filter(version=version).select_related('draft_node', 'base_node')
        b_to_d = {m.base_node_id: m.draft_node for m in mappings}
        d_to_b = {m.draft_node_id: m.base_node for m in mappings}
        
        rows = []
        consumed_draft_ids = set()
        
        # Duyệt theo danh sách Gốc
        for b_node in base_nodes:
            d_node = b_to_d.get(b_node.id)
            
            if d_node:
                for skip_d in draft_nodes:
                    if skip_d.order_index < d_node.order_index and skip_d.id not in consumed_draft_ids:
                        if skip_d.id not in d_to_b:
                            rows.append(self._build_row(None, skip_d, nodes_by_parent))
                            consumed_draft_ids.add(skip_d.id)
                
                rows.append(self._build_row(b_node, d_node, nodes_by_parent))
                consumed_draft_ids.add(d_node.id)
            else:
                rows.append(self._build_row(b_node, None, nodes_by_parent))
                
        for last_d in draft_nodes:
            if last_d.id not in consumed_draft_ids:
                rows.append(self._build_row(None, last_d, nodes_by_parent))
                
        return rows

    def _build_row(self, b_node, d_node, nodes_by_parent=None):
        b_full = self._get_full_content(b_node, nodes_by_parent) if b_node else ""
        d_full = self._get_full_content(d_node, nodes_by_parent) if d_node else ""
        b_exp = self._get_full_explanation(b_node, nodes_by_parent) if b_node else ""
        d_exp = self._get_full_explanation(d_node, nodes_by_parent) if d_node else ""
        
        # Nếu không có dự thảo, lấy thuyết minh của bản gốc (dành cho trường hợp bãi bỏ)
        final_exp = d_exp if d_node else b_exp
        
        return {
            "base_node": {
                "id": b_node.id, "node_label": b_node.node_label, "content": b_full, "node_type": b_node.node_type,
                "explanation": b_exp
            } if b_node else None,
            "draft_node": {
                "id": d_node.id, "node_label": d_node.node_label, "content": d_full, "node_type": d_node.node_type,
                "explanation": d_exp
            } if d_node else None,
            "display_explanation": final_exp,
            "diff_content": legislative_diff(b_full, d_full) if d_node and b_node else (d_full if d_node else "")
        }

    @action(detail=True, methods=['get'])
    def workspace_data(self, request, pk=None):
        """Lấy dữ liệu so sánh căn hàng ngang - Logic Interleaving (Base-centric)"""
        try:
            version = self.get_object()
            rows = self._get_interleaved_rows(version)
            return Response({
                "project_name": version.project.name,
                "version_label": version.version_label,
                "explanation_sheet_url": version.explanation_sheet_url,
                "rows": rows
            })
        except Exception as e:
            return Response({"error": f"Lỗi nạp workspace: {str(e)}"}, status=500)

    @action(detail=True, methods=['post'])
    def add_manual_row(self, request, pk=None):
        """Thêm một hàng trống vào bản gốc để người dùng tự ghép nối"""
        version = self.get_object()
        project = version.project
        insert_after_id = request.data.get('insert_after_id')
        
        with transaction.atomic():
            if insert_after_id:
                try:
                    after_node = ComparisonNode.objects.get(id=insert_after_id, project=project, version__isnull=True)
                    new_order = after_node.order_index + 1
                    # Đẩy các node phía sau xuống
                    ComparisonNode.objects.filter(
                        project=project, 
                        version__isnull=True, 
                        order_index__gte=new_order
                    ).update(order_index=models.F('order_index') + 1)
                except ComparisonNode.DoesNotExist:
                    from django.db.models import Max
                    new_order = (ComparisonNode.objects.filter(project=project, version__isnull=True).aggregate(Max('order_index'))['order_index__max'] or 0) + 1
            else:
                from django.db.models import Max
                new_order = (ComparisonNode.objects.filter(project=project, version__isnull=True).aggregate(Max('order_index'))['order_index__max'] or 0) + 1
            
            node = ComparisonNode.objects.create(
                project=project,
                node_type='Vấn đề khác',
                node_label='Hàng thủ công',
                content='',
                order_index=new_order
            )
        return Response(ComparisonNodeSerializer(node).data)

    @action(detail=True, methods=['post'])
    def reorder_node(self, request, pk=None):
        """Di chuyển node lên/xuống (chỉ áp dụng cho base nodes)"""
        version = self.get_object()
        project = version.project
        node_id = request.data.get('node_id')
        direction = request.data.get('direction') # 'up' or 'down'
        
        if not node_id or direction not in ['up', 'down']:
            return Response({"error": "Dữ liệu không hợp lệ."}, status=400)
            
        with transaction.atomic():
            try:
                node = ComparisonNode.objects.get(id=node_id, project=project, version__isnull=True)
                
                if direction == 'up':
                    swap_node = ComparisonNode.objects.filter(
                        project=project, version__isnull=True, order_index__lt=node.order_index
                    ).order_by('-order_index').first()
                else:
                    swap_node = ComparisonNode.objects.filter(
                        project=project, version__isnull=True, order_index__gt=node.order_index
                    ).order_by('order_index').first()
                    
                if swap_node:
                    temp_order = node.order_index
                    node.order_index = swap_node.order_index
                    swap_node.order_index = temp_order
                    node.save()
                    swap_node.save()
                    return Response({"message": "Đổi vị trí thành công"})
                else:
                    return Response({"message": "Không thể di chuyển thêm"})
            except ComparisonNode.DoesNotExist:
                return Response({"error": "Không tìm thấy Node."}, status=404)

    @action(detail=True, methods=['post'])
    def delete_manual_node(self, request, pk=None):
        """Xóa hàng thủ công"""
        version = self.get_object()
        project = version.project
        node_id = request.data.get('node_id')
        
        if not node_id:
            return Response({"error": "Thiếu node_id"}, status=400)
            
        try:
            # Chỉ cho phép xóa node 'Hàng thủ công'
            node = ComparisonNode.objects.get(id=node_id, project=project, version__isnull=True, node_label='Hàng thủ công')
            node.delete()
            return Response({"message": "Đã xóa hàng thủ công"})
        except ComparisonNode.DoesNotExist:
            return Response({"error": "Không tìm thấy hàng thủ công hợp lệ"}, status=404)

    @action(detail=True, methods=['get'])
    def export_word(self, request, pk=None):
        """Xuất file Word bảng đối chiếu - Logic Interleaving (Base-centric)"""
        user = request.user
        if not user.is_authenticated:
            token = request.query_params.get('token')
            if not token: return Response({"error": "Vui lòng cung cấp token"}, status=401)
            from rest_framework_simplejwt.tokens import AccessToken
            from django.contrib.auth import get_user_model
            try:
                access_token = AccessToken(token)
                user = get_user_model().objects.get(id=access_token['user_id'])
                request.user = user
            except Exception: return Response({"error": "Token không hợp lệ"}, status=401)
        
        version = self.get_object()
        project = version.project
        rows = self._get_interleaved_rows(version)
        word_rows = []
        for r in rows:
            word_rows.append({
                "base_node": {"node_label": r["base_node"]["node_label"], "content": r["base_node"]["content"]} if r["base_node"] else None,
                "draft_node": {"node_label": r["draft_node"]["node_label"], "content": r["draft_node"]["content"]} if r["draft_node"] else None,
                "diff_content": r["diff_content"]
            })
        
        stream = export_comparison_table(
            project.name, 
            version.version_label, 
            word_rows,
            base_name=project.base_document_name,
            draft_name=project.draft_document_name
        )
        
        response = HttpResponse(
            stream.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        safe_name = f"Bang_doi_chieu_V{version.id}.docx"
        response['Content-Disposition'] = f'attachment; filename="{safe_name}"'
        return response

    @action(detail=True, methods=['post'])
    def update_mapping(self, request, pk=None):
        """Cập nhật ánh xạ thủ công"""
        version = self.get_object()
        base_node_id = request.data.get('base_node_id')
        draft_node_id = request.data.get('draft_node_id') # Có thể null nếu muốn hủy map
        
        if not base_node_id:
            return Response({"error": "Thiếu base_node_id"}, status=400)
            
        with transaction.atomic():
            # 1. Xóa map cũ của base_node này
            ComparisonMapping.objects.filter(version=version, base_node_id=base_node_id).delete()
            
            if draft_node_id:
                # 2. Đảm bảo ánh xạ 1:1: Một node dự thảo chỉ được map vào DUY NHẤT một node gốc
                # Xóa các ánh xạ khác đang trỏ tới draft_node_id này (nếu có)
                ComparisonMapping.objects.filter(version=version, draft_node_id=draft_node_id).delete()
                
                ComparisonMapping.objects.create(
                    project=version.project,
                    version=version,
                    base_node_id=base_node_id,
                    draft_node_id=draft_node_id
                )
        
        return Response({"message": "Cập nhật ánh xạ thành công!"})

    @action(detail=True, methods=['get'])
    def nodes(self, request, pk=None):
        """Lấy tất cả node của phiên bản dự thảo này"""
        version = self.get_object()
        nodes = ComparisonNode.objects.filter(version=version).order_by('order_index')
        serializer = ComparisonNodeSerializer(nodes, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def ai_check_references(self, request, pk=None):
        """Action: Rà soát dẫn chiếu chéo sử dụng AI"""
        version = self.get_object()
        # Lấy toàn bộ text dự thảo
        nodes = ComparisonNode.objects.filter(version=version).order_by('order_index')
        text_parts = []
        for n in nodes:
            text_parts.append(f"{n.node_label}: {n.content}")
        full_text = "\n".join(text_parts)

        ai_service = UnifiedAIService()
        result_json = ai_service.check_internal_references(full_text)
        
        if "error" in result_json:
            return Response(result_json, status=500)

        # Lưu kết quả
        ai_res = ComparisonAIResult.objects.create(
            version=version,
            result_type='reference_check',
            content=json.dumps(result_json, ensure_ascii=False),
            agent_info=f"{ai_service.provider} ({ai_service.model_name or 'default'})"
        )
        return Response(ComparisonAIResultSerializer(ai_res).data)

    @action(detail=True, methods=['post'])
    def ai_generate_report(self, request, pk=None):
        """Action: Đối chiếu & Tạo báo cáo tự động sử dụng AI"""
        version = self.get_object()
        summary_text = request.data.get('summary_text', '')
        custom_request = request.data.get('custom_request', 'Tạo báo cáo tóm tắt nội dung mới')

        # Lấy toàn bộ text dự thảo
        nodes = ComparisonNode.objects.filter(version=version).order_by('order_index')
        full_text = "\n".join([f"{n.node_label}: {n.content}" for n in nodes])

        ai_service = UnifiedAIService()
        report_md = ai_service.generate_automated_report(summary_text, full_text, custom_request)

        if report_md.startswith("Lỗi"):
            return Response({"error": report_md}, status=500)

        # Lưu kết quả
        ai_res = ComparisonAIResult.objects.create(
            version=version,
            result_type='automated_report',
            content=report_md,
            agent_info=f"{ai_service.provider} ({ai_service.model_name or 'default'})"
        )
        return Response(ComparisonAIResultSerializer(ai_res).data)

    @action(detail=True, methods=['post'])
    def upload_explanation(self, request, pk=None):
        """Action: Nạp thuyết minh từ file Word (.docx)"""
        version = self.get_object()
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "Vui lòng tải lên tệp thuyết minh (.docx)."}, status=400)
            
        from .utils.parser_engine import ExplanationParser
        try:
            parser = ExplanationParser(file_obj)
            exp_dict = parser.parse_to_dict()
            
            # Cập nhật thuyết minh cho các Điều tương ứng
            count = 0
            # Chỉ khớp với type='Điều' như yêu cầu
            nodes = ComparisonNode.objects.filter(version=version, node_type='Điều')
            for node in nodes:
                # Trích xuất số điều từ nhãn (Ví dụ: "Điều 1" -> "1")
                m = re.search(r'\d+', node.node_label)
                if m:
                    num = m.group()
                    if num in exp_dict:
                        node.explanation = exp_dict[num]
                        node.save()
                        count += 1
            
            return Response({"message": f"Đã cập nhật thuyết minh cho {count} Điều từ file Word."})
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    @action(detail=True, methods=['post'])
    def save_gsheet_url(self, request, pk=None):
        """Action: Chỉ lưu link Google Sheet không đồng bộ"""
        version = self.get_object()
        sheet_url = request.data.get('sheet_url')
        if not sheet_url:
            return Response({"error": "Vui lòng nhập đường dẫn Google Sheet."}, status=400)
        
        version.explanation_sheet_url = sheet_url
        version.save()
        return Response({"message": "Đã lưu đường dẫn Google Sheet thành công.", "sheet_url": sheet_url})

    @action(detail=True, methods=['post'])
    def sync_gsheet_explanation(self, request, pk=None):
        """Action: Đồng bộ thuyết minh từ Google Sheets"""
        version = self.get_object()
        sheet_url = request.data.get('sheet_url')
        if not sheet_url:
            return Response({"error": "Vui lòng nhập đường dẫn Google Sheet."}, status=400)
            
        try:
            count = self._perform_gsheet_sync(version, sheet_url)
            
            # Lưu lại link GSheet vào version nếu chưa có hoặc thay đổi
            if version.explanation_sheet_url != sheet_url:
                version.explanation_sheet_url = sheet_url
                version.save()
                
            return Response({"message": f"Đã đồng bộ thuyết minh cho {count} mục từ Google Sheet."})
        except Exception as e:
            return Response({"error": str(e)}, status=400)

    def _normalize_text_for_compare(self, text):
        """Chuẩn hóa văn bản để so sánh: loại bỏ HTML, gộp khoảng trắng thừa, xử lý placeholder trống"""
        if not text: return ""
        import unicodedata
        # 0. Chuẩn hóa NFC để đảm bảo tiếng Việt giống nhau trên mọi OS (Quan trọng: Windows/Mac vs Linux)
        text = unicodedata.normalize('NFC', str(text))
        
        # 1. Loại bỏ tag HTML
        text = re.sub(r'<[^>]+>', ' ', text)
        
        # 2. Quy đổi whitespace (newline, tab, nhiều space, non-breaking space) thành 1 space duy nhất
        text = re.sub(r'[\s\xa0\u200b\ufeff]+', ' ', text)
        text = text.strip()

        # 3. Chuẩn hóa các cụm từ tương đương với "Trống" hoặc "Mới"
        # Giúp tránh báo lỗi khi bên hệ thống là null/trống còn GSheet điền text ghi chú
        placeholders = [
            '(trống)', 'trống', '(trong)', 'trong',
            '(dòng mới)', 'dòng mới', '(dong moi)', 'dong moi',
            '(mới)', 'mới', '(moi)', 'moi',
            '(bãi bỏ)', 'bãi bỏ', '(bai bo)', 'bai bo',
            'x', '-', '...', 'none', 'null'
        ]
        if text.lower() in placeholders:
            return ""

        return text

    @action(detail=True, methods=['get'])
    def gsheet_compare_explanation(self, request, pk=None):
        """So sánh dữ liệu 3 cột (Gốc, Dự thảo, Thuyết minh) giữa Hệ thống và Google Sheet"""
        version = self.get_object()
        
        # Ưu tiên lấy URL từ param gửi lên để phản hồi tức thì khi đổi link
        sheet_url = request.query_params.get('url') or version.explanation_sheet_url
        
        if not sheet_url:
            return Response({"error": "Chưa cài đặt URL Google Sheet Thuyết minh."}, status=400)
            
        from .utils.gsheet_sync import sync_explanation_from_gsheet, extract_norm_label, get_content_fingerprint
        try:
            # Lấy dữ liệu gsheet (đã bao gồm ID, Label và Fingerprint keys)
            gsheet_rows = sync_explanation_from_gsheet(sheet_url)
            
            # Sử dụng logic trộn hàng chuẩn của hệ thống
            rows = self._get_interleaved_rows(version)
            
            comparison = []
            for r in rows:
                b_node = r.get("base_node")
                d_node = r.get("draft_node")
                
                # ID tham chiếu: node_{id}
                primary_node = d_node if d_node else b_node
                if not primary_node: continue
                
                row_id = f"node_{primary_node.get('id')}"
                label = primary_node.get("node_label")
                label_key = extract_norm_label(label)
                fp_key = f"fp_{get_content_fingerprint(primary_node.get('content'))}"
                
                # Dữ liệu trong hệ thống
                db_base = f"{b_node.get('node_label')}\n{b_node.get('content')}".strip() if b_node else ""
                db_draft = f"{d_node.get('node_label')}\n{d_node.get('content')}".strip() if d_node else ""
                db_exp = r.get("display_explanation") or ""
                
                # Dữ liệu trên GSheet - Đối soát theo Cặp Mapping (ID -> Pair Key -> Fallback)
                gs_data = gsheet_rows.get(row_id)
                
                if not gs_data:
                    # Tạo Pair-Key từ hệ thống để đối soát GSheet
                    b_lab_norm = extract_norm_label(b_node.get('node_label')) if b_node else ""
                    d_lab_norm = extract_norm_label(d_node.get('node_label')) if d_node else ""
                    pair_key = f"{b_lab_norm}|{d_lab_norm}"
                    gs_data = gsheet_rows.get(pair_key)
                
                if not gs_data:
                    # Fallback 1: Thử tìm theo Draft Label đơn lẻ
                    d_lab_norm = extract_norm_label(d_node.get('node_label')) if d_node else ""
                    if d_lab_norm:
                        gs_data = gsheet_rows.get(f"only_draft_{d_lab_norm}")
                
                if not gs_data:
                    # Fallback 2: Thử tìm theo Fingerprint của Draft (Chính xác cao cho nội dung)
                    if d_node:
                        fp_d = f"fp_d_{get_content_fingerprint(d_node.get('content'))}"
                        gs_data = gsheet_rows.get(fp_d)
                
                if not gs_data:
                    # Fallback 3: Thử tìm theo Base Label đơn lẻ
                    b_lab_norm = extract_norm_label(b_node.get('node_label')) if b_node else ""
                    if b_lab_norm:
                        gs_data = gsheet_rows.get(f"only_base_{b_lab_norm}")

                if not gs_data:
                    # Fallback 4: Thử tìm theo Fingerprint của Base (Cho các dòng bãi bỏ hoặc giữ nguyên)
                    if b_node:
                        fp_b = f"fp_b_{get_content_fingerprint(b_node.get('content'))}"
                        gs_data = gsheet_rows.get(fp_b)

                if not gs_data: gs_data = {}
                
                gs_base = gs_data.get('base', '')
                gs_draft = gs_data.get('draft', '')
                gs_exp = gs_data.get('exp', '')
                
                status_val = "match"
                if not gs_data:
                    status_val = "missing_gsheet"
                else:
                    # So sánh thông minh sau khi chuẩn hóa
                    db_exp_norm = self._normalize_text_for_compare(db_exp)
                    gs_exp_norm = self._normalize_text_for_compare(gs_exp)
                    
                    db_base_norm = self._normalize_text_for_compare(db_base)
                    gs_base_norm = self._normalize_text_for_compare(gs_base)
                    
                    db_draft_norm = self._normalize_text_for_compare(db_draft)
                    gs_draft_norm = self._normalize_text_for_compare(gs_draft)

                    if db_exp_norm != gs_exp_norm:
                        status_val = "mismatch"
                    elif db_base_norm != gs_base_norm or db_draft_norm != gs_draft_norm:
                        status_val = "mismatch"
                
                comparison.append({
                    "id": primary_node.get("id"),
                    "row_id": row_id, # ID tham chiếu cho GSheet
                    "label": label,
                    "db_content": db_exp,
                    "gsheet_content": gs_exp,
                    "db_data": {"base": db_base, "draft": db_draft, "exp": db_exp},
                    "gs_data": {"base": gs_base, "draft": gs_draft, "exp": gs_exp},
                    "status": status_val
                })
            
            return Response(comparison)
        except Exception as e:
            return Response({"error": str(e)}, status=400)

    @action(detail=True, methods=['post'])
    def gsheet_sync_selected_explanation(self, request, pk=None):
        """Kéo dữ liệu từ GSheet về những node được chọn"""
        version = self.get_object()
        items = request.data.get('items', []) # [{id, content}]
        
        updated_count = 0
        with transaction.atomic():
            project = version.project
            for item in items:
                # Cập nhật node: có thể là base node (version=None) hoặc draft node (version=version)
                # Phải thuộc Version này hoặc Project này
                res = ComparisonNode.objects.filter(id=item['id']).filter(
                    models.Q(version=version) | models.Q(project=version.project)
                ).update(explanation=item['content'])
                updated_count += res
                
        return Response({
            "message": f"Đã cập nhật thành công {updated_count}/{len(items)} mục từ Google Sheet.",
            "updated_count": updated_count
        })

    @action(detail=True, methods=['post'])
    def gsheet_push_selected_explanation(self, request, pk=None):
        """Đẩy dữ liệu từ hệ thống lên GSheet cho những node được chọn - Hỗ trợ chuẩn 4 cột"""
        try:
            version = self.get_object()
            if not version.explanation_sheet_url:
                return Response({"error": "Chưa cài đặt URL Google Sheet Thuyết minh."}, status=400)
                
            selected_ids = request.data.get('node_ids', [])
            if not selected_ids:
                return Response({"error": "Vui lòng chọn ít nhất một mục để đẩy."}, status=400)
            
            # Sử dụng logic trộn hàng để có đầy đủ nội dung Gốc/Dự thảo/Thuyết minh
            all_rows = self._get_interleaved_rows(version)
            
            items_to_push = []
            for r in all_rows:
                b_n = r.get("base_node")
                d_n = r.get("draft_node")
                primary = d_n if d_n else b_n
                
                if not primary or primary.get("id") not in selected_ids:
                    continue
                
                # Xây dựng item đẩy lên GSheet
                items_to_push.append({
                    'id': f"node_{primary.get('id')}",
                    'label': primary.get("node_label"),
                    'base_content': f"{b_n.get('node_label')}\n{b_n.get('content')}".strip() if b_n else "",
                    'draft_content': f"{d_n.get('node_label')}\n{d_n.get('content')}".strip() if d_n else "",
                    'explanation': r.get("display_explanation") or ""
                })
            
            if not items_to_push:
                return Response({"error": "Không tìm thấy dữ liệu hợp lệ để đẩy."}, status=400)

            from .utils.gsheet_sync import push_explanations_to_gsheet
            push_explanations_to_gsheet(version.explanation_sheet_url, items_to_push)
            
            return Response({"message": f"Đã đẩy thành công {len(items_to_push)} mục lên Google Sheet (Cấu trúc 4 cột)."})
        except Exception as e:
            return Response({"error": f"Lỗi đồng bộ: {str(e)}"}, status=500)


    @action(detail=True, methods=['get'])
    def export_mappings(self, request, pk=None):
        """Xuất bộ nhớ ánh xạ (Mapping Memory) ra file Excel"""
        import pandas as pd
        import io
        from django.http import HttpResponse
        from rest_framework_simplejwt.authentication import JWTAuthentication

        # Hỗ trợ xác thực qua query param để window.open hoạt động
        if not request.user.is_authenticated:
            query_token = request.query_params.get('token')
            if query_token:
                try:
                    auth = JWTAuthentication()
                    validated_token = auth.get_validated_token(query_token)
                    user = auth.get_user(validated_token)
                    request.user = user
                except Exception as e:
                    print(f"Query token auth error: {e}")

        if not request.user.is_authenticated:
               return Response({"detail": "Authentication credentials were not provided."}, status=401)

        version = self.get_object()
        mappings = ComparisonMapping.objects.filter(version=version).select_related('base_node', 'draft_node')
        
        data = []
        for m in mappings:
            data.append({
                'Loại mục': m.base_node.node_type,
                'Nhãn Gốc': m.base_node.node_label,
                'Nội dung Gốc': m.base_node.content[:200] + '...' if len(m.base_node.content) > 200 else m.base_node.content,
                'Nhãn Dự thảo': m.draft_node.node_label,
                'Nội dung Dự thảo': m.draft_node.content[:200] + '...' if len(m.draft_node.content) > 200 else m.draft_node.content,
            })
            
        if not data:
            return Response({"error": "Chưa có dữ liệu ánh xạ để xuất."}, status=400)

        df = pd.DataFrame(data)
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False, sheet_name='Mapping Memory')
        
        output.seek(0)
        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="Mapping_Memory_{version.id}.xlsx"'
        return response

    @action(detail=False, methods=['get'], url_path='export_ai_report/(?P<result_id>[^/.]+)')
    def export_ai_report(self, request, result_id=None):
        """Xuất báo cáo AI sang file Word (.docx)"""
        import io
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        ai_res = ComparisonAIResult.objects.get(id=result_id)
        if ai_res.result_type != 'automated_report':
            return Response({"error": "Chỉ hỗ trợ xuất báo cáo tự động."}, status=400)

        doc = Document()
        # Chuyển đổi Markdown đơn giản sang docx
        lines = ai_res.content.split('\n')
        for line in lines:
            if line.startswith('###'):
                p = doc.add_heading(line.replace('###', '').strip(), level=3)
            elif line.startswith('##'):
                p = doc.add_heading(line.replace('##', '').strip(), level=2)
            elif line.startswith('#'):
                p = doc.add_heading(line.replace('#', '').strip(), level=1)
            elif line.startswith('*') or line.startswith('-'):
                p = doc.add_paragraph(line[1:].strip(), style='List Bullet')
            elif '|' in line and '--' not in line: # Bảng cơ bản
                cells = [c.strip() for c in line.split('|') if c.strip()]
                if cells:
                    p = doc.add_paragraph(" | ".join(cells))
            else:
                p = doc.add_paragraph(line.strip())

        stream = io.BytesIO()
        doc.save(stream)
        stream.seek(0)
        
        response = HttpResponse(
            stream.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        safe_name = f"Bao_cao_AI_{ai_res.id}.docx"
        response['Content-Disposition'] = f'attachment; filename="{safe_name}"'
        return response

import json
class ComparisonNodeViewSet(viewsets.ModelViewSet):
    queryset = ComparisonNode.objects.all().order_by('order_index')
    serializer_class = ComparisonNodeSerializer

    @action(detail=True, methods=['post'])
    def insert_node(self, request, pk=None):
        """Chèn một node mới vào phía trên hoặc phía dưới node hiện tại"""
        target_node = self.get_object()
        position = request.data.get('position', 'below') # 'above' or 'below'
        node_label = request.data.get('node_label', 'Mục mới')
        content = request.data.get('content', '')
        node_type = request.data.get('node_type', 'Điều')
        
        # Xác định phạm vi (scope)
        scope_filter = {
            'project': target_node.project,
            'version': target_node.version,
            'standalone_review': target_node.standalone_review,
            'parent': target_node.parent
        }
        
        with transaction.atomic():
            if position == 'above':
                new_order = target_node.order_index
            else: # 'below'
                new_order = target_node.order_index + 1
                
            # Đẩy các node phía sau lên
            ComparisonNode.objects.filter(
                **scope_filter,
                order_index__gte=new_order
            ).update(order_index=models.F('order_index') + 1)
            
            # Tạo node mới
            new_node = ComparisonNode.objects.create(
                **scope_filter,
                node_type=node_type,
                node_label=node_label,
                content=content,
                order_index=new_order
            )
            
        return Response(ComparisonNodeSerializer(new_node).data)

class StandaloneReviewViewSet(viewsets.ModelViewSet):
    queryset = StandaloneReview.objects.all().order_by('-created_at')
    serializer_class = StandaloneReviewSerializer

    def perform_create(self, serializer):
        review = serializer.save(uploaded_by=self.request.user)
        # Tự động bóc tách sau khi tải lên
        # Sử dụng đúng constructor và method của ComparisonParser
        parser = ComparisonParser(review.file.path)
        structure = parser.parse()
        
        def save_nodes_recursive(node_list, parent=None):
            for node_data in node_list:
                children = node_data.pop('children', [])
                node = ComparisonNode.objects.create(
                    standalone_review=review,
                    parent=parent,
                    **node_data
                )
                if children:
                    save_nodes_recursive(children, parent=node)
        
        with transaction.atomic():
            save_nodes_recursive(structure)

    @action(detail=True, methods=['get'])
    def data(self, request, pk=None):
        review = self.get_object()
        nodes = ComparisonNode.objects.filter(standalone_review=review).order_by('order_index')
        ai_results = ComparisonAIResult.objects.filter(standalone_review=review)
        
        return Response({
            "review": StandaloneReviewSerializer(review).data,
            "nodes": ComparisonNodeSerializer(nodes, many=True).data,
            "ai_results": ComparisonAIResultSerializer(ai_results, many=True).data
        })

    @action(detail=True, methods=['post'])
    def ai_check(self, request, pk=None):
        review = self.get_object()
        nodes = ComparisonNode.objects.filter(standalone_review=review).order_by('order_index')
        full_text = "\n".join([f"{n.node_label}: {n.content}" for n in nodes])

        ai_service = UnifiedAIService()
        result_json = ai_service.check_internal_references(full_text)
        
        if "error" in result_json:
            error_msg = result_json.get("error", "Lỗi không xác định từ AI")
            return Response({"error": error_msg}, status=500)

        ai_res = ComparisonAIResult.objects.create(
            standalone_review=review,
            result_type='reference_check',
            content=json.dumps(result_json, ensure_ascii=False),
            agent_info=f"{ai_service.provider} ({ai_service.model_name or 'default'})"
        )
        return Response(ComparisonAIResultSerializer(ai_res).data)

    @action(detail=True, methods=['get'])
    def export_excel(self, request, pk=None):
        review = self.get_object()
        ai_result = ComparisonAIResult.objects.filter(standalone_review=review, result_type='reference_check').first()
        if not ai_result:
            return Response({"error": "Chưa có kết quả rà soát AI để xuất."}, status=400)
            
        output = export_reference_excel(review, ai_result)
        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
        )
        response['Content-Disposition'] = f'attachment; filename="Loi_dan_chieu_{review.id}.xlsx"'
        return response

    @action(detail=True, methods=['get'])
    def export_word(self, request, pk=None):
        review = self.get_object()
        ai_result = ComparisonAIResult.objects.filter(standalone_review=review, result_type='reference_check').first()
        if not ai_result:
            return Response({"error": "Chưa có kết quả rà soát AI để xuất."}, status=400)
            
        output = export_reference_word(review, ai_result)
        response = HttpResponse(
            output.read(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = f'attachment; filename="Bao_cao_dan_chieu_{review.id}.docx"'
        return response
    permission_classes = [permissions.IsAuthenticated]
