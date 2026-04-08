from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml import OxmlElement, ns
import io
import re

def create_element(name):
    return OxmlElement(name)

def create_attribute(element, name, value):
    element.set(ns.qn(name), value)

def add_page_number(run):
    fldChar1 = create_element('w:fldChar')
    create_attribute(fldChar1, 'w:fldCharType', 'begin')

    instrText = create_element('w:instrText')
    create_attribute(instrText, 'xml:space', 'preserve')
    instrText.text = "PAGE"

    fldChar2 = create_element('w:fldChar')
    create_attribute(fldChar2, 'w:fldCharType', 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)

def export_comparison_table(project_name, version_label, rows, base_name=None, draft_name=None):
    doc = Document()
    
    # Thiết lập trang nằm ngang (Landscape)
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    # Thêm số trang vào Footer
    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_p.add_run()
    footer_run.font.name = 'Times New Roman'
    footer_run.font.size = Pt(14)
    add_page_number(footer_run)

    # Cài đặt font chữ mặc định (Times New Roman)
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(13)

    # 1. Header hành chính (Bảng ẩn không viền)
    header_table = doc.add_table(rows=2, cols=2)
    header_table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Ô bên trái: Tên cơ quan
    c1 = header_table.rows[0].cells[0]
    p1 = c1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.add_run("BỘ XÂY DỰNG\n").bold = True
    p1.add_run("-------")
    
    # Ô bên phải: Quốc hiệu
    c2 = header_table.rows[0].cells[1]
    p2 = c2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\n").bold = True
    p2.add_run("Độc lập - Tự do - Hạnh phúc\n").bold = True
    p2.add_run("---------------")
    
    # Dòng ngày tháng
    c4 = header_table.rows[1].cells[1]
    p4 = c4.paragraphs[0]
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.add_run("Hà Nội, ngày ... tháng ... năm ...")
    
    doc.add_paragraph() # Dòng trống

    # 2. Tiêu đề chính
    b_display = base_name if base_name else "VĂN BẢN GỐC"
    d_display = draft_name if draft_name else "VĂN BẢN DỰ THẢO"
    
    title_text = f"BẢN SO SÁNH, THUYẾT MINH DỰ THẢO {d_display.upper()} VỚI {b_display.upper()}\n(Kèm theo Công văn số: .../... ngày ... tháng ... năm ...)"
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run(title_text)
    run_title.bold = True
    run_title.font.size = Pt(14)

    doc.add_paragraph() # Dòng trống

    # 3. Tạo bảng so sánh chính (3 cột)
    table = doc.add_table(rows=1, cols=3)
    table.style = 'Table Grid'
    
    # Định nghĩa độ rộng cột (Tổng ngang ~ 9.5 inches cho khổ Landscape A4)
    widths = [Inches(3.3), Inches(3.3), Inches(2.9)]
    
    # Dòng tiêu đề bảng
    hdr_cells = table.rows[0].cells
    headers = [b_display.upper(), d_display.upper(), 'THUYẾT MINH']
    
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].width = widths[i]
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    # 4. Nội dung các dòng
    for row_data in rows:
        row = table.add_row().cells
        for i in range(3):
            row[i].width = widths[i]

        # Cột 1: Văn bản Gốc
        b_node = row_data.get('base_node')
        if b_node:
            p_b = row[0].paragraphs[0]
            p_b.add_run(f"{b_node.get('node_label', '')}\n").bold = True
            p_b.add_run(b_node.get('content', ''))
        else:
            p_b = row[0].paragraphs[0]
            p_b.add_run("(Dòng mới)").italic = True
        
        # Cột 2: Dự thảo
        d_node = row_data.get('draft_node')
        if d_node:
            p_d = row[1].paragraphs[0]
            p_d.add_run(f"{d_node.get('node_label', '')}\n").bold = True
            
            diff_html = row_data.get('diff_content', d_node.get('content', ''))
            parts = re.split(r'(<i>.*?</i>)', diff_html, flags=re.DOTALL)
            for part in parts:
                if part.startswith('<i>') and part.endswith('</i>'):
                    clean_text = part[3:-4]
                    run = p_d.add_run(clean_text)
                    run.italic = True
                else:
                    p_d.add_run(part)
        else:
            p_d = row[1].paragraphs[0]
            p_d.add_run("(Bãi bỏ)").italic = True
            
        # Cột 3: Thuyết minh
        row[2].text = ""

    # 5. Phần ghi chú cuối trang
    doc.add_paragraph()
    note_p = doc.add_paragraph()
    note_p.add_run("Ghi chú:").bold = True
    doc.add_paragraph("(1) Tên cơ quan chủ trì soạn thảo văn bản quy phạm pháp luật.")
    doc.add_paragraph("(2) Địa danh nơi cơ quan, tổ chức chủ trì soạn thảo văn bản đóng trụ sở.")
    doc.add_paragraph("(3) Tên dự thảo văn bản quy phạm pháp luật.")

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream
