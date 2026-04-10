import docx
import re

class ComparisonParser:
    def __init__(self, file_path_or_obj):
        self.doc = docx.Document(file_path_or_obj)
        self.nodes = []
        self.current_part = None
        self.current_chapter = None
        self.current_section = None
        self.current_subsection = None
        self.current_article = None
        self.current_clause = None
        self.index = 0

    def parse(self):
        """
        Bóc tách cấu trúc văn bản pháp luật: Phần -> Chương -> Mục -> Tiểu mục -> Điều -> Khoản -> Điểm.
        """
        for para in self.doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # 1. Dừng hoàn toàn nếu gặp "Phụ lục"
            if re.match(r'^Phụ\s+lục', text, re.IGNORECASE):
                break

            # 2. Loại bỏ Footnote
            if re.match(r'^\[\d+\]', text) or (len(text) < 110 and re.match(r'^\d+\.?\s', text) and not self.current_article):
                if not re.match(r'^Điều\s+\d+', text, re.IGNORECASE): 
                    continue

            # 3. Nhận diện Phần
            part_match = re.match(r'^Phần\s+([IVXLCDM\d]+|[A-ZĐĂÂÊÔƠƯa-zđăâêôơư\s]+)[:\.]?\s*(.*)', text, re.IGNORECASE)
            if part_match:
                label = f"Phần {part_match.group(1).strip()}"
                node = self._create_node('Phần', label, part_match.group(2).strip())
                self.nodes.append(node)
                self.current_part = node
                self.current_chapter = None
                self.current_section = None
                self.current_subsection = None
                self.current_article = None
                self.current_clause = None
                continue

            # 4. Nhận diện Chương
            chapter_match = re.match(r'^Chương\s+([IVXLCDM\d]+|[A-ZĐĂÂÊÔƠƯa-zđăâêôơư\s]+)[:\.]?\s*(.*)', text, re.IGNORECASE)
            if chapter_match:
                label = f"Chương {chapter_match.group(1).strip()}"
                node = self._create_node('Chương', label, chapter_match.group(2).strip())
                if self.current_part:
                    self.current_part['children'].append(node)
                else:
                    self.nodes.append(node)
                self.current_chapter = node
                self.current_section = None
                self.current_subsection = None
                self.current_article = None
                self.current_clause = None
                continue

            # 5. Nhận diện Mục
            section_match = re.match(r'^Mục\s+(\d+|[IVXLCDM]+|[A-ZĐĂÂÊÔƠƯa-zđăâêôơư\s]+)[:\.]?\s*(.*)', text, re.IGNORECASE)
            if section_match:
                label = f"Mục {section_match.group(1).strip()}"
                node = self._create_node('Mục', label, section_match.group(2).strip())
                if self.current_chapter:
                    self.current_chapter['children'].append(node)
                elif self.current_part:
                    self.current_part['children'].append(node)
                else:
                    self.nodes.append(node)
                self.current_section = node
                self.current_subsection = None
                self.current_article = None
                self.current_clause = None
                continue
            
            # 6. Nhận diện Tiểu mục
            sub_section_match = re.match(r'^Tiểu\s+mục\s+(\d+|[IVXLCDM]+|[A-ZĐĂÂÊÔƠƯa-zđăâêôơư\s]+)[:\.]?\s*(.*)', text, re.IGNORECASE)
            if sub_section_match:
                label = f"Tiểu mục {sub_section_match.group(1).strip()}"
                node = self._create_node('Tiểu mục', label, sub_section_match.group(2).strip())
                if self.current_section:
                    self.current_section['children'].append(node)
                elif self.current_chapter:
                    self.current_chapter['children'].append(node)
                else:
                    self.nodes.append(node)
                self.current_subsection = node
                self.current_article = None
                self.current_clause = None
                continue

            # 7. Nhận diện Điều
            article_match = re.match(r'^Điều\s+([\d\.\-]+[a-z]?)\s*[\.:-]?\s*(.*)', text, re.IGNORECASE)
            if article_match:
                number = article_match.group(1).strip()
                if number.endswith('.'): number = number[:-1]
                title = article_match.group(2).strip()
                label = f"Điều {number}. {title}" if title else f"Điều {number}"
                node = self._create_node('Điều', label, "")
                
                # Tìm cha phù hợp nhất
                if self.current_subsection:
                    self.current_subsection['children'].append(node)
                elif self.current_section:
                    self.current_section['children'].append(node)
                elif self.current_chapter:
                    self.current_chapter['children'].append(node)
                elif self.current_part:
                    self.current_part['children'].append(node)
                else:
                    self.nodes.append(node)
                
                self.current_article = node
                self.current_clause = None
                continue

            # 8. Nhận diện Khoản
            clause_match = re.match(r'^(?:Khoản\s+)?(\d+)[\.\:\s]*\s*(.*)', text, re.IGNORECASE)
            if clause_match and self.current_article:
                node = self._create_node('Khoản', f"{clause_match.group(1)}.", clause_match.group(2).strip())
                self.current_article['children'].append(node)
                self.current_clause = node
                continue

            # 9. Nhận diện Điểm
            point_match = re.match(r'^(?:Điểm\s+)?([a-zđ])[\)\.\:\s]*\s*(.*)', text, re.IGNORECASE)
            if point_match and self.current_clause:
                node = self._create_node('Điểm', f"{point_match.group(1)})", point_match.group(2).strip())
                self.current_clause['children'].append(node)
                continue

            # 10. Nội dung cộng dồn
            if self.current_clause:
                self.current_clause['content'] += f"\n{text}"
            elif self.current_article:
                self.current_article['content'] += f"\n{text}"
            elif self.current_subsection:
                self.current_subsection['content'] += f"\n{text}"
            elif self.current_section:
                self.current_section['content'] += f"\n{text}"
            elif self.current_chapter:
                self.current_chapter['content'] += f"\n{text}"
            elif self.current_part:
                self.current_part['content'] += f"\n{text}"
            else:
                if self.nodes and self.nodes[-1]['node_type'] == 'Vấn đề khác':
                    self.nodes[-1]['content'] += f"\n{text}"
                else:
                    node = self._create_node('Vấn đề khác', 'Phần mở đầu', text)
                    self.nodes.append(node)

        return self.nodes

    def _is_bold(self, para):
        """Kiểm tra xem toàn bộ đoạn văn có được in đậm không"""
        if para.runs:
            # Chỉ kiểm tra các run có chứa text thực sự
            real_runs = [r for r in para.runs if r.text.strip()]
            if real_runs and all(r.bold for r in real_runs):
                return True
        return False

    def _create_node(self, node_type, label, content):
        node = {
            'node_type': node_type,
            'node_label': label,
            'content': content,
            'children': [],
            'order_index': self.index
        }
        self.index += 1
        return node

class ExplanationParser:
    def __init__(self, file_path_or_obj):
        try:
            self.doc = docx.Document(file_path_or_obj)
        except Exception as e:
            raise Exception(f"Lỗi khi đọc file Word: {str(e)}")

    def parse_to_dict(self):
        """
        Bóc tách các bảng trong Word: Cột 1 là số Điều, Cột 2 là Thuyết minh.
        Trả về dictionary { số_điều_chuỗi: nội_dung_thuyết_minh }
        """
        results = {}
        for table in self.doc.tables:
            for row in table.rows:
                # Bảng thuyết minh phải có ít nhất 2 cột
                if len(row.cells) >= 2:
                    c1 = row.cells[0].text.strip()
                    c2 = row.cells[1].text.strip()
                    
                    if not c1:
                        continue
                        
                    # Tìm số Điều (Ví dụ: "Điều 1", "Điều 1." -> "1")
                    match = re.search(r'[\u0110\u0111]i\u1ec1u\s+(\d+)', c1, re.IGNORECASE)
                    if match:
                        article_num = match.group(1)
                        # Nếu đã có dữ liệu rồi thì cộng dồn (trường hợp 1 điều chia nhiều hàng)
                        if article_num in results:
                            results[article_num] += f"\n{c2}"
                        else:
                            results[article_num] = c2
        return results
