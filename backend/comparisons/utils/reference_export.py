import pandas as pd
import io
import json
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from django.http import HttpResponse

def export_reference_excel(review, ai_result):
    """Xuất danh sách lỗi dẫn chiếu ra file Excel"""
    try:
        data = json.loads(ai_result.content)
        refs = data.get('references', [])
    except:
        refs = []

    df = pd.DataFrame(refs)
    
    # Việt hóa tiêu đề cột
    column_mapping = {
        'source_location': 'Vị trí nguồn',
        'extracted_text': 'Nội dung trích xuất',
        'target_location': 'Vị trí đích',
        'status': 'Trạng thái',
        'reason': 'Lý do/Đánh giá',
        'suggestion': 'Gợi ý sửa đổi'
    }
    df = df.rename(columns=column_mapping)
    
    # Đảm bảo các cột hiển thị theo thứ tự đẹp
    cols = ['Vị trí nguồn', 'Nội dung trích xuất', 'Vị trí đích', 'Trạng thái', 'Lý do/Đánh giá', 'Gợi ý sửa đổi']
    df = df.reindex(columns=[c for c in cols if c in df.columns])

    output = io.BytesIO()
    with pd.ExcelWriter(output, engine='openpyxl') as writer:
        df.to_excel(writer, index=False, sheet_name='Danh sách lỗi dẫn chiếu')
        
        # Auto-adjust columns width
        worksheet = writer.sheets['Danh sách lỗi dẫn chiếu']
        for idx, col in enumerate(df.columns):
            max_len = max(df[col].astype(str).map(len).max(), len(col)) + 2
            worksheet.column_dimensions[chr(65 + idx)].width = min(max_len, 50)

    output.seek(0)
    return output

def export_reference_word(review, ai_result):
    """Xuất báo cáo đánh giá dẫn chiếu ra file Word"""
    try:
        data = json.loads(ai_result.content)
        refs = data.get('references', [])
    except:
        refs = []

    doc = Document()
    
    # Tiêu đề báo cáo
    title = doc.add_heading('BÁO CÁO RÀ SOÁT DẪN CHIẾU PHÁP LÝ', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Thông tin chung
    p = doc.add_paragraph()
    p.add_run(f'Tên văn bản: ').bold = True
    p.add_run(review.name)
    
    p = doc.add_paragraph()
    p.add_run(f'Ngày rà soát: ').bold = True
    p.add_run(ai_result.created_at.strftime('%d/%m/%Y %H:%M'))
    
    p = doc.add_paragraph()
    p.add_run(f'Công cụ rà soát: ').bold = True
    p.add_run(ai_result.agent_info)

    doc.add_heading('1. Tổng hợp kết quả', level=1)
    
    valid_count = len([r for r in refs if r.get('status') == 'Valid'])
    error_count = len([r for r in refs if r.get('status') != 'Valid'])
    
    p = doc.add_paragraph()
    p.add_run(f'- Tổng số dẫn chiếu nội bộ phát hiện: ').bold = True
    p.add_run(f'{len(refs)}')
    
    p = doc.add_paragraph()
    p.add_run(f'- Số dẫn chiếu hợp lệ: ').bold = True
    p.add_run(f'{valid_count}')
    
    p = doc.add_paragraph()
    p.add_run(f'- Số dẫn chiếu cần lưu ý/lỗi: ').bold = True
    p.add_run(f'{error_count}')

    doc.add_heading('2. Chi tiết các dẫn chiếu cần lưu ý', level=1)
    
    if error_count == 0:
        doc.add_paragraph('Không phát hiện lỗi dẫn chiếu nghiêm trọng.')
    else:
        # Tạo bảng lỗi
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Vị trí'
        hdr_cells[1].text = 'Trích đoạn'
        hdr_cells[2].text = 'Trạng thái'
        hdr_cells[3].text = 'Đánh giá & Gợi ý'
        
        for r in refs:
            if r.get('status') == 'Valid':
                continue
            row_cells = table.add_row().cells
            row_cells[0].text = r.get('source_location', '')
            row_cells[1].text = r.get('extracted_text', '')
            row_cells[2].text = r.get('status', '')
            
            eval_text = f"Lý do: {r.get('reason', '')}"
            if r.get('suggestion'):
                eval_text += f"\nGợi ý: {r.get('suggestion')}"
            row_cells[3].text = eval_text

    doc.add_heading('3. Kết luận', level=1)
    if error_count > 0:
        doc.add_paragraph('Văn bản còn tồn tại một số sai sót về kỹ thuật lập pháp trong dẫn chiếu chéo nội bộ. Đề nghị đơn vị soạn thảo rà soát và chỉnh sửa theo các gợi ý nêu trên trước khi ban hành.')
    else:
        doc.add_paragraph('Kỹ thuật lập pháp về dẫn chiếu nội bộ của văn bản đảm bảo tính logic và chính xác.')

    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output
