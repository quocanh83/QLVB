from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from django.db import models
from rest_framework.decorators import action
from .models import Feedback, Explanation, ActionLog, ConsultationResponse, FeedbackAssignment
from documents.models import Document, DocumentNode, DocumentAppendix
from django.contrib.contenttypes.models import ContentType
from .serializers import FeedbackSerializer, ConsultationResponseSerializer
import docx
import re
import os
from openai import OpenAI
from django.http import FileResponse
from .utils.v2_template_generator import generate_from_v2_template, _get_field_value
from .utils.v2_template_generator import generate_from_v2_template
from django.core.files.storage import FileSystemStorage
from django.conf import settings
import shutil
from django.utils.text import get_valid_filename
import traceback
import unicodedata
import pandas as pd
import numpy as np
import io
import requests
import json



class FeedbackViewSet(viewsets.ModelViewSet):
    queryset = Feedback.objects.all()
    serializer_class = FeedbackSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_permissions(self):
        if self.action == 'export_mau_10':
            return [permissions.AllowAny()]
        return super().get_permissions()

    def create(self, request, *args, **kwargs):
        node_id = request.data.get('node')
        appendix_id = request.data.get('appendix')
        
        document_id = None
        if node_id:
            try:
                node = DocumentNode.objects.get(id=node_id)
                document_id = node.document_id
            except DocumentNode.DoesNotExist:
                return Response({"error": "Điều/Khoản không tồn tại."}, status=404)
        elif appendix_id:
            from documents.models import DocumentAppendix
            try:
                app = DocumentAppendix.objects.get(id=appendix_id)
                document_id = app.document_id
            except DocumentAppendix.DoesNotExist:
                return Response({"error": "Phụ lục không tồn tại."}, status=404)
        
        if document_id and not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Nhập góp ý cho dự thảo này."}, status=403)
            
        return super().create(request, *args, **kwargs)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    def _check_permission(self, request, document_id, required_role=None):
        user = request.user
        # 1. Admin?
        is_admin = user.is_staff or user.is_superuser or (hasattr(user, 'roles') and user.roles.filter(role_name='Admin').exists())
        if is_admin: return True
        
        # 2. Lead?
        try:
            if not document_id: return False
            doc = Document.objects.get(id=document_id)
            if doc.leads.filter(id=user.id).exists(): return True
        except (Document.DoesNotExist, ValueError, TypeError, AttributeError):
            pass

        # 3. Chuyên viên Role?
        allowed_roles = ['Contributor', 'Explainer', 'Admin', 'Chuyên viên Góp ý', 'Chuyên viên Giải trình']
        return user.roles.filter(role_name__in=allowed_roles).exists()
        
    def _find_best_node_match(self, document_id, label, parent_context=None):
        """
        Tìm node phù hợp nhất dựa trên label (Điều 1, Khoản 2, Điểm a...)
        Hỗ trợ phân cấp: nếu label chỉ là 'Khoản 2', sẽ ưu tiên tìm Khoản 2 thuộc parent_context (Điều x).
        """
        if not label or label == "Vấn đề khác": return None
        label = label.strip()
        
        # Thử tìm khớp hoàn toàn trước (case-insensitive)
        exact_match = DocumentNode.objects.filter(document_id=document_id, node_label__iexact=label).first()
        if exact_match: return exact_match

        # Regex linh hoạt hơn để tách Điều/Khoản/Điểm
        dieu_match = re.search(r'(?:Điều|D|Art|Art\.|Chương|Mục)\s*(\d+)', label, re.IGNORECASE)
        khoan_match = re.search(r'(?:Khoản|K|Clause|K\.)\s*(\d+)', label, re.IGNORECASE)
        diem_match = re.search(r'(?:Điểm|Diem|Item|Point|P\.|p)\s*([a-zđ0-9]+)', label, re.IGNORECASE)
        
        # Nếu nhãn chỉ là số
        if not dieu_match and not khoan_match and label.replace('.', '').isdigit():
            clean_num = label.replace('.', '').strip()
            if not parent_context:
                dieu_match = re.search(r'(\d+)', clean_num)
            else:
                khoan_match = re.search(r'(\d+)', clean_num)

        # 1. Xử lý có Điều hoặc Chương/Mục
        if dieu_match:
            d_num_int = int(dieu_match.group(1))
            possible_nodes = DocumentNode.objects.filter(document_id=document_id, node_type='Điều', node_label__icontains=str(d_num_int))
            
            node_dieu = None
            for n in possible_nodes:
                actual_nums = re.findall(r'\d+', n.node_label)
                if any(int(x) == d_num_int for x in actual_nums):
                    node_dieu = n
                    break
            
            if node_dieu:
                if khoan_match:
                    k_num_int = int(khoan_match.group(1))
                    possible_khoans = DocumentNode.objects.filter(document_id=document_id, parent_id=node_dieu.id)
                    for k in possible_khoans:
                        k_actual_nums = re.findall(r'\d+', k.node_label)
                        if any(int(x) == k_num_int for x in k_actual_nums):
                            if diem_match:
                                d_label = diem_match.group(1).lower()
                                possible_diems = DocumentNode.objects.filter(document_id=document_id, parent_id=k.id, node_label__icontains=d_label)
    def _find_best_node_match(self, document_id, label):
        if not label: return None
        
        import re
        import unicodedata
        # Chuẩn hóa đầu vào sang NFC
        label = unicodedata.normalize('NFC', label).strip()
        label_lower = label.lower()
        label_upper = label.upper()
        
        # --- 1. XỬ LÝ PHỤ LỤC (Số La Mã VII, VIII... hoặc 7, 8...) ---
        # Regex tìm số La Mã hoặc số thường
        roman_pattern = r'^(?:M{0,3}(?:CM|CD|D?C{0,3})(?:XC|XL|L?X{0,3})(?:IX|IV|V?I{0,3})|\d+[a-z]?)$'
        
        app_num = ""
        is_appendix_prefix = any(kw in label_lower for kw in ['phụ lục', 'pl', 'phu luc'])
        if is_appendix_prefix:
            m = re.search(r'(?:phụ\s*lục|pl|phu\s*luc)\s*([ivxlcdm\d]+[a-z]?)', label_lower)
            if m: app_num = m.group(1).upper()
        elif re.match(roman_pattern, label_upper):
            app_num = label_upper

        if app_num:
            from documents.models import DocumentAppendix
            # Tìm trong DocumentNode
            possible_nodes = DocumentNode.objects.filter(document_id=document_id, node_label__icontains=app_num)
            for n in possible_nodes:
                db_label = unicodedata.normalize('NFC', n.node_label).lower()
                # Bóc tách phần số từ nhãn của DB (ví dụ từ "Phụ lục VII: abc" lấy ra "VII")
                db_match = re.search(r'(?:phụ\s*lục|pl|phu\s*luc)\s*([ivxlcdm\d]+[a-z]?)', db_label)
                if db_match:
                    db_extracted_num = db_match.group(1).upper()
                    # SO KHỚP TUYỆT ĐỐI: Phải giống hệt nhau về ký tự và độ dài (Chống VII khớp VIII)
                    if db_extracted_num == app_num:
                        return n
            
            # Tìm trong DocumentAppendix
            possible_apps = DocumentAppendix.objects.filter(document_id=document_id, name__icontains=app_num)
            for a in possible_apps:
                db_name = unicodedata.normalize('NFC', a.name).lower()
                db_match = re.search(r'(?:phụ\s*lục|pl|phu\s*luc)\s*([ivxlcdm\d]+[a-z]?)', db_name)
                if db_match:
                    db_extracted_num = db_match.group(1).upper()
                    if db_extracted_num == app_num:
                        return a
                elif app_num in db_name.upper(): # Fallback cho appendix không có chữ "Phụ lục"
                    # Kiểm tra ranh giới từ để tránh VII vào VIII
                    if re.search(rf'\b{re.escape(app_num.lower())}\b', db_name):
                        return a

        # --- 2. XỬ LÝ ĐIỀU / KHOẢN / ĐIỂM (Phân cấp) ---
        d_match = re.search(r'(?:điều|đ)\.?\s*(\d+)', label_lower)
        k_match = re.search(r'(?:khoản|k)\.?\s*(\d+)', label_lower)
        p_match = re.search(r'(?:điểm|đ)\.?\s*([a-zđ])\b', label_lower)
        
        target_node = None
        if d_match:
            d_num = d_match.group(1)
            # Dùng regex \b cho Điều
            d_node = DocumentNode.objects.filter(document_id=document_id, node_type='Điều').filter(node_label__regex=rf'(?i)điều\s*{d_num}(?:\b|:|\s|$)').first()
            if d_node:
                target_node = d_node
                if k_match:
                    k_num = k_match.group(1)
                    k_node = DocumentNode.objects.filter(parent=d_node, node_type='Khoản').filter(node_label__regex=rf'(?i)khoản\s*{k_num}(?:\b|:|\s|$)').first()
                    if k_node:
                        target_node = k_node
                        if p_match:
                            p_char = p_match.group(1)
                            p_node = DocumentNode.objects.filter(parent=k_node, node_type='Điểm').filter(node_label__regex=rf'(?i)điểm\s*{p_char}(?:\b|:|\s|$)').first()
                            if p_node: target_node = p_node
        
        if target_node: return target_node

        # --- 3. FALLBACK ---
        node = DocumentNode.objects.filter(document_id=document_id, node_label=label).first()
        if node: return node
        return DocumentNode.objects.filter(document_id=document_id, node_label__icontains=label).first()

    def _find_similar_agencies(self, name, limit=5):
        from core.models import Agency
        from difflib import get_close_matches
        import unicodedata
        
        def normalize(s):
            if not s: return ""
            return unicodedata.normalize('NFC', str(s)).strip()
            
        norm_input = normalize(name)
        if not norm_input or norm_input.lower() == 'nan': return []
        
        agencies = Agency.objects.all()
        agency_map = {normalize(a.name): a for a in agencies}
        
        if norm_input in agency_map:
            a = agency_map[norm_input]
            return [{"id": a.id, "name": a.name, "is_exact": True}]
        
        # Case-insensitive exact match
        agency_map_lower = {k.lower(): v for k, v in agency_map.items()}
        if norm_input.lower() in agency_map_lower:
            a = agency_map_lower[norm_input.lower()]
            return [{"id": a.id, "name": a.name, "is_exact": True}]

        matches = get_close_matches(norm_input, agency_map.keys(), n=limit, cutoff=0.5)
        return [{"id": agency_map[m].id, "name": agency_map[m].name, "is_exact": False} for m in matches]

    @action(detail=False, methods=['post'])
    def analyze_import(self, request):
        document_id = request.data.get('document_id')
        file_obj = request.FILES.get('file')
        gs_url = request.data.get('google_sheets_url')
        
        if not file_obj and not gs_url:
            return Response({"error": "Không tìm thấy tệp tải lên hoặc đường dẫn Google Sheets."}, status=400)
            
        try:
            import pandas as pd
            import io
            import requests
            import numpy as np
            
            if gs_url:
                # Regex trích xuất sheet_id và gid
                sheet_id_match = re.search(r'/d/([a-zA-Z0-9-_]+)', gs_url)
                gid_match = re.search(r'gid=([0-9]+)', gs_url)
                if not sheet_id_match:
                    return Response({"error": "Đường dẫn Google Sheets không hợp lệ."}, status=400)
                
                sheet_id = sheet_id_match.group(1)
                gid = gid_match.group(1) if gid_match else '0'
                export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=xlsx&gid={gid}"
                
                response = requests.get(export_url)
                if response.status_code != 200:
                    return Response({"error": f"Không thể tải dữ liệu từ Google Sheets (Status: {response.status_code}). Hãy đảm bảo sheet ở chế độ công khai."}, status=400)
                
                df = pd.read_excel(io.BytesIO(response.content))
            else:
                if file_obj.name.endswith('.xlsx') or file_obj.name.endswith('.xls'):
                    df = pd.read_excel(file_obj)
                else:
                    df = pd.read_csv(file_obj, encoding='utf-8-sig')
            
            # Map columns based on keywords (Scoring System)
            col_map = {}
            # Khởi tạo điểm là 0, chỉ nhận cột nếu điểm > 0 để tránh vơ nhầm cột STT
            col_scores = {'node': 0, 'agency': 0, 'content': 0, 'doc_number': 0, 'doc_date': 0, 'explanation': 0, 'reason': 0, 'note': 0}
            
            import unicodedata
            for col in df.columns:
                # Chuẩn chuẩn hóa Unicode NFC để so khớp tiếng Việt chính xác
                c = unicodedata.normalize('NFC', str(col).lower().replace('\n', ' ').strip())
                
                # 1. Logic cho Điều/Khoản/Phụ lục
                score_node = 0
                if 'điều, khoản' in c: score_node = 100
                elif any(kw in c for kw in ['điều', 'khoản', 'phụ lục']): score_node = 80
                if score_node > col_scores['node']:
                    col_map['node'] = col
                    col_scores['node'] = score_node
                
                # 2. Logic cho Chủ thể/Cơ quan góp ý
                score_agency = 0
                if 'chủ thể góp ý' in c: score_agency = 100
                elif any(kw in c for kw in ['chủ thể', 'agency']): score_agency = 80
                elif 'cơ quan' in c: score_agency = 50
                # Ghi chú: Đã loại bỏ 'bộ', 'tỉnh' để tránh nhận nhầm 'Cán bộ' hay cột địa danh
                if score_agency > col_scores['agency']:
                    col_map['agency'] = col
                    col_scores['agency'] = score_agency
                
                # 3. Logic cho Số hiệu văn bản
                score_doc = 0
                if 'số văn bản góp ý' in c: score_doc = 100
                elif any(kw in c for kw in ['số hiệu', 'số văn bản', 'số vb']): score_doc = 50
                if score_doc > col_scores['doc_number']:
                    col_map['doc_number'] = col
                    col_scores['doc_number'] = score_doc

                # 4. Logic cho Ngày
                score_date = 0
                if c.strip() == 'ngày': score_date = 100
                elif 'ngày' in c: score_date = 50
                if score_date > col_scores['doc_date']:
                    col_map['doc_date'] = col
                    col_scores['doc_date'] = score_date

                # 5. Logic cho Giải trình/Tiếp thu
                score_exp = 0
                if any(kw in c for kw in ['tiếp thu giải trình', 'giải trình, tiếp thu', 'tiếp thu, giải trình']): score_exp = 100
                elif any(kw in c for kw in ['giải trình', 'tiếp thu']): score_exp = 80
                elif any(kw in c for kw in ['xử lý', 'ý kiến', 'phương án']): score_exp = 50
                if score_exp > col_scores['explanation']:
                    col_map['explanation'] = col
                    col_scores['explanation'] = score_exp
                
                # 6. Lý do & Ghi chú
                score_reason = 0
                if c.strip() == 'lý do': score_reason = 100
                if score_reason > col_scores['reason']:
                    col_map['reason'] = col
                    col_scores['reason'] = score_reason

                if any(kw in c for kw in ['ghi chú', 'note']): col_map['note'] = col

            # Tìm cột Nội dung - CỰC KỲ QUAN TRỌNG
            for col in df.columns:
                c = unicodedata.normalize('NFC', str(col).lower().replace('\n', ' ').strip())
                if 'nội dung góp ý' in c:
                    col_map['content'] = col
                    break
                elif 'nội dung' in c and any(kw in c for kw in ['góp ý', 'tham vấn', 'phản biện']):
                    col_map['content'] = col
                    break
            if 'content' not in col_map:
                for col in df.columns:
                    c = unicodedata.normalize('NFC', str(col).lower().replace('\n', ' ').strip())
                    if 'góp ý' in c and col != col_map.get('doc_number'):
                        col_map['content'] = col
                        break
            
            # Fallback cuối cùng: Nếu vẫn không tìm thấy, lấy cột nào có chữ 'nội dung' duy nhất
            if 'content' not in col_map:
                for col in df.columns:
                    c = unicodedata.normalize('NFC', str(col).lower().replace('\n', ' ').strip())
                    if c == 'nội dung':
                        col_map['content'] = col
                        break
            
            if 'content' not in col_map:
                return Response({"error": "Không tìm thấy cột 'Nội dung góp ý' trong bảng dữ liệu."}, status=400)

            local_feedbacks = Feedback.objects.filter(document_id=document_id).prefetch_related('explanations')
            fb_map = {}
            for fb in local_feedbacks:
                norm_c = self._normalize_text(fb.content)
                key = (fb.agency_id, norm_c)
                if key not in fb_map or (not fb_map[key].explanations.exists() and fb.explanations.exists()):
                    fb_map[key] = fb

            results = []
            
            # XỬ LÝ MERGE CELLS
            import numpy as np
            columns_to_ffill = [
                col_map.get('node'), col_map.get('agency'), 
                col_map.get('doc_number'), col_map.get('doc_date')
            ]
            for c in columns_to_ffill:
                if c and c in df.columns:
                    df[c] = df[c].replace(r'^\s*$', np.nan, regex=True)
                    df[c] = df[c].ffill()
            
            for index, row in df.iterrows():
                def safe_str(val):
                    import unicodedata
                    s = str(val).strip() if pd.notnull(val) else ""
                    return unicodedata.normalize('NFC', s)
                
                raw_node = safe_str(row.get(col_map.get('node'), ''))
                raw_agency = safe_str(row.get(col_map.get('agency'), ''))
                raw_content = safe_str(row.get(col_map['content'], ''))
                raw_reason = safe_str(row.get(col_map.get('reason'), ''))
                raw_note = safe_str(row.get(col_map.get('note'), ''))
                exp_col = col_map.get('explanation')
                raw_explanation = safe_str(row.get(exp_col)) if exp_col else ''
                raw_doc_num = safe_str(row.get(col_map.get('doc_number'), ''))
                
                # Check ND 'OK' và Giải trình trống -> Skip (Giảm nhiễu cho người dùng)
                if raw_content.upper() == "OK" and (not raw_explanation or raw_explanation.upper() in ["OK", "NAN", "NONE"]):
                    continue

                val_doc_date = row.get(col_map.get('doc_date'), '')
                raw_doc_date = ""
                if pd.notnull(val_doc_date):
                    if hasattr(val_doc_date, 'strftime'):
                        raw_doc_date = val_doc_date.strftime('%d/%m/%Y')
                    else:
                        raw_doc_date = str(val_doc_date).split(' ')[0]
                
                if not raw_content or raw_content.lower() in ['nan', 'none']: continue
                
                # Khớp Cơ quan & Node
                agency_matches = self._find_similar_agencies(raw_agency)
                matched_agency = agency_matches[0] if agency_matches and agency_matches[0]['is_exact'] else None
                agency_id = matched_agency['id'] if matched_agency else None
                
                suggested_match = self._find_best_node_match(document_id, raw_node)
                
                target_node_id = None
                appendix_id = None
                display_label = raw_node
                
                if suggested_match:
                    from documents.models import DocumentAppendix
                    if isinstance(suggested_match, DocumentAppendix):
                        appendix_id = suggested_match.id
                        display_label = re.split(r'[:.\n]', suggested_match.name)[0].strip()
                    else:
                        target_node_id = suggested_match.id
                        display_label = suggested_match.node_label
                
                # --- KIỂM TRA TRÙNG LẶP (SỬ DỤNG MAPPING ĐÃ TỐI ƯU) ---
                is_duplicate = False
                duplicate_id = None
                existing_fb = None
                explanation_status = "none"
                existing_exp_content = ""
                
                if agency_id:
                    norm_raw_content = self._normalize_text(raw_content)
                    key = (agency_id, norm_raw_content)
                    
                    if key in fb_map:
                        existing_fb = fb_map[key]
                        is_duplicate = True
                        duplicate_id = existing_fb.id
                
                # 2. Logic xử lý Giải trình (Explanation)
                if raw_explanation and raw_explanation.lower() not in ['nan', 'none', '']:
                    if existing_fb:
                        exp = existing_fb.explanations.first()
                        if exp:
                            # Chuẩn hóa cả 2 bên sang NFC trước khi so sánh
                            existing_exp_content = unicodedata.normalize('NFC', exp.content).strip()
                            raw_exp_normalized = raw_explanation.strip()
                            explanation_status = "identical" if existing_exp_content == raw_exp_normalized else "conflict"
                        else:
                            explanation_status = "new"
                    else:
                        explanation_status = "new"
                else:
                    # QUY TẮC MỚI: Nếu hệ thống có nhưng GG sheet rỗng thì vẫn coi là khác biệt (conflict)
                    if existing_fb and existing_fb.explanations.exists():
                        explanation_status = "conflict"
                        existing_exp_content = unicodedata.normalize('NFC', existing_fb.explanations.first().content).strip()
                    else:
                        explanation_status = "none"

                results.append({
                    "key": f"import-{index}",
                    "original_node": raw_node,
                    "node_id": target_node_id,
                    "node_label": display_label,
                    "appendix_id": appendix_id,
                    "original_agency": raw_agency,
                    "agency_id": agency_id,
                    "agency_name": matched_agency['name'] if matched_agency else raw_agency,
                    "agency_suggestions": agency_matches,
                    "content": raw_content,
                    "reason": raw_reason if raw_reason.lower() not in ['nan', 'none'] else "",
                    "note": raw_note if raw_note.lower() not in ['nan', 'none'] else "",
                    "is_duplicate": is_duplicate,
                    "duplicate_id": duplicate_id,
                    "import_mode": "explanation_only" if (is_duplicate and explanation_status == "new") else ("skip" if is_duplicate else "add_new"),
                    "explanation_content": raw_explanation if explanation_status not in ["none", "identical"] else "",
                    "explanation_status": explanation_status,
                    "existing_explanation": existing_exp_content,
                    "explanation_import_mode": "overwrite" if explanation_status in ["new", "conflict"] else "skip",
                    "official_number": raw_doc_num if raw_doc_num.lower() not in ['nan', 'none'] else "",
                    "official_date": raw_doc_date,
                    "row_index": index + 2,
                })
                
            return Response({"rows": results})
            
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def confirm_import(self, request):
        from django.db import transaction
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Nhập góp ý cho dự thảo này."}, status=403)
            
        rows = request.data.get('rows', [])
        created_count = 0
        updated_count = 0
        skipped_count = 0
        exp_count = 0
        consultation_count = 0
        
        try:
            doc = Document.objects.get(id=document_id)
        except Document.DoesNotExist:
            return Response({"error": "Dự thảo không tồn tại."}, status=404)
        
        def clean_id(val):
            if val is None or str(val).strip() == "" or str(val).lower() == 'nan':
                return None
            try:
                if isinstance(val, str) and '-' in val:
                    return int(val.split('-')[-1])
                return int(val)
            except (ValueError, TypeError):
                return None

        try:
            with transaction.atomic():
                for row in rows:
                    mode = row.get('import_mode', 'add_new')
                    node_id = clean_id(row.get('node_id'))
                    appendix_id = clean_id(row.get('appendix_id'))
                    agency_id = clean_id(row.get('agency_id'))
                    content = row.get('content', '')
                    reason = row.get('reason', '')
                    note = row.get('note', '')
                    need_opinion = row.get('need_opinion', False)
                    duplicate_id = row.get('duplicate_id')
                    
                    # 1. Xử lý Feedback (Góp ý)
                    feedback_obj = None
                    if mode == 'skip' or (mode == 'add_if_diff' and row.get('is_duplicate')):
                        skipped_count += 1
                        if duplicate_id:
                            feedback_obj = Feedback.objects.filter(id=duplicate_id).first()
                    elif mode == 'explanation_only' and duplicate_id:
                        feedback_obj = Feedback.objects.filter(id=duplicate_id).first()
                    elif mode == 'overwrite' and duplicate_id:
                        Feedback.objects.filter(id=duplicate_id).update(
                            node_id=node_id,
                            appendix_id=appendix_id,
                            agency_id=agency_id,
                            content=content,
                            reason=reason,
                            note=note,
                            need_opinion=need_opinion,
                            user=request.user
                        )
                        feedback_obj = Feedback.objects.get(id=duplicate_id)
                        updated_count += 1
                    else:
                        feedback_obj = Feedback.objects.create(
                            document_id=document_id,
                            node_id=node_id,
                            appendix_id=appendix_id,
                            user=request.user,
                            agency_id=agency_id,
                            content=content,
                            reason=reason,
                            note=note,
                            need_opinion=need_opinion
                        )
                        created_count += 1
        
                    # 2. Xử lý Giải trình (Explanation)
                    exp_content = row.get('explanation_content', '')
                    exp_mode = row.get('explanation_import_mode', 'skip')
                    
                    if feedback_obj and exp_content and exp_mode == 'overwrite':
                        feedback_obj.explanations.all().delete()
                        from django.contrib.contenttypes.models import ContentType
                        Explanation.objects.create(
                            target_type='Feedback',
                            content_type=ContentType.objects.get_for_model(Feedback),
                            object_id=feedback_obj.id,
                            content=exp_content,
                            user=request.user
                        )
                        exp_count += 1
        
                    # 3. Xử lý Công văn (ConsultationResponse)
                    doc_num = row.get('official_number', '')
                    doc_date_raw = row.get('official_date', '')
                    if agency_id and doc_num and str(doc_num).lower() != 'nan':
                        official_date = None
                        if doc_date_raw and str(doc_date_raw).lower() != 'nan':
                            try:
                                from datetime import datetime
                                date_str = str(doc_date_raw).split(' ')[0]
                                for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y'):
                                    try:
                                        official_date = datetime.strptime(date_str, fmt).date()
                                        break
                                    except: continue
                            except: pass
                        
                        cr, created = ConsultationResponse.objects.get_or_create(
                            document=doc,
                            agency_id=agency_id,
                            official_number=doc_num,
                            defaults={'official_date': official_date}
                        )
                        if not created and official_date:
                            cr.official_date = official_date
                            cr.save()
                        if created: consultation_count += 1
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            return Response({"error": f"Lỗi khi lưu dữ liệu tại bản ghi '{index}': {str(e)}"}, status=500)

        # MỚI: Cập nhật Google Sheets nếu có link
        gs_url = request.data.get('gs_url')
        
        # MỚI: Lưu lại đường link vào Dự thảo nếu được yêu cầu
        save_gs_flag = request.data.get('save_gs_url')
        print(f"DEBUG: gs_url={gs_url}, save_gs_flag={save_gs_flag}, doc_id={document_id}")
        
        if gs_url and str(save_gs_flag).lower() in ['true', 'on', '1', 'yes'] and document_id:
            try:
                Document.objects.filter(id=document_id).update(google_sheets_url=gs_url)
                print(f"✅ Đã cập nhật link Google Sheets vào Dự thảo: {document_id}")
            except Exception as e:
                print(f"❌ Lỗi khi cập nhật link GS vào Dự thảo: {str(e)}")

        if gs_url and rows:
            self._update_google_sheet_status(gs_url, rows)

        # Tránh in trực tiếp ra console trên Windows nếu có tiếng Việt để không gây UnicodeEncodeError
        try:
            msg = f"Nhập hoàn tất: {created_count} góp ý mới, {updated_count} ghi đè. Đã xử lý {exp_count} giải trình và {consultation_count} công văn."
            print(msg.encode('ascii', 'ignore').decode('ascii')) # Safe print for console
        except:
            pass

        return Response({
            "message": f"Nhập hoàn tất: {created_count} góp ý mới, {updated_count} ghi đè. Đã xử lý {exp_count} giải trình và {consultation_count} công văn.",
            "created": created_count,
            "updated": updated_count,
            "skipped": skipped_count,
            "explanations": exp_count,
            "consultations": consultation_count
        })

    @action(detail=False, methods=['post'])
    def delete_all(self, request):
        document_id = request.data.get('document_id')
        if not document_id:
            return Response({"error": "Chưa chọn dự thảo để xóa."}, status=400)
            
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền xóa toàn bộ góp ý của dự thảo này."}, status=403)
            
        count = Feedback.objects.filter(document_id=document_id).count()
        Feedback.objects.filter(document_id=document_id).delete()
        
        return Response({"message": f"Đã xóa thành công toàn bộ {count} nội dung góp ý của dự thảo."}, status=200)

    def _get_inherited_assignments(self, node):
        """Lấy danh sách cán bộ được phân công cho một node, bao gồm cả kế thừa từ cha"""
        users_map = {} # id -> name
        curr = node
        while curr:
            for a in curr.assignments.all():
                if a.user_id not in users_map:
                    # Ưu tiên trường full_name tùy chỉnh, sau đó đến get_full_name, cuối cùng là username
                    users_map[a.user_id] = a.user.full_name or a.user.get_full_name() or a.user.username
            curr = curr.parent
        return [{"id": uid, "full_name": name} for uid, name in users_map.items()]

    def _get_gsheet_worksheet(self, gs_url):
        try:
            import gspread
            from google.oauth2.service_account import Credentials
            from google.auth.exceptions import GoogleAuthError
            import re
            
            # 1. Auth (Dùng credentials file)
            scopes = ['https://www.googleapis.com/auth/spreadsheets']
            key_path = os.path.join(settings.BASE_DIR, 'keys.json')
            if not os.path.exists(key_path):
                key_path = os.path.join(settings.BASE_DIR, 'google_keys.json')
                
            if not os.path.exists(key_path):
                return None, f"File keys.json hoặc google_keys.json không tồn tại tại {settings.BASE_DIR}. Vui lòng liên hệ quản trị viên."
                
            try:
                creds = Credentials.from_service_account_file(key_path, scopes=scopes)
                client = gspread.authorize(creds)
            except GoogleAuthError as auth_e:
                return None, f"Lỗi xác thực Google API: {str(auth_e)}. Kiểm tra tính hợp lệ của file keys."
            except Exception as e:
                return None, f"Lỗi không xác định khi khởi tạo Google Client: {str(e)}"
            
            # 2. Open Sheet (Lấy ID và GID từ link)
            match = re.search(r'/spreadsheets/d/([a-zA-Z0-9-_]+)', gs_url)
            if not match: 
                return None, "Link Google Sheet không hợp lệ. Phải có định dạng: https://docs.google.com/spreadsheets/d/SPREADSHEET_ID/..."
            
            sheet_id = match.group(1)
            
            gid_match = re.search(r'gid=([0-9]+)', gs_url)
            gid = gid_match.group(1) if gid_match else "0"
            
            try:
                spreadsheet = client.open_by_key(sheet_id)
            except gspread.exceptions.SpreadsheetNotFound:
                return None, f"Không tìm thấy Spreadsheet với ID: {sheet_id}. Kiểm tra lại URL."
            except gspread.exceptions.APIError as api_e:
                return None, f"Lỗi API Google Sheets (Quyền truy cập?): {str(api_e)}. Hãy đảm bảo đã chia sẻ Sheet cho email trong file keys."
            except Exception as open_e:
                return None, f"Không thể mở Spreadsheet. Chi tiết: {str(open_e)}"

            # Tìm worksheet theo GID
            worksheet = None
            try:
                for ws in spreadsheet.worksheets():
                    if str(ws.id) == gid:
                        worksheet = ws
                        break
                if not worksheet: 
                    worksheet = spreadsheet.get_worksheet(0)
            except Exception as ws_e:
                 return None, f"Lỗi khi truy cập Worksheet: {str(ws_e)}"
            
            return worksheet, None
        except Exception as e:
            return None, f"Lỗi hệ thống khi kết nối Google Sheets: {str(e)}"

    def _update_google_sheet_status(self, gs_url, rows):
        try:
            import gspread
            worksheet, err = self._get_gsheet_worksheet(gs_url)
            if not worksheet:
                print(f"❌ _update_google_sheet_status failed: {err}")
                return

            # 3. Headers check & Identify columns
            headers = worksheet.row_values(1)
            
            def get_or_create_col(name):
                try:
                    return headers.index(name) + 1
                except ValueError:
                    new_col = len(headers) + 1
                    worksheet.update_cell(1, new_col, name)
                    headers.append(name)
                    return new_col

            col_content = get_or_create_col("ND")
            col_exp = get_or_create_col("GT")
            
            # 4. Batch Update (Ghi giá trị 'OK' cho các dòng tương ứng)
            cells_to_update = []
            for row in rows:
                r_idx = row.get('row_index')
                if not r_idx: continue
                
                if row.get('has_imported_content'):
                    cells_to_update.append(gspread.Cell(row=r_idx, col=col_content, value="OK"))
                if row.get('has_imported_explanation'):
                    cells_to_update.append(gspread.Cell(row=r_idx, col=col_exp, value="OK"))
            
            if cells_to_update:
                worksheet.update_cells(cells_to_update)
                print(f"✅ Đã cập nhật xong {len(cells_to_update)} trạng thái vào Google Sheets: {gs_url}")
            else:
                print("⚠️ Không có dòng nào cần cập nhật trạng thái vào Google Sheets.")
                
        except Exception as e:
            print(f"❌ Lỗi khi cập nhật Google Sheets: {str(e)}")
            import traceback
            traceback.print_exc()

    @action(detail=False, methods=['post'])
    def parse_file(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Nhập góp ý cho dự thảo này."}, status=403)
        
        file_obj = request.FILES.get('file')
        agency_default = request.data.get('contributing_agency', '')
        agency_id = request.data.get('agency_id')
        
        if not file_obj:
            return Response({"error": "No file uploaded"}, status=400)
            
        doc = docx.Document(file_obj)
        results = []
        metadata = {
            "drafting_agency": "",
            "agency_location": "",
            "total_consulted_doc": 0,
            "total_feedbacks_doc": 0
        }
        
        # Sơ bộ tìm metadata ở 20 đoạn đầu
        for i, para in enumerate(doc.paragraphs[:20]):
            text = para.text.strip()
            if not text: continue
            
            loc_match = re.search(r'^([\w\s]+),\s*ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+', text)
            if loc_match and not metadata["agency_location"]:
                metadata["agency_location"] = loc_match.group(1).strip()
            
            if i < 7 and any(kw in text.upper() for kw in ["BỘ ", "TỔNG CỤC ", "ỦY BAN ", "SỞ "]):
                if not metadata["drafting_agency"]:
                    metadata["drafting_agency"] = text
            
            consult_match = re.search(r'tổng số cơ quan.*?là\s*(\d+)', text.lower()) or re.search(r'lấy ý kiến.*?:?\s*(\d+)', text.lower())
            if consult_match:
                metadata["total_consulted_doc"] = int(consult_match.group(1))
            
            feedback_match = re.search(r'tổng số ý kiến.*?là\s*(\d+)', text.lower()) or re.search(r'số ý kiến nhận được.*?:?\s*(\d+)', text.lower())
            if feedback_match:
                metadata["total_feedbacks_doc"] = int(feedback_match.group(1))

        # Phân rã từ BẢNG (thường có trong phụ lục)
        table_context_node = None
        for table in doc.tables:
            # Kiểm tra xem bảng có phải bảng góp ý không (dựa vào header)
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if not any(kw in " ".join(header_cells) for kw in ["góp ý", "ý kiến", "nội dung", "điều", "khoản"]):
                continue
            
            # Tìm cột chứa nội dung và cột chứa node_label
            content_col_idx = -1
            node_col_idx = -1
            agency_col_idx = -1
            
            for idx, text in enumerate(header_cells):
                if any(kw in text for kw in ["góp ý", "nội dung", "ý kiến"]): content_col_idx = idx
                if any(kw in text for kw in ["điều", "khoản", "mục", "vị trí"]): node_col_idx = idx
                if any(kw in text for kw in ["cơ quan", "đơn vị", "chủ thể"]): agency_col_idx = idx
            
            # Nếu không tìm thấy cột nội dung, lấy cột rộng nhất hoặc mặc định cột 2
            if content_col_idx == -1: continue
            
            for row in table.rows[1:]: # Bỏ qua header
                cells = row.cells
                if len(cells) <= content_col_idx: continue
                
                row_content = cells[content_col_idx].text.strip()
                if not row_content: continue
                
                row_node_label = cells[node_col_idx].text.strip() if node_col_idx != -1 and len(cells) > node_col_idx else ""
                row_agency = cells[agency_col_idx].text.strip() if agency_col_idx != -1 and len(cells) > agency_col_idx else agency_default
                
                # Tìm node_id tự động
                suggested_node = self._find_best_node_match(document_id, row_node_label, parent_context=table_context_node)
                if suggested_node and suggested_node.node_type in ['Điều', 'Khoản']:
                    table_context_node = suggested_node # Cập nhật context cho dòng sau
                
                results.append({
                    "key": f"table-{len(results)}",
                    "node_label": row_node_label if row_node_label else "Vấn đề khác",
                    "node_id": suggested_node.id if suggested_node else None,
                    "contributing_agency": row_agency,
                    "content": row_content
                })

        # Phân rã nội dung từ Văn bản (Paragraphs)
        current_node_label = "Vấn đề khác"
        current_node_obj = None # Context cho phân cấp
        current_content = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text: continue
            
            node_match = re.search(r'(Điều\s+\d+|Khoản\s+\d+|Điểm\s+[a-zđ])', text, re.IGNORECASE)
            
            if node_match:
                if current_content.strip():
                    results.append({
                        "key": f"para-{len(results)}",
                        "node_label": current_node_label,
                        "node_id": current_node_obj.id if current_node_obj else None,
                        "contributing_agency": agency_default,
                        "agency_id": agency_id,
                        "content": current_content.strip()
                    })
                
                new_label = node_match.group(1).title()
                # Thử tìm node và cập nhật context
                new_node = self._find_best_node_match(document_id, new_label, parent_context=current_node_obj if 'Điều' not in new_label else None)
                
                current_node_label = new_label
                current_node_obj = new_node
                current_content = text
            else:
                if current_content:
                    current_content += f"\n{text}"
                else:
                    current_content = text
                    
        if current_content.strip():
            results.append({
                "key": f"para-{len(results)}",
                "node_label": current_node_label,
                "node_id": current_node_obj.id if current_node_obj else None,
                "contributing_agency": agency_default,
                "agency_id": agency_id,
                "content": current_content.strip()
            })
            
        return Response({
            "metadata": metadata,
            "feedbacks": results
        })

    @action(detail=False, methods=['post'])
    def ocr_parse(self, request):
        """
        Xử lý OCR cho file Ảnh/PDF, trả về văn bản đã qua AI sửa lỗi và highlight diff.
        """
        try:
            document_id = request.data.get('document_id')
            if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
                return Response({"error": "Bạn không có quyền hoặc ID văn bản không hợp lệ."}, status=403)
            
            file_obj = request.FILES.get('file')
            if not file_obj:
                return Response({"error": "Không tìm thấy tệp tin được tải lên."}, status=400)
                
            # Save temporary file
            fs = FileSystemStorage(location=os.path.join(settings.MEDIA_ROOT, 'uploads_ocr'))
            
            # Làm sạch tên file để tránh lỗi tiếng Việt/ký tự lạ trên Windows
            safe_name = get_valid_filename(file_obj.name)
            print(f"--- [OCR API] Đang lưu tệp tin: {safe_name} ---")
            filename = fs.save(safe_name, file_obj)
            file_path = fs.path(filename)
            
            from .utils.ocr_service import OCRService # Lazy load
            print(f"--- [OCR API] Bắt đầu xử lý AI cho {filename} (đợi nạp bộ não)... ---")
            
            selected_pages = request.data.get('selected_pages') # VD: "2, 4" hoặc "1-3, 5"
            
            service = OCRService()
            ocr_results = service.process_file(file_path, target_pages=selected_pages)
            print(f"--- [OCR API] Bóc tách hoàn tất {len(ocr_results)} trang. ---")
            
            # Clean up raw upload if desired, or keep for audit
            # os.remove(file_path) 
            
            return Response({
                "pages": ocr_results
            })
        except Exception as e:
            import traceback
            with open('error_log_ocr.txt', 'w', encoding='utf-8') as f:
                traceback.print_exc(file=f)
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def get_pdf_info(self, request):
        """
        Lấy thông tin số trang và tạo thumbnails cho file PDF/Ảnh.
        """
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "No file"}, status=400)
            
        fs = FileSystemStorage(location=os.path.join(settings.MEDIA_ROOT, 'temp_previews'))
        filename = fs.save(get_valid_filename(file_obj.name), file_obj)
        file_path = fs.path(filename)
        ext = os.path.splitext(file_path)[1].lower()
        
        previews = []
        total_pages = 0
        
        import fitz
        try:
            if ext == '.pdf':
                doc = fitz.open(file_path)
                total_pages = len(doc)
                # Chỉ tạo thumbnail cho tất cả các trang
                for i in range(total_pages):
                    page = doc.load_page(i)
                    # Tạo ảnh nhỏ (thumbnail)
                    pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5)) 
                    img_name = f"thumb_{os.urandom(4).hex()}_{i}.jpg"
                    img_p = os.path.join(fs.location, img_name)
                    pix.save(img_p)
                    rel_p = os.path.relpath(img_p, settings.MEDIA_ROOT).replace('\\', '/')
                    previews.append(settings.MEDIA_URL + rel_p)
                doc.close()
            else:
                total_pages = 1
                rel_p = os.path.relpath(file_path, settings.MEDIA_ROOT).replace('\\', '/')
                previews.append(settings.MEDIA_URL + rel_p)
                
            return Response({
                "total_pages": total_pages,
                "previews": previews,
                "temp_file_path": filename # Để frontend dùng lại nếu cần (tuy nhiên hiện tại vẫn upload lại ocr_parse cho đơn giản)
            })
        except Exception as e:
            return Response({"error": str(e)}, status=500)
        finally:
            # Lưu ý: Không xóa file ở đây vì preview cần URL. 
            # Hệ thống nên có job dọn dẹp temp định kỳ.
            pass

    @action(detail=False, methods=['post'])
    def ocr_finalize_parse(self, request):
        """
        Giai đoạn cuối: biến văn bản sau OCR thành danh sách góp ý có cấu trúc.
        """
        document_id = request.data.get('document_id')
        text = request.data.get('text', '')
        
        if not text:
            return Response({"error": "Văn bản trống"}, status=400)
            
        from .utils.ai_parser import AIParser
        parser = AIParser()
        raw_feedbacks = parser.parse_text_to_feedbacks(text)
        
        if not raw_feedbacks:
            return Response({"error": "Không thể phân tích văn bản này bằng AI."}, status=500)
            
        # Làm sạch và khớp NodeID
        results = []
        for i, f in enumerate(raw_feedbacks):
            node_label = f.get('node_label', 'Chung')
            content = f.get('content', '')
            agency = f.get('agency', '')
            
            # Sử dụng logic khớp thông minh của hệ thống
            suggested_node = self._find_best_node_match(document_id, node_label)
            
            results.append({
                "key": f"ocr-finalize-{i}-{os.urandom(4).hex()}",
                "node_label": suggested_node.node_label if suggested_node else node_label,
                "node_id": suggested_node.id if suggested_node else None,
                "contributing_agency": agency or "",
                "content": content
            })
            
        return Response({
            "feedbacks": results
        })

    @action(detail=False, methods=['get'])
    def get_document_nodes(self, request):
        """Lấy danh sách node và Phụ lục để phục vụ mapping ở frontend với nhãn đầy đủ"""
        doc_id = request.query_params.get('document_id')
        if not doc_id: return Response([])
        
        nodes = DocumentNode.objects.filter(document_id=doc_id).select_related('parent', 'parent__parent').order_by('order_index')
        
        results = []
        for n in nodes:
            # Tạo nhãn phân tầng: Điều 1 > Khoản 2 > Điểm a
            path_parts = []
            curr = n
            while curr:
                if curr.node_type != 'Văn bản':
                    path_parts.insert(0, curr.node_label)
                curr = curr.parent
            
            label = " > ".join(path_parts) if path_parts else n.node_label
            
            # Nếu là Phụ lục, thực hiện rút gọn nhãn theo yêu cầu
            if n.node_type == 'Phụ lục' or n.node_label.lower().startswith('phụ lục'):
                label = re.split(r'[:.\n]', label)[0].strip()
                
            results.append({
                "id": n.id,
                "unique_id": f"node-{n.id}",
                "value": f"node-{n.id}",
                "label": label, 
                "type": n.node_type
            })
            
        # Thêm Phụ lục từ bảng DocumentAppendix
        from documents.models import DocumentAppendix
        appendices = DocumentAppendix.objects.filter(document_id=doc_id).order_by('name')
        for app in appendices:
            # Rút gọn nhãn hiển thị triệt để theo yêu cầu: "Phụ lục I"
            display_name = re.split(r'[:.\n]', app.name)[0].strip()
            if not display_name.lower().startswith('phụ lục'):
                display_name = f"Phụ lục {display_name}"
            
            results.append({
                "id": app.id,
                "unique_id": f"app-{app.id}",
                "value": f"app-{app.id}",
                "label": display_name,
                "type": "Appendix"
            })
            
        return Response(results)

    @action(detail=False, methods=['post'])
    def bulk_create(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Nhập góp ý cho dự thảo này."}, status=403)
            
        feedbacks_data = request.data.get('feedbacks', [])
        metadata = request.data.get('metadata', {})
        
        # Cập nhật Document Metadata
        from documents.models import Document
        doc = Document.objects.filter(id=document_id).first()
        if doc and metadata:
            if metadata.get('drafting_agency'): doc.drafting_agency = metadata['drafting_agency']
            if metadata.get('agency_location'): doc.agency_location = metadata['agency_location']
            if metadata.get('total_consulted_doc'): doc.total_consulted_doc = metadata['total_consulted_doc']
            if metadata.get('total_feedbacks_doc'): doc.total_feedbacks_doc = metadata['total_feedbacks_doc']
            doc.save()

        # Tìm node làm fallback (Ưu tiên nhãn "Chung")
        fallback_node = DocumentNode.objects.filter(document_id=document_id, node_label='Chung').first()
        if not fallback_node:
            fallback_node = DocumentNode.objects.filter(document_id=document_id, node_type='Vấn đề khác').first()

        created_count = 0
        for item in feedbacks_data:
            node_id = item.get('node_id') # Ưu tiên id nếu frontend đã map
            node_label = item.get('node_label', '')
            
            node = None
            if node_id:
                node = DocumentNode.objects.filter(id=node_id, document_id=document_id).first()
            
            if not node and node_label:
                # Tìm theo label (icontains)
                node = DocumentNode.objects.filter(document_id=document_id, node_label__icontains=node_label).first()
            
            # Fallback
            target_node = node or fallback_node
            
            if target_node:
                # Ưu tiên lấy cơ quan góp ý và số công văn từ metadata (chung cho cả lần nhập)
                global_agency = metadata.get('contributing_agency')
                global_agency_id = metadata.get('agency_id')
                global_doc_number = metadata.get('official_doc_number')

                Feedback.objects.create(
                    document_id=document_id,
                    node=target_node,
                    user=request.user,
                    contributing_agency=global_agency or item.get('contributing_agency', ''),
                    agency_id=global_agency_id or item.get('agency_id'),
                    official_doc_number=global_doc_number or item.get('official_doc_number', ''),
                    content=item.get('content', ''),
                    reason=item.get('reason', ''),
                    note=item.get('note', ''),
                    need_opinion=item.get('need_opinion', "")
                )
                created_count += 1
                
        return Response({"message": f"Đã lưu thành công {created_count} góp ý.", "count": created_count})

    @action(detail=False, methods=['post'])
    def ai_suggest(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Giải trình'):
            return Response({"error": "Bạn không có quyền sử dụng AI Giải trình cho dự thảo này."}, status=403)

        node_content = request.data.get('node_content', '')
        feedback_content = request.data.get('feedback_content', '')
        
        system_prompt = """Bạn là một Chuyên viên Pháp chế cấp cao của cơ quan Nhà nước trực thuộc Chính phủ Việt Nam.
NGUYÊN TẮC LẬP LUẬN:
- Bạn mang tâm thế của cơ quan chủ trì soạn thảo. ƯU TIÊN MẶC ĐỊNH là bảo vệ, giữ nguyên nội dung dự thảo ban đầu, trừ khi ý kiến góp ý chỉ ra sai sót pháp lý nghiêm trọng không thể chối cãi.
- Lập luận từ chối tiếp thu phải cực kỳ khéo léo, khách quan, hợp tình hợp lý. Dựa vào các lý do như: tính khả thi, đánh giá tác động thủ tục hành chính, bối cảnh thực tiễn...
- Văn phong mang ĐẬM TÍNH NGÔN NGỮ HÀNH CHÍNH NHÀ NƯỚC.
FORMAT TRẢ LỜI CỐ ĐỊNH:
- KHÔNG XƯNG HÔ "tôi", "bạn", "chúng tôi", "AI". Chỉ dùng đại từ "Cơ quan chủ trì soạn thảo".
- KHÔNG CÓ lời chào hỏi, dạo đầu. TRẢ VỀ TRỰC TIẾP kết quả giải trình để điền thẳng vào khung báo cáo."""
        
        user_prompt = f"Nội dung gốc:\n{node_content}\n\nNội dung góp ý:\n{feedback_content}\n\nHãy đưa ra nội dung giải trình ngắn gọn, đanh thép để bảo vệ dự thảo."
        
        from core.models import SystemSetting
        db_setting = SystemSetting.objects.filter(key='OPENAI_API_KEY').first()
        api_key = db_setting.value if db_setting and db_setting.value else os.environ.get('OPENAI_API_KEY')
        
        if not api_key:
            return Response({
                "suggestion": "Cơ quan chủ trì soạn thảo xin ghi nhận ý kiến. Tuy nhiên, qua rà soát bối cảnh thực tiễn và tính bộ khung của toàn Hệ thống điều khoản, kính đề nghị giữ nguyên như dự thảo để đảm bảo tính khả thi.\n\n(Hệ thống đang chạy MOCK AI do Môi trường Backend chưa nạp biến OPENAI_API_KEY. Để nạp key, vui lòng Set Environment OPENAI_API_KEY=sk-...)"
            })
            
        try:
            client = OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=600
            )
            return Response({
                "suggestion": response.choices[0].message.content.strip()
            })
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def save_explanation(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Giải trình'):
            return Response({"error": "Bạn không có quyền Lưu giải trình cho dự thảo này."}, status=403)
            
        target_type = request.data.get('target_type')
        object_id = request.data.get('object_id')
        content = request.data.get('content')
        
        if target_type == 'Feedback':
            content_type = ContentType.objects.get_for_model(Feedback)
        else:
            content_type = ContentType.objects.get_for_model(DocumentNode)
            
        explanation, created = Explanation.objects.update_or_create(
            target_type=target_type,
            content_type=content_type,
            object_id=object_id,
            defaults={'content': content, 'user': request.user}
        )
        
        # Log action and update feedback status if target is Feedback
        if target_type == 'Feedback':
            fb = Feedback.objects.filter(id=object_id).first()
            if fb:
                if fb.status == 'pending':
                    fb.status = 'reviewed'
                    fb.save()
                ActionLog.objects.create(
                    user=request.user,
                    feedback=fb,
                    action="Lưu giải trình",
                    details=f"Nội dung: {str(content)[:100]}..." if content else "Cập nhật giải trình."
                )
            
        return Response({"message": "Đã lưu giải trình"})

    @action(detail=False, methods=['post'])
    def delete_explanation(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Giải trình'):
            return Response({"error": "Bạn không có quyền Xóa giải trình cho dự thảo này."}, status=403)
            
        target_type = request.data.get('target_type')
        object_id = request.data.get('object_id')
        
        if target_type == 'Feedback':
            content_type = ContentType.objects.get_for_model(Feedback)
        else:
            content_type = ContentType.objects.get_for_model(DocumentNode)
            
        deleted_count, _ = Explanation.objects.filter(
            target_type=target_type,
            content_type=content_type,
            object_id=object_id
        ).delete()
        
        if target_type == 'Feedback' and deleted_count > 0:
            fb = Feedback.objects.filter(id=object_id).first()
            if fb:
                fb.status = 'pending'
                fb.save()
                ActionLog.objects.create(
                    user=request.user,
                    feedback=fb,
                    action="Xóa giải trình",
                    details="Chuyên viên đã xóa nội dung giải trình."
                )
            
        return Response({"message": "Đã xóa giải trình thành công."})

    @action(detail=True, methods=['post'])
    def submit_for_review(self, request, pk=None):
        fb = self.get_object()
        if not self._check_permission(request, fb.document_id, 'Chuyên viên Giải trình'):
            return Response({"error": "No permission"}, status=403)
        
        fb.status = 'reviewed'
        fb.save()
        ActionLog.objects.create(
            user=request.user,
            feedback=fb,
            action="Gửi duyệt",
            details="Chuyên viên đã gửi duyệt nội dung giải trình."
        )
        return Response({"message": "Đã gửi duyệt thành công"})

    @action(detail=True, methods=['post'])
    def approve_feedback(self, request, pk=None):
        fb = self.get_object()
        # Admin or Lead only
        user = request.user
        is_lead = fb.document.lead == user or user.is_staff or user.is_superuser
        if not is_lead:
            return Response({"error": "Chỉ Lãnh đạo/Admin mới có quyền phê duyệt."}, status=403)
        
        fb.status = 'approved'
        fb.save()
        ActionLog.objects.create(
            user=request.user,
            feedback=fb,
            action="Phê duyệt",
            details="Lãnh đạo đã phê duyệt nội dung giải trình."
        )
        return Response({"message": "Đã phê duyệt thành công"})

    @action(detail=False, methods=['get'])
    def by_node(self, request):
        """Lấy danh sách góp ý phẳng cho một node (và các con của nó)"""
        node_id = request.query_params.get('node_id')
        filter_type = request.query_params.get('filter_type', 'has_feedback')
        if not node_id: return Response([])
        
        from documents.models import DocumentNode
        try:
            root_node = DocumentNode.objects.get(id=node_id)
            
            def get_all_descendant_ids(node):
                ids = [node.id]
                for child in node.children.all():
                    ids.extend(get_all_descendant_ids(child))
                return ids
            
            target_ids = get_all_descendant_ids(root_node)
            
            feedbacks = Feedback.objects.filter(node_id__in=target_ids)\
                .select_related('node', 'agency')\
                .prefetch_related('explanations', 'logs', 'logs__user')\
                .order_by('created_at')
            
            results = []
            for fb in feedbacks:
                path = []
                curr = fb.node
                while curr:
                    path.insert(0, curr.node_label)
                    curr = curr.parent
                
                explanation_obj = fb.explanations.first()
                if filter_type == 'resolved' and not explanation_obj: continue
                if filter_type == 'unresolved' and explanation_obj: continue
                
                results.append({
                    "id": fb.id,
                    "node_id": fb.node_id,
                    "node_path": ", ".join(path),
                    "node_content": fb.node.content if fb.node else "",
                    "contributing_agency": fb.contributing_agency or (fb.agency.name if fb.agency else "Ẩn danh"),
                    "agency_category": fb.agency.category if fb.agency else "other",
                    "content": fb.content,
                    "explanation": explanation_obj.content if explanation_obj else "",
                    "status": fb.status,
                    "created_at": fb.created_at.isoformat(),
                    "logs": [
                        {
                            "username": log.user.username,
                            "action": log.action,
                            "time": log.created_at.isoformat(),
                            "details": log.details
                        } for log in fb.logs.all()
                    ]
                })
            return Response(results)
        except DocumentNode.DoesNotExist:
            return Response({"error": "Node not found"}, status=404)

    @action(detail=False, methods=['get'])
    def subject_stats(self, request):
        document_id = request.query_params.get('document_id')
        query = Feedback.objects.all()
        if document_id:
            query = query.filter(document_id=document_id)
            
        from django.db.models.functions import Coalesce
        agency_counts = query.annotate(
            display_name=Coalesce('agency__name', 'contributing_agency')
        ).values('display_name', 'agency__category', 'agency__agency_category__name').annotate(
            total=models.Count('id'),
            resolved=models.Count('id', filter=models.Q(explanations__isnull=False))
        ).order_by('-total')
        
        DEFAULT_CAT = 'Phân loại khác'
        agency_stats_list = []
        found_categories = set()
        
        for item in agency_counts:
            raw_cat = item['agency__agency_category__name'] or item['agency__category'] or 'other'
            cat_label = raw_cat if raw_cat != 'other' else DEFAULT_CAT
            found_categories.add(cat_label)
            
            total = item['total']
            resolved = item['resolved']
            agency_stats_list.append({
                'agency': item['display_name'] or 'Ẩn danh',
                'total': total,
                'resolved': resolved,
                'resolve_rate': round((resolved / total * 100), 1) if total > 0 else 0,
                'category': cat_label
            })

        # 5. Thống kê cấp ý kiến (Quy tắc mới)
        AGREED_PHRASE = "thống nhất với nội dung dự thảo Nghị định"
        
        # Tính số lượng cơ quan thống nhất (từ toàn bộ query)
        count_agreed = query.filter(content__icontains=AGREED_PHRASE)\
            .values('agency_id', 'contributing_agency').distinct().count()
        
        # Tập dữ liệu loại trừ các ý kiến thống nhất để tính các thông số khác
        active_query = query.exclude(content__icontains=AGREED_PHRASE)
        
        total_fbs = active_query.count()
        
        from .models import Explanation
        # Lấy danh sách ID các ý kiến có giải trình tiếp thu
        accepted_fb_ids = Explanation.objects.filter(
            target_type='Feedback',
            content__iregex=r'tiếp\s+thu'
        ).values_list('object_id', flat=True)
        
        # Lấy danh sách ID các ý kiến có giải trình tiếp thu một phần
        partial_fb_ids = Explanation.objects.filter(
            target_type='Feedback',
            content__iregex=r'tiếp\s+thu\s+một\s+phần'
        ).values_list('object_id', flat=True)
        
        count_accepted = active_query.filter(id__in=accepted_fb_ids).distinct().count()
        count_partial = active_query.filter(id__in=partial_fb_ids).distinct().count()
        
        # Feedback có giải trình nhưng không tiếp thu (và không thống nhất)
        count_explained_no_acc = active_query.filter(
            explanations__isnull=False
        ).exclude(
            id__in=accepted_fb_ids
        ).distinct().count()
        
        # Feedback chưa giải trình (và không thống nhất)
        count_pending = active_query.filter(explanations__isnull=True).count()


        category_stats = {}
        for stat in query.values('agency__agency_category__name', 'agency__category').annotate(count=models.Count('agency', distinct=True)):
            raw_cat = stat['agency__agency_category__name'] or stat['agency__category'] or 'other'
            label = raw_cat if raw_cat != 'other' else DEFAULT_CAT
            category_stats[label] = category_stats.get(label, 0) + stat['count']
        
        invited_category_stats = {}
        if document_id:
            from documents.models import Document
            try:
                doc = Document.objects.get(id=document_id)
                for cat_stat in doc.consulted_agencies.values('agency_category__name', 'category').annotate(count=models.Count('id')):
                    raw_cat = cat_stat['agency_category__name'] or cat_stat['category'] or 'other'
                    label = raw_cat if raw_cat != 'other' else DEFAULT_CAT
                    invited_category_stats[label] = invited_category_stats.get(label, 0) + cat_stat['count']
            except Document.DoesNotExist:
                pass

        # Count feedbacks needing leadership opinion (where need_opinion is documented)
        count_need_opinion = query.exclude(need_opinion__isnull=True).exclude(need_opinion='').count()

        return Response({
            'agency_stats': agency_stats_list,
            'category_stats': category_stats,
            'invited_category_stats': invited_category_stats,
            'available_categories': sorted(list(found_categories)),
            'summary': {
                'total_fbs': query.count(),
                'total_pending': count_pending,
                'total_agreed': count_agreed,
                'total_accepted': count_accepted,
                'total_partial': count_partial,
                'total_explained_no_acc': count_explained_no_acc,
                'total_need_opinion': count_need_opinion,
            }
        })

    @action(detail=False, methods=['get'])
    def by_appendix(self, request):
        """Lấy danh sách góp ý cho một Phụ lục"""
        appendix_id = request.query_params.get('appendix_id')
        if not appendix_id: return Response([])
        try:
            feedbacks = Feedback.objects.filter(appendix_id=appendix_id)\
                .exclude(content__icontains="thống nhất với nội dung dự thảo Nghị định")\
                .select_related('appendix', 'agency')\
                .prefetch_related('explanations', 'logs', 'logs__user').order_by('created_at')
            results = []
            # Xác định số thứ tự phụ lục
            appendix = DocumentAppendix.objects.get(id=appendix_id)
            all_apps = list(DocumentAppendix.objects.filter(document_id=appendix.document_id).order_by('created_at').values_list('id', flat=True))
            try:
                app_num = all_apps.index(appendix.id) + 1
            except ValueError:
                app_num = "?"
            app_label = f"Phụ lục {app_num}"

            results = []
            for fb in feedbacks:
                explanation_obj = fb.explanations.first()
                results.append({
                    "id": fb.id,
                    "appendix_id": fb.appendix_id,
                    "node_path": app_label,
                    "node_label": app_label,
                    "node_content": fb.appendix.content,
                    "contributing_agency": fb.contributing_agency or (fb.agency.name if fb.agency else "Ẩn danh"),
                    "agency_category": fb.agency.category if fb.agency else "other",
                    "content": fb.content,
                    "explanation": explanation_obj.content if explanation_obj else "",
                    "status": fb.status,
                    "created_at": fb.created_at.isoformat(),
                    "logs": [{"username": log.user.username, "action": log.action, "time": log.created_at.isoformat(), "details": log.details} for log in fb.logs.all()]
                })
            return Response(results)
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['get'])
    @action(detail=False, methods=['get'])
    def by_document(self, request):
        """Lấy danh sách góp ý của một Dự thảo với bộ lọc và phân trang"""
        doc_id = request.query_params.get('document_id')
        if not doc_id:
            return Response({"error": "Vui lòng cung cấp document_id"}, status=400)
            
        # Các tham số lọc
        node_id = request.query_params.get('node_id')
        agency_id = request.query_params.get('agency')
        specialist = request.query_params.get('specialist') # Cán bộ thụ lý (dựa trên giải trình)
        status_filter = request.query_params.get('status')
        search = request.query_params.get('search')
        
        # Query cơ bản
        queryset = Feedback.objects.filter(document_id=doc_id)\
            .select_related('node', 'agency', 'appendix')\
            .prefetch_related('explanations', 'logs', 'logs__user', 'node__assignments__user')\
            .order_by('node__order_index', 'created_at')

        # 1. Lọc theo Điều/Khoản (Bao gồm cả các cấp con)
        if node_id and node_id != 'all':
            # Loại bỏ tiền tố 'node-' hoặc 'app-' nếu có
            clean_node_id = node_id
            if isinstance(node_id, str):
                if node_id.startswith('node-'): clean_node_id = int(node_id.replace('node-', ''))
                elif node_id.startswith('app-'): 
                    queryset = queryset.filter(appendix_id=node_id.replace('app-', ''))
                    clean_node_id = None
            
            if clean_node_id:
                # Tìm tất cả nút con (Khoản, Điểm, Tiết...) thuộc Điều/Khoản này để bao hàm trong bộ lọc
                all_node_ids = {clean_node_id}
                current_layer_ids = [clean_node_id]
                for _ in range(3): # Duyệt tối đa 3 cấp con
                    child_ids = list(DocumentNode.objects.filter(parent_id__in=current_layer_ids).values_list('id', flat=True))
                    if not child_ids:
                        break
                    all_node_ids.update(child_ids)
                    current_layer_ids = child_ids
                
                queryset = queryset.filter(node_id__in=all_node_ids)
            
        # 1.1 Lọc theo Cơ quan
        if agency_id and agency_id != 'all':
            queryset = queryset.filter(agency_id=agency_id)
            
        # 2. Tìm kiếm nội dung
        if search:
            queryset = queryset.filter(
                models.Q(content__icontains=search) | 
                models.Q(contributing_agency__icontains=search) |
                models.Q(node__node_label__icontains=search)
            )

        # 3. Lọc theo Cán bộ (Specialist) - Bao gồm phân công kế thừa
        if specialist and specialist != 'all':
            if specialist == 'none':
                # Lọc những góp ý hoàn toàn chưa được gán (ở bất kỳ cấp nào trong phân cấp node)
                queryset = queryset.filter(
                    node__assignments__isnull=True,
                    node__parent__assignments__isnull=True,
                    node__parent__parent__assignments__isnull=True
                )
            else:
                queryset = queryset.filter(
                    models.Q(node__assignments__user_id=specialist) |
                    models.Q(node__parent__assignments__user_id=specialist) |
                    models.Q(node__parent__parent__assignments__user_id=specialist)
                )

        # 4. Lọc theo Trạng thái giải trình
        # Lưu ý: Do trạng thái tính toán dựa trên text trong giải trình nên ta cần xử lý sau khi lấy queryset 
        # hoặc tối ưu bằng cách lọc DB nếu có thể. Ở đây ta áp dụng logic lọc sau/trước tùy độ phức tạp.
        if status_filter and status_filter != 'all':
            import re
            AGREED_REGEX = r'thống\s+nhất\s+với\s+nội\s+dung\s+dự\s+thảo\s+Nghị\s+định'
            
            # Chuyển thành list để lọc phức tạp (giống report)
            feedbacks_list = list(queryset)
            filtered_list = []
            for fb in feedbacks_list:
                exps = fb.explanations.all()
                exp_text = " ".join([e.content for e in exps if e.content]).lower()
                content_text = (fb.content or "").lower()
                
                match = False
                if status_filter == 'pending':
                    match = not exps
                elif status_filter == 'explained':
                    match = exps and 'tiếp thu' not in exp_text
                elif status_filter == 'accepted':
                    match = exps and 'tiếp thu' in exp_text and 'tiếp thu một phần' not in exp_text
                elif status_filter == 'partially_accepted':
                    match = exps and 'tiếp thu một phần' in exp_text
                elif status_filter == 'agreed':
                    match = re.search(AGREED_REGEX, content_text, re.IGNORECASE)
                
                if match:
                    filtered_list.append(fb)
            queryset = filtered_list
        else:
            queryset = list(queryset)

        # 5. Phân trang Thủ công (vì queryset có thể đã bị ép thành list sau khi lọc status)
        try:
            page = int(request.query_params.get('page', 1))
            page_size = int(request.query_params.get('page_size', 20))
        except ValueError:
            page = 1
            page_size = 20
            
        total_count = len(queryset)
        start = (page - 1) * page_size
        end = start + page_size
        paginated_data = queryset[start:end]

        # Map appendix_id sang số thứ tự
        appendices = DocumentAppendix.objects.filter(document_id=doc_id).order_by('created_at')
        appendix_map = {app.id: idx + 1 for idx, app in enumerate(appendices)}

        results = []
        for fb in paginated_data:
            node_label = "Chung"
            node_path = "Chung"
            if fb.appendix_id:
                app_num = appendix_map.get(fb.appendix_id, "?")
                node_label = f"Phụ lục {app_num}"
                node_path = f"Phụ lục {app_num}"
            elif fb.node:
                # Ưu tiên hiển thị Điều làm nhãn chính nếu là Khoản/Điểm
                if fb.node.node_type in ['Khoản', 'Điểm'] and fb.node.parent:
                    node_label = fb.node.parent.node_label
                else:
                    node_label = fb.node.node_label
                node_path = self._get_full_node_path(fb.node)

            explanation_obj = fb.explanations.first()

            # Lấy danh sách cán bộ được phân công (bao gồm quy tắc tự động và kế thừa)
            assigned_users, _ = self._get_feedback_assignments(fb)
            
            results.append({
                "id": fb.id,
                "document_id": fb.document_id,
                "node_id": fb.node_id,
                "appendix_id": fb.appendix_id,
                "node_label": node_label,
                "node_path": node_path,
                "agency": fb.agency_id,
                "contributing_agency": fb.contributing_agency or (fb.agency.name if fb.agency else "Ẩn danh"),
                "official_doc_number": fb.official_doc_number,
                "content": fb.content,
                "explanation": explanation_obj.content if explanation_obj else "",
                "status": fb.status,
                "assigned_users": assigned_users,
                "created_at": fb.created_at.isoformat(),
            })
            
        return Response({
            "count": total_count,
            "page": page,
            "page_size": page_size,
            "results": results
        })

    @action(detail=False, methods=['get'])
    def uncontributed(self, request):
        """Trả về danh sách các cơ quan CHƯA có góp ý cho dự thảo cụ thể"""
        doc_id = request.query_params.get('document_id')
        if not doc_id:
            return Response({"error": "Vui lòng cung cấp document_id"}, status=400)
            
        from core.models import Agency
        # Lấy ID của các Agency đã góp ý cho dự thảo này
        contributed_agency_ids = Feedback.objects.filter(
            document_id=doc_id, agency__isnull=False
        ).values_list('agency_id', flat=True).distinct()
        
        # Các Agency còn lại là chưa góp ý
        uncontributed = Agency.objects.exclude(id__in=contributed_agency_ids).values('id', 'name', 'category')
        
        return Response(list(uncontributed))

    @action(detail=False, methods=['get'])
    def custom_report_preview(self, request):
        doc_id = request.query_params.get('document_id')
        agency = request.query_params.get('agency')
        status_filter = request.query_params.get('status')
        specialist = request.query_params.get('specialist')
        report_type = request.query_params.get('report_type', 'mau10')
        only_opinion = request.query_params.get('only_opinion') == 'true'
        show_agreed_text = request.query_params.get('show_agreed_text') == 'true'
        
        if not doc_id: return Response([])
        
        feedbacks = Feedback.objects.filter(document_id=doc_id).select_related('node', 'agency').prefetch_related('explanations', 'user').order_by('node__order_index')
        
        # Lọc bỏ các ý kiến thống nhất nếu không bật tùy chọn hiển thị
        from .utils.v2_template_generator import is_exact_agreement
        if not show_agreed_text:
            feedbacks = [f for f in feedbacks if not is_exact_agreement(f.content)]
        
        if only_opinion:
            feedbacks = feedbacks.exclude(need_opinion__isnull=True).exclude(need_opinion='')
        
        if agency and agency != 'all':
            import unicodedata
            norm_target = unicodedata.normalize('NFC', agency).strip().lower()
            feedbacks = [
                fb for fb in feedbacks 
                if (fb.contributing_agency and unicodedata.normalize('NFC', fb.contributing_agency).strip().lower() == norm_target) or
                (fb.agency and unicodedata.normalize('NFC', fb.agency.name).strip().lower() == norm_target)
            ]
            
        if status_filter and status_filter != 'all':
            import re
            AGREED_REGEX = r'thống\s+nhất\s+với\s+nội\s+dung\s+dự\s+thảo\s+Nghị\s+định'
            
            new_feedbacks = []
            for fb in feedbacks:
                exps = fb.explanations.all()
                exp_text = " ".join([e.content for e in exps if e.content]).lower()
                content_text = (fb.content or "").lower()
                
                match = False
                if status_filter == 'pending' or status_filter == 'unresolved':
                    match = not exps
                elif status_filter == 'explained':
                    match = exps and 'tiếp thu' not in exp_text
                elif status_filter == 'accepted' or status_filter == 'resolved':
                    match = exps and 'tiếp thu' in exp_text
                elif status_filter == 'partially_accepted':
                    match = exps and 'tiếp thu một phần' in exp_text
                elif status_filter == 'agreed':
                    match = re.search(AGREED_REGEX, content_text, re.IGNORECASE)
                
                if match:
                    new_feedbacks.append(fb)
            feedbacks = new_feedbacks

        if specialist and specialist != 'all':
            if specialist == 'none':
                if isinstance(feedbacks, list):
                    feedbacks = [fb for fb in feedbacks if not any(ex.user for ex in fb.explanations.all())]
                else:
                    feedbacks = feedbacks.filter(explanations__user__isnull=True)
            else:
                if isinstance(feedbacks, list):
                    feedbacks = [fb for fb in feedbacks if any(str(ex.user_id) == str(specialist) for ex in fb.explanations.all())]
                else:
                    feedbacks = feedbacks.filter(explanations__user_id=specialist)
            
        # Lay cau hinh truong tu template
        from reports.models import ReportTemplate
        template = ReportTemplate.objects.filter(template_type=report_type, is_active=True).first()
        fields = []
        if template:
            fields = template.field_configs.filter(is_enabled=True).order_by('column_order')

        from .utils.v2_template_generator import _get_field_value
        
        results = []
        for i, fb in enumerate(feedbacks, 1):
            explanation = fb.explanations.first()
            if fields:
                row = { f.field_key: _get_field_value(f.field_key, i, fb, explanation, show_agreed_text=show_agreed_text) for f in fields }
            else:
                # Fallback mac dinh
                dieu_khoan = f"{fb.node.node_label}" if fb.node else ""
                if fb.node and fb.node.parent:
                    dieu_khoan = f"{fb.node.parent.node_label}, {fb.node.node_label}"
                row = {
                    "stt": i,
                    "dieu_khoan": dieu_khoan,
                    "co_quan": (fb.agency.name if fb.agency else fb.contributing_agency) or "Khác",
                    "noi_dung_gop_y": fb.content,
                    "noi_dung_giai_trinh": explanation.content if explanation else "",
                    "xin_y_kien": fb.need_opinion or "",
                }
            results.append(row)
            
        return Response(results)

    @action(detail=False, methods=['get'], permission_classes=[])
    def export_mau_10(self, request):
        doc_id = request.query_params.get('document_id')
        agency = request.query_params.get('agency')
        status_filter = request.query_params.get('status')
        only_opinion = request.query_params.get('only_opinion') == 'true'
        show_agreed_text = request.query_params.get('show_agreed_text') == 'true'
        
        # Xác thực qua token ở URL (Cho phép tải file trực tiếp từ trình duyệt)
        user = request.user
        if not user.is_authenticated:
            token = request.query_params.get('token')
            if not token:
                return Response({"error": "Vui lòng cung cấp token"}, status=401)
            from rest_framework_simplejwt.tokens import AccessToken
            from django.contrib.auth import get_user_model
            try:
                access_token = AccessToken(token)
                user = get_user_model().objects.get(id=access_token['user_id'])
            except Exception:
                return Response({"error": "Token không hợp lệ"}, status=401)
        
        if not doc_id:
            return Response({"error": "Vui lòng cung cấp document_id"}, status=400)
            
        try:
            document = Document.objects.get(id=doc_id)
            feedbacks = Feedback.objects.filter(document_id=doc_id).select_related('node', 'agency').prefetch_related('explanations').order_by('node__order_index')
            
            # Lọc bỏ các ý kiến thống nhất nếu không bật tùy chọn hiển thị
            from .utils.v2_template_generator import is_exact_agreement
            if not show_agreed_text:
                feedbacks = [f for f in feedbacks if not is_exact_agreement(f.content)]
            
            if only_opinion:
                feedbacks = feedbacks.exclude(need_opinion__isnull=True).exclude(need_opinion='')
            
            if agency and agency != 'all':
                import unicodedata
                norm_target = unicodedata.normalize('NFC', agency).strip().lower()
                feedbacks = [
                    fb for fb in feedbacks 
                    if (fb.contributing_agency and unicodedata.normalize('NFC', fb.contributing_agency).strip().lower() == norm_target) or
                    (fb.agency and unicodedata.normalize('NFC', fb.agency.name).strip().lower() == norm_target)
                ]
                
            if status_filter and status_filter != 'all':
                import re
                AGREED_REGEX = r'thống\s+nhất\s+với\s+nội\s+dung\s+dự\s+thảo\s+Nghị\s+định'
                
                new_feedbacks = []
                for fb in feedbacks:
                    exps = fb.explanations.all()
                    exp_text = " ".join([e.content for e in exps if e.content]).lower()
                    content_text = (fb.content or "").lower()
                    
                    match = False
                    if status_filter == 'pending' or status_filter == 'unresolved':
                        match = not exps
                    elif status_filter == 'explained':
                        match = exps and 'tiếp thu' not in exp_text
                    elif status_filter == 'accepted' or status_filter == 'resolved':
                        match = exps and 'tiếp thu' in exp_text
                    elif status_filter == 'partially_accepted':
                        match = exps and 'tiếp thu một phần' in exp_text
                    elif status_filter == 'agreed':
                        match = re.search(AGREED_REGEX, content_text, re.IGNORECASE)
                    
                    if match:
                        new_feedbacks.append(fb)
                feedbacks = new_feedbacks
            
            # Kiểm tra an toàn xem feedbacks (list hoặc queryset) có trống không
            has_data = feedbacks.exists() if hasattr(feedbacks, 'exists') else (len(feedbacks) > 0)
            if not has_data:
                return Response({"error": "Không có ý kiến góp ý nào thỏa mãn bộ lọc để xuất báo cáo."}, status=404)
            
            # Phan nhanh theo loai bao cao (report_type param)
            report_type = request.query_params.get('report_type', 'mau10')
            if report_type == 'mau10': report_type = 'mau_10' # Chuẩn hóa về format DB
            
            # Doc cau hinh tu DB cho loai tuong ung
            template_config = {}
            tpl_db = None
            try:
                from reports.models import ReportTemplate, ReportFieldConfig
                tpl_db = ReportTemplate.objects.filter(template_type=report_type, is_active=True).first()
                if tpl_db:
                    fields_qs = tpl_db.field_configs.filter(is_enabled=True).order_by('column_order')
                    fields = [
                        {
                            'field_key': f.field_key,
                            'field_label': f.field_label,
                            'column_width_cm': f.column_width_cm
                        } for f in fields_qs
                    ]
                    template_config = {
                        'header_org_name': tpl_db.header_org_name,
                        'header_org_location': tpl_db.header_org_location,
                        'footer_signer_name': tpl_db.footer_signer_name,
                        'footer_signer_title': tpl_db.footer_signer_title,
                        'fields': fields
                    }
            except Exception as e:
                print(f"Error loading template config: {e}")

            from .utils.v2_template_generator import generate_from_v2_template
            from django.http import HttpResponse

            # HỢP NHẤT TUYỆT ĐỐI: Dùng V2 generator (Landscape) cho mọi chế độ xuất
            file_stream = generate_from_v2_template(
                document, 
                feedbacks, 
                template_config=template_config, 
                template_type=report_type,
                show_agreed_text=show_agreed_text
            )
            
            # Cấu hình tên file
            type_label = 'Mau_10' if report_type == 'mau_10' else 'Tuy_chinh'
            filename = f"Bao_cao_{type_label}_{document.id}.docx"
            
            # Trả về file Word cho người dùng (SỬ DỤNG HTTPRESPONSE ĐỂ ĐẢM BẢO ỔN ĐỊNH DỮ LIỆU)
            response = HttpResponse(
                file_stream.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Document.DoesNotExist:
            return Response({"error": "Văn bản không tồn tại"}, status=404)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": f"Lỗi xuất báo cáo: {str(e)}"}, status=500)

    @action(detail=True, methods=['post'])
    def reassign_node(self, request, pk=None):
        fb = self.get_object()
        new_node_id = request.data.get('node_id')
        if not new_node_id:
            return Response({"error": "node_id is required"}, status=400)
        
        try:
            from documents.models import DocumentNode
            new_node = DocumentNode.objects.get(id=new_node_id, document_id=fb.document_id)
            old_label = fb.node.node_label if fb.node else "N/A"
            fb.node = new_node
            fb.save()
            
            ActionLog.objects.create(
                user=request.user,
                feedback=fb,
                action="Gắn lại điều khoản",
                details=f"Chuyển từ '{old_label}' sang '{new_node.node_label}'"
            )
            
            return Response({"message": "Đã chuyển điều khoản thành công."})
        except DocumentNode.DoesNotExist:
            return Response({"error": "Điều khoản mới không hợp lệ hoặc không thuộc dự thảo này."}, status=404)

    def _get_short_label(self, label):
        """Cắt bỏ tiêu đề dài sau dấu ':' hoặc ' -', chỉ giữ lại phần ngắn gọn.
        Ví dụ: 'Phụ lục I: PHÂN LOẠI CÔNG TRÌNH...' -> 'Phụ lục I'
        """
        if not label:
            return label
        short = label.split(':')[0].split(' -')[0].strip()
        return short if short else label

    def _get_full_node_path(self, node):
        if not node: return "Chung"
        path = []
        current = node
        while current:
            path.insert(0, self._get_short_label(current.node_label))
            current = current.parent
        return ", ".join(path)

    def _normalize_text(self, s):
        import unicodedata
        import re
        if not s: return ""
        # Bỏ HTML tags
        s = re.sub(r'<[^>]+>', ' ', str(s))
        s = str(s).strip()
        # Xoá các ký tự đặc biệt, dấu câu, nhưng vẫn giữ khoảng trắng để hiển thị/so sánh cơ bản
        s = re.sub(r'[^\w\s]', ' ', s)
        s = re.sub(r'\s+', ' ', s)
        return unicodedata.normalize('NFC', s).strip().lower()

    def _normalize_match_key(self, s):
        """Dạng siêu chuẩn hóa để so khớp (không khoảng trắng, không dấu câu, mở rộng từ viết tắt)"""
        if not s: return ""
        s = self._normalize_text(s)
        # Expansion of common acronyms
        s = s.replace("tvgs", "tư vấn giám sát")
        s = s.replace("tp", "thành phố")
        s = s.replace("tw", "trung ương")
        s = s.replace("ubnd", "ủy ban nhân dân")
        # Xóa hoàn toàn khoảng trắng
        import re
        s = re.sub(r'\s+', '', s)
        return s

    def _get_feedback_assignments(self, fb):
        """
        Logic tính toán phân công chuyên viên tập trung với 3 mức ưu tiên:
        1. Phân công cá nhân (FeedbackAssignment)
        2. Quy tắc tự động (Thống nhất với nội dung dự thảo Nghị định -> Quốc Anh)
        3. Phân công kế thừa (DocumentNodeAssignment)
        """
        # 1. Ưu tiên Phân công cá nhân
        individual = [{"id": a.user_id, "full_name": a.user.full_name or a.user.username} 
                     for a in fb.individual_assignments.all()]
        if individual:
            return individual, individual # Trả về cả list đầy đủ và list cá nhân để phân biệt

        # 2. Quy tắc tự động
        norm_content = self._normalize_text(fb.content)
        if norm_content == "thống nhất với nội dung dự thảo nghị định":
            auto_assigned = [{"id": 9, "full_name": "Quốc Anh"}]
            return auto_assigned, []

        # 3. Phân công kế thừa từ Điều/Khoản/Phụ lục
        node_assigned = self._get_inherited_assignments(fb.node) if fb.node else []
        return node_assigned, []

    @action(detail=False, methods=['post'])
    def assign_feedbacks(self, request):
        """
        Gán chuyên viên cho danh sách các góp ý (Feedback) cụ thể.
        Payload: { "assignments": [ { "feedback_id": 1, "user_ids": [1, 2] }, ... ] }
        """
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Phân công cho dự thảo này."}, status=403)
            
        assignments_data = request.data.get('assignments', [])
        if not assignments_data:
            return Response({"error": "Thiếu dữ liệu phân công."}, status=400)
            
        from django.db import transaction
        try:
            with transaction.atomic():
                for item in assignments_data:
                    fb_id = item.get('feedback_id')
                    user_ids = item.get('user_ids', [])
                    
                    # Xóa phân công cũ của góp ý này
                    FeedbackAssignment.objects.filter(feedback_id=fb_id).delete()
                    
                    # Thêm phân công mới
                    new_assignments = [
                        FeedbackAssignment(feedback_id=fb_id, user_id=uid, assigned_by=request.user)
                        for uid in user_ids
                    ]
                    FeedbackAssignment.objects.bulk_create(new_assignments)
                    
            return Response({"message": "Đã cập nhật phân công riêng biệt thành công."})
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def gsheet_compare(self, request):
        document_id = request.data.get('document_id')
        gs_url = request.data.get('gs_url')
        
        if not document_id or not gs_url:
            return Response({"error": "Thiếu document_id hoặc gs_url."}, status=400)
            
        try:
            # 1. Fetch local feedbacks
            local_feedbacks = Feedback.objects.filter(document_id=document_id).select_related('node', 'agency', 'appendix').prefetch_related('explanations', 'individual_assignments', 'individual_assignments__user')
            
            # 2. Fetch GS rows
            worksheet, error_msg = self._get_gsheet_worksheet(gs_url)
            if not worksheet:
                return Response({"error": error_msg}, status=400)
            
            all_values = worksheet.get_all_values()
            if not all_values:
                return Response({"error": "Google Sheet trống hoặc không có dòng tiêu đề."}, status=400)
                
            headers = [h.strip().lower() for h in all_values[0]]
            rows = all_values[1:]
            
            # Identify columns using fuzzy matching keywords
            col_map = {}
            for idx, h in enumerate(headers):
                if 'node' not in col_map and any(kw in h for kw in ['điều', 'khoản', 'mục', 'vị trí']): 
                    col_map['node'] = idx
                elif 'explanation' not in col_map and any(kw in h for kw in ['giải trình', 'tiếp thu', 'ý kiến giải trình']): 
                    col_map['explanation'] = idx
                elif 'agency' not in col_map and any(kw in h for kw in ['cơ quan', 'chủ thể', 'đơn vị', 'người']): 
                    col_map['agency'] = idx
                elif 'need_opinion' not in col_map and any(kw in h for kw in ['xin ý kiến', 'xyk', 'cần xin ý kiến', 'vấn đề còn ý kiến khác nhau']):
                    col_map['need_opinion'] = idx
                elif 'content' not in col_map and any(kw in h for kw in ['nội dung', 'ý kiến', 'góp ý']): 
                    col_map['content'] = idx
                elif 'specialist' not in col_map and any(kw in h for kw in ['cán bộ', 'chuyên viên', 'thụ lý', 'phân công']): 
                    col_map['specialist'] = idx

            if 'content' not in col_map:
                 # Thử tìm kiếm sâu hơn hoặc báo lỗi rõ ràng
                 return Response({
                     "error": "Không tìm thấy cột 'Nội dung góp ý' trong Google Sheet.",
                     "found_headers": headers,
                     "suggestion": "Hãy đảm bảo Sheet có dòng tiêu đề và chứa cột 'Nội dung', 'Ý kiến' hoặc 'Góp ý'."
                 }, status=400)

            # 3. Reference data for labels
            from documents.models import DocumentAppendix
            appendices = DocumentAppendix.objects.filter(document_id=document_id)
            
            def _to_roman(num):
                val = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
                syms = ['M','CM','D','CD','C','XC','L','XL','X','IX','V','IV','I']
                result = ''
                for i, v in enumerate(val):
                    while num >= v:
                        result += syms[i]
                        num -= v
                return result
            
            # Lấy tên phụ lục dạng ngắn gọn: chỉ lấy phần trước ':' hoặc '-'
            # Ví dụ: 'Phụ lục I: PHÂN LOẠI...' -> 'Phụ lục I'
            appendix_map = {}
            for idx, app in enumerate(appendices):
                name = (app.name or '').strip()
                # Cắt bỏ phần sau dấu ':' hoặc ' -' để lấy tên ngắn
                short_name = name.split(':')[0].split(' -')[0].strip()
                if short_name:
                    appendix_map[app.id] = short_name
                else:
                    # Fallback: Phụ lục + số La Mã theo thứ tự
                    appendix_map[app.id] = f"Ph\u1ee5 l\u1ee5c {_to_roman(idx + 1)}"
            
            # 4. Compare logic
            from difflib import SequenceMatcher

            def get_key(content, agency):
                return f"{self._normalize_text(content)}|{self._normalize_text(agency)}"

            from collections import defaultdict
            gs_data = defaultdict(list)
            gs_data_robust = defaultdict(list) # Key là robust-content|robust-agency
            gs_data_content_only = defaultdict(list) # Key là robust-content (dành cho nội dung dài)
            
            unmatched_gs = []
            for i, row in enumerate(rows):
                raw_content = row[col_map['content']] if len(row) > col_map['content'] else ""
                raw_agency = row[col_map['agency']] if 'agency' in col_map and len(row) > col_map['agency'] else ""
                raw_exp = row[col_map['explanation']] if 'explanation' in col_map and len(row) > col_map['explanation'] else ""
                raw_specialist = row[col_map['specialist']] if 'specialist' in col_map and len(row) > col_map['specialist'] else ""
                raw_node = row[col_map['node']] if 'node' in col_map and len(row) > col_map['node'] else ""
                
                norm_c = self._normalize_text(raw_content)
                norm_a = self._normalize_text(raw_agency)
                robust_c = self._normalize_match_key(raw_content)
                robust_a = self._normalize_match_key(raw_agency)
                
                entry = {
                    "row": i + 2,
                    "content": raw_content,
                    "agency": raw_agency,
                    "explanation": raw_exp,
                    "specialist": raw_specialist,
                    "node": raw_node,
                    "norm_content": norm_c,
                    "norm_agency": norm_a,
                    "robust_content": robust_c,
                    "robust_agency": robust_a,
                    "need_opinion": row[col_map['need_opinion']].strip() if 'need_opinion' in col_map and len(row) > col_map['need_opinion'] else ""
                }

                if norm_c or norm_a:
                    key = f"{norm_c}|{norm_a}"
                    gs_data[key].append(entry)
                    
                    robust_key = f"{robust_c}|{robust_a}"
                    gs_data_robust[robust_key].append(entry)
                    
                    if len(robust_c) > 60: # Chỉ dùng mapping content-only nếu nội dung đủ dài
                        gs_data_content_only[robust_c].append(entry)
                        
                    unmatched_gs.append(entry)

            results = []
            for fb in local_feedbacks:
                norm_fb_content = self._normalize_text(fb.content)
                norm_fb_agency = self._normalize_text(fb.contributing_agency)
                robust_fb_content = self._normalize_match_key(fb.content)
                robust_fb_agency = self._normalize_match_key(fb.contributing_agency)
                
                key = f"{norm_fb_content}|{norm_fb_agency}"
                robust_key = f"{robust_fb_content}|{robust_fb_agency}"
                
                gs_info = None
                
                # PASS 1: Exact Match (Content + Agency, with spaces)
                if key in gs_data and gs_data[key]:
                    gs_info = gs_data[key].pop(0)
                
                # PASS 2: Robust Match (Content + Agency, without spaces/acronyms)
                if not gs_info and robust_key in gs_data_robust and gs_data_robust[robust_key]:
                    gs_info = gs_data_robust[robust_key].pop(0)
                
                # PASS 3: Content-Only Robust Match (Large contents)
                if not gs_info and len(robust_fb_content) > 60 and robust_fb_content in gs_data_content_only and gs_data_content_only[robust_fb_content]:
                    gs_info = gs_data_content_only[robust_fb_content].pop(0)

                # Cleanup references in other maps and unmatched list
                if gs_info:
                    if gs_info in unmatched_gs: unmatched_gs.remove(gs_info)
                    # Sync removal across maps
                    k1 = f"{gs_info['norm_content']}|{gs_info['norm_agency']}"
                    rk = f"{gs_info['robust_content']}|{gs_info['robust_agency']}"
                    ck = gs_info['robust_content']
                    if k1 in gs_data and gs_info in gs_data[k1]: gs_data[k1].remove(gs_info)
                    if rk in gs_data_robust and gs_info in gs_data_robust[rk]: gs_data_robust[rk].remove(gs_info)
                    if ck in gs_data_content_only and gs_info in gs_data_content_only[ck]: gs_data_content_only[ck].remove(gs_info)
                
                # FALLBACK: Fuzzy matching if no exact/robust match
                if not gs_info and norm_fb_content:
                    best_match = None
                    best_ratio = 0
                    for gs in unmatched_gs:
                        # Thử fuzzy trên nội dung chính. 
                        # Yêu cầu cơ quan đóng góp trùng khớp hoặc gần giống nếu có khai báo.
                        valid_agency = (not gs["robust_agency"] or not robust_fb_agency or 
                                        gs["robust_agency"] == robust_fb_agency or 
                                        SequenceMatcher(None, robust_fb_agency, gs["robust_agency"]).ratio() > 0.8)
                        
                        if valid_agency:
                            ratio = SequenceMatcher(None, norm_fb_content, gs["norm_content"]).ratio()
                            if ratio > best_ratio and ratio >= 0.85:
                                best_ratio = ratio
                                best_match = gs
                    
                    if best_match:
                        gs_info = best_match
                        # Remove from all pools
                        unmatched_gs.remove(gs_info)
                        k1 = f"{gs_info['norm_content']}|{gs_info['norm_agency']}"
                        rk = f"{gs_info['robust_content']}|{gs_info['robust_agency']}"
                        ck = gs_info['robust_content']
                        if k1 in gs_data and gs_info in gs_data[k1]: gs_data[k1].remove(gs_info)
                        if rk in gs_data_robust and gs_info in gs_data_robust[rk]: gs_data_robust[rk].remove(gs_info)
                        if ck in gs_data_content_only and gs_info in gs_data_content_only[ck]: gs_data_content_only[ck].remove(gs_info)
                
                exp = fb.explanations.first()
                exp_content = exp.content if exp else ""
                
                # Check for specific diffs even if matched
                is_content_diff = False
                is_exp_diff = False
                
                if gs_info:
                    # Sử dụng robust match key để bỏ qua sai sót khoảng trắng/dấu câu trong so sánh nội dung
                    is_content_diff = robust_fb_content != gs_info["robust_content"]
                    
                    db_robust_exp = self._normalize_match_key(exp_content)
                    gs_robust_exp = self._normalize_match_key(gs_info["explanation"])
                    is_exp_diff = db_robust_exp != gs_robust_exp

                # Lấy danh sách cán bộ: Ưu tiên phân công riêng, nếu rỗng thì lấy kế thừa từ node
                # Logic hiển thị theo thứ tự ưu tiên tập trung (Cá nhân > Tự động > Kế thừa)
                final_assignments, individual_assignments = self._get_feedback_assignments(fb)
                node_assignments = self._get_inherited_assignments(fb.node) if fb.node else []

                gs_specialist = gs_info["specialist"] if gs_info else ""
                is_specialist_diff = False
                if gs_info and 'specialist' in col_map:
                    # So khớp nâng cao: So sánh tập hợp tên (Set-based)
                    import re
                    db_names = set(self._normalize_match_key(u['full_name']) for u in final_assignments if u.get('full_name'))
                    # Tách tên chuyên viên trên Sheet bằng các dấu phân cách phổ biến
                    gs_names_raw = re.split(r'[,;|\n]', gs_specialist)
                    gs_names = set(self._normalize_match_key(name) for name in gs_names_raw if name.strip())
                    
                    if db_names != gs_names:
                        is_specialist_diff = True
                                
                # So sánh Vị trí (Node)
                gs_node = gs_info["node"] if gs_info else ""
                is_node_diff = False
                
                # Rút gọn nhãn phụ lục cho đồng bộ
                if fb.appendix_id:
                    node_label = appendix_map.get(fb.appendix_id, f"Phụ lục ?")
                else:
                    node_label = self._get_full_node_path(fb.node)
                
                is_node_diff = False
                if gs_info and 'node' in col_map:
                    # So sánh sau khi chuẩn hóa để bỏ qua sai sót khoảng trắng/chữ hoa
                    raw_node_diff = self._normalize_text(node_label) != self._normalize_text(gs_node)
                    # Nếu nội dung và giải trình đã khớp hoàn toàn, thì không coi là khác biệt vị trí (đáp ứng yêu cầu người dùng)
                    if raw_node_diff and (is_content_diff or is_exp_diff):
                        is_node_diff = True
                    
                results.append({
                    "id": fb.id,
                    "node_label": node_label,
                    "gs_node": gs_node,
                    "is_node_diff": is_node_diff,
                    "agency": fb.contributing_agency,
                    "content": fb.content,
                    "explanation": exp_content,
                    "assigned_users": final_assignments,
                    "individual_assignments": individual_assignments, 
                    "node_assignments": node_assignments,
                    "gs_specialist": gs_specialist,
                    "is_specialist_diff": is_specialist_diff,
                    "is_in_gs": gs_info is not None,
                    "gs_row": gs_info["row"] if gs_info else None,
                    "gs_content": gs_info["content"] if gs_info else "",
                    "gs_explanation": gs_info["explanation"] if gs_info else "",
                    "need_opinion": fb.need_opinion or "",
                    "gs_need_opinion": gs_info["need_opinion"] if gs_info else "",
                    "is_content_diff": is_content_diff,
                    "is_exp_diff": is_exp_diff,
                    "is_opinion_diff": (fb.need_opinion or "").strip() != (gs_info["need_opinion"] if gs_info else "").strip() if gs_info and 'need_opinion' in col_map else False,
                    "status": "synced" if (gs_info and not (is_content_diff or is_exp_diff or is_node_diff or is_specialist_diff or ((fb.need_opinion or "").strip() != (gs_info["need_opinion"] if gs_info else "").strip() if 'need_opinion' in col_map else False))) else ("diff" if (gs_info and (is_content_diff or is_exp_diff or is_node_diff or is_specialist_diff or ((fb.need_opinion or "").strip() != (gs_info["need_opinion"] if gs_info else "").strip() if 'need_opinion' in col_map else False))) else "new_in_db")
                })
                
            return Response({"feedbacks": results})
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def gsheet_pull_assignments(self, request):
        """Cập nhật phân công trong DB từ dữ liệu trên Google Sheet"""
        document_id = request.data.get('document_id')
        pull_items = request.data.get('pull_items', []) # [{id, gs_specialist}]
        
        if not document_id or not pull_items:
            return Response({"error": "Thiếu document_id hoặc dữ liệu cập nhật."}, status=400)
            
        try:
            from accounts.models import User
            # 1. Fetch all users to build a name-to-id map
            users = User.objects.all().only('id', 'full_name', 'username')
            name_map = {}
            for u in users:
                if u.full_name:
                    name_map[self._normalize_text(u.full_name)] = u.id
                name_map[self._normalize_text(u.username)] = u.id
            
            updated_count = 0
            for item in pull_items:
                fb_id = item.get('id')
                gs_spec_str = item.get('gs_specialist', '')
                
                if not fb_id or not gs_spec_str: continue
                
                # Parse specialist names (supports comma, semicolon, newline)
                names = re.split(r'[,;\n]', gs_spec_str)
                user_ids = []
                for name in names:
                    clean_name = name.strip()
                    if not clean_name: continue
                    
                    # Try direct match with normalized text
                    norm_name = self._normalize_text(clean_name)
                    if norm_name in name_map:
                        user_ids.append(name_map[norm_name])
                    else:
                        # Try case-insensitive comparison on raw names if normalization was too aggressive
                        raw_name_lower = clean_name.lower()
                        for u in users:
                            u_name = u.full_name.lower() if u.full_name else u.username.lower()
                            if raw_name_lower == u_name:
                                user_ids.append(u.id)
                                break
                
                if user_ids:
                    # Update DB
                    FeedbackAssignment.objects.filter(feedback_id=fb_id).delete()
                    new_assignments = [
                        FeedbackAssignment(feedback_id=fb_id, user_id=uid, assigned_by=request.user)
                        for uid in user_ids
                    ]
                    FeedbackAssignment.objects.bulk_create(new_assignments)
                    updated_count += 1
            
            return Response({"message": f"Đã cập nhật phân công cho {updated_count} góp ý từ Google Sheet."})
            
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def gsheet_pull_positions(self, request):
        """Cập nhật Vị trí trong DB từ dữ liệu trên Google Sheet"""
        document_id = request.data.get('document_id')
        pull_items = request.data.get('pull_items', []) # [{id, gs_node}]
        
        if not document_id or not pull_items:
            return Response({"error": "Thiếu document_id hoặc dữ liệu cập nhật."}, status=400)
            
        try:
            from documents.models import DocumentNode, DocumentAppendix
            from .models import Feedback
            
            # 1. Build node_map (norm_path -> node_id)
            nodes = DocumentNode.objects.filter(document_id=document_id)
            node_map = {}
            for n in nodes:
                path = self._get_full_node_path(n)
                node_map[self._normalize_text(path)] = n.id
                # Thử map thêm label ngắn lỡ người dùng gõ thiếu chương
                node_map[self._normalize_text(self._get_short_label(n.node_label))] = n.id
                
            # 2. Build appendix_map (norm_name -> appendix_id)
            def _to_roman(num):
                val = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
                syms = ['M','CM','D','CD','C','XC','L','XL','X','IX','V','IV','I']
                result = ''
                for i, v in enumerate(val):
                    while num >= v:
                        result += syms[i]
                        num -= v
                return result
                
            appendices = DocumentAppendix.objects.filter(document_id=document_id)
            appendix_map = {}
            for idx, app in enumerate(appendices):
                name = (app.name or '').strip()
                short_name = name.split(':')[0].split(' -')[0].strip()
                if short_name:
                    appendix_map[self._normalize_text(short_name)] = app.id
                else:
                    appendix_map[self._normalize_text(f"Phụ lục {_to_roman(idx + 1)}")] = app.id

            updated_count = 0
            for item in pull_items:
                fb_id = item.get('id')
                gs_node = item.get('gs_node', '').strip()
                
                if not fb_id: continue
                
                fb = Feedback.objects.filter(id=fb_id).first()
                if not fb: continue

                norm_gs = self._normalize_text(gs_node)
                
                if norm_gs == self._normalize_text("Chung") or not norm_gs:
                    fb.node_id = None
                    fb.appendix_id = None
                    fb.save()
                    updated_count += 1
                elif norm_gs in appendix_map:
                    fb.appendix_id = appendix_map[norm_gs]
                    fb.node_id = None
                    fb.save()
                    updated_count += 1
                elif norm_gs in node_map:
                    fb.node_id = node_map[norm_gs]
                    fb.appendix_id = None
                    fb.save()
                    updated_count += 1
            
            return Response({"message": f"Đã cập nhật vị trí cho {updated_count} góp ý từ Google Sheet."})
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def gsheet_pull_data(self, request):
        """Cập nhật Nội dung và Giải trình trong DB từ dữ liệu trên Google Sheet"""
        document_id = request.data.get('document_id')
        pull_items = request.data.get('pull_items', []) # [{id, content, explanation}]
        
        if not document_id or not pull_items:
            return Response({"error": "Thiếu dữ liệu cập nhật."}, status=400)
            
        try:
            from django.contrib.contenttypes.models import ContentType
            from .models import Feedback, Explanation
            feedback_ct = ContentType.objects.get_for_model(Feedback)
            
            updated_count = 0
            for item in pull_items:
                fb_id = item.get('id')
                if not fb_id: continue
                
                fb = Feedback.objects.filter(id=fb_id).first()
                if not fb: continue
                
                # 1. Cập nhật Giải trình (ưu tiên cao)
                gs_exp = item.get('explanation')
                if gs_exp is not None:
                    # Xóa tất cả giải trình cũ của góp ý này và thay bằng nội dung mới từ Sheet
                    fb.explanations.all().delete()
                    if gs_exp.strip():
                        Explanation.objects.create(
                            target_type='Feedback',
                            content_type=feedback_ct,
                            object_id=fb.id,
                            content=gs_exp.strip(),
                            user=request.user
                        )
                
                # 2. Cập nhật Nội dung (nếu có sự thay đổi)
                gs_content = item.get('content')
                if gs_content and gs_content.strip() != fb.content:
                    fb.content = gs_content.strip()
                    fb.save()
                
                # 3. Cập nhật Trạng thái "Cần xin ý kiến" (đã đổi sang TextField)
                gs_need_opinion = item.get('need_opinion')
                if gs_need_opinion is not None:
                    fb.need_opinion = gs_need_opinion
                    fb.save()
                
                updated_count += 1
                
            return Response({"message": f"Đã cập nhật {updated_count} bản ghi từ Google Sheet vào hệ thống thành công."})
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def gsheet_push(self, request):
        document_id = request.data.get('document_id')
        gs_url = request.data.get('gs_url')
        push_items = request.data.get('push_items', [])
        update_mode = request.data.get('update_mode', 'all') # 'all' or 'specialist_only'
        
        # Hỗ trợ legacy (trường hợp api gửi danh sách id thô)
        feedback_ids = request.data.get('feedback_ids', [])
        
        if not document_id or not gs_url or (not push_items and not feedback_ids):
            return Response({"error": "Thiếu tham số bắt buộc (document_id, gs_url, danh sách đẩy)."}, status=400)
            
        try:
            worksheet, error_msg = self._get_gsheet_worksheet(gs_url)
            if not worksheet:
                return Response({"error": error_msg}, status=400)
            
            if not push_items and feedback_ids:
                push_items = [{"id": fid, "gs_row": None} for fid in feedback_ids]
                
            all_fids = [item['id'] for item in push_items]
            feedbacks = Feedback.objects.filter(id__in=all_fids).select_related('node', 'appendix').prefetch_related('explanations', 'individual_assignments', 'individual_assignments__user')
            feedback_dict = {fb.id: fb for fb in feedbacks}
            
            # Build appendix_map cho gsheet_push (dùng La Mã, đồng bộ với gsheet_compare)
            def _to_roman(num):
                val = [1000,900,500,400,100,90,50,40,10,9,5,4,1]
                syms = ['M','CM','D','CD','C','XC','L','XL','X','IX','V','IV','I']
                result = ''
                for i, v in enumerate(val):
                    while num >= v:
                        result += syms[i]
                        num -= v
                return result
            
            from documents.models import DocumentAppendix
            _push_appendices = DocumentAppendix.objects.filter(document_id=document_id)
            appendix_map = {}
            for idx, app in enumerate(_push_appendices):
                name = (app.name or '').strip()
                # Cắt bỏ phần sau ':' hoặc ' -' để chỉ lấy 'Phụ lục I'
                short_name = name.split(':')[0].split(' -')[0].strip()
                if short_name:
                    appendix_map[app.id] = short_name
                else:
                    appendix_map[app.id] = f"Ph\u1ee5 l\u1ee5c {_to_roman(idx + 1)}"
            
            # Map headers to know where to write
            raw_headers = worksheet.row_values(1)
            headers = [h.strip().lower() for h in raw_headers]
            
            col_node = col_agency = col_content = col_reason = col_explanation = col_note = col_nd = col_gt = col_specialist = col_opinion = None
            
            for idx, h in enumerate(headers):
                # Ưu tiên các cột trạng thái viết tắt trước để tránh bị hốt vào cột nội dung chính
                if not col_nd and h == 'nd': col_nd = idx + 1
                elif not col_gt and h == 'gt': col_gt = idx + 1
                elif not col_node and any(kw in h for kw in ['điều', 'khoản', 'mục', 'vị trí']): col_node = idx + 1
                elif not col_opinion and any(kw in h for kw in ['xin ý kiến', 'xyk', 'cần xin ý kiến', 'vấn đề còn ý kiến khác nhau']): col_opinion = idx + 1
                elif not col_explanation and any(kw in h for kw in ['giải trình', 'tiếp thu', 'ý kiến giải trình']): col_explanation = idx + 1
                elif not col_agency and any(kw in h for kw in ['cơ quan', 'chủ thể', 'đơn vị', 'người']): col_agency = idx + 1
                elif not col_content and any(kw in h for kw in ['nội dung', 'ý kiến', 'góp ý']): col_content = idx + 1
                elif not col_reason and any(kw in h for kw in ['lý do', 'cơ sở']): col_reason = idx + 1
                elif not col_note and any(kw in h for kw in ['ghi chú', 'note']): col_note = idx + 1
                elif not col_specialist and any(kw in h for kw in ['cán bộ', 'chuyên viên', 'thụ lý', 'phân công', 'specialist']): col_specialist = idx + 1

            if not col_content:
                return Response({"error": "Google Sheet không có cột 'Nội dung góp ý' để định danh dữ liệu."}, status=400)
            
            rows_to_append = []
            cells_to_update = []
            from gspread.cell import Cell
            
            for item in push_items:
                fb = feedback_dict.get(item['id'])
                if not fb: continue
                
                if fb.appendix_id:
                    node_val = appendix_map.get(fb.appendix_id, "Phụ lục ?")
                else:
                    node_val = self._get_full_node_path(fb.node) if fb.node else "Chung"
                exp = fb.explanations.first()
                exp_val = exp.content if exp else ""

                gs_row = item.get('gs_row')
                if gs_row: # Cập nhật
                    if update_mode == 'all':
                        if col_node: cells_to_update.append(Cell(row=gs_row, col=col_node, value=node_val))
                        if col_agency: cells_to_update.append(Cell(row=gs_row, col=col_agency, value=fb.contributing_agency))
                        if col_content: cells_to_update.append(Cell(row=gs_row, col=col_content, value=fb.content))
                        if col_reason: cells_to_update.append(Cell(row=gs_row, col=col_reason, value=fb.reason or ""))
                        if col_explanation: cells_to_update.append(Cell(row=gs_row, col=col_explanation, value=exp_val))
                        if col_note: cells_to_update.append(Cell(row=gs_row, col=col_note, value=fb.note or ""))
                        
                        # Tự động điền trạng thái OK
                        if col_nd and fb.content: cells_to_update.append(Cell(row=gs_row, col=col_nd, value="OK"))
                        if col_gt and exp_val: cells_to_update.append(Cell(row=gs_row, col=col_gt, value="OK"))
                        if col_opinion: cells_to_update.append(Cell(row=gs_row, col=col_opinion, value=fb.need_opinion or ""))
                    
                    elif update_mode == 'node_only':
                        if col_node: cells_to_update.append(Cell(row=gs_row, col=col_node, value=node_val))
                    
                    if col_specialist:
                        # Thứ tự ưu tiên đồng bộ toàn hệ thống
                        assigned_list, _ = self._get_feedback_assignments(fb)
                        specialists_val = ", ".join([u['full_name'] for u in assigned_list])
                        cells_to_update.append(Cell(row=gs_row, col=col_specialist, value=specialists_val))
                    
                    # Logic cũ (Tự động điền trạng thái OK đã được cho vào trong block update_mode == 'all')
                else: # Thêm mới
                    row = [""] * len(headers)
                    if col_node: row[col_node-1] = node_val
                    if col_agency: row[col_agency-1] = fb.contributing_agency
                    if col_content: row[col_content-1] = fb.content
                    if col_reason: row[col_reason-1] = fb.reason or ""
                    if col_explanation: row[col_explanation-1] = exp_val
                    if col_note: row[col_note-1] = fb.note or ""
                    if col_opinion: row[col_opinion-1] = fb.need_opinion or ""
                    
                    if col_specialist:
                        # Thứ tự ưu tiên đồng bộ toàn hệ thống
                        assigned_list, _ = self._get_feedback_assignments(fb)
                        row[col_specialist-1] = ", ".join([u['full_name'] for u in assigned_list])
                    
                    # Trạng thái OK cho dòng mới
                    if col_nd and fb.content: row[col_nd-1] = "OK"
                    if col_gt and exp_val: row[col_gt-1] = "OK"
                    
                    rows_to_append.append(row)
                
            if rows_to_append:
                worksheet.append_rows(rows_to_append)
                
            if cells_to_update:
                worksheet.update_cells(cells_to_update)
                
            msg = []
            if rows_to_append: msg.append(f"Thêm mới {len(rows_to_append)} dòng.")
            if cells_to_update: 
                updated_count = len(push_items) - len(rows_to_append)
                msg.append(f"Cập nhật {updated_count} dòng có sẵn.")
                
            return Response({"message": " ".join(msg) or "Không có thay đổi nào."})
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": str(e)}, status=500)

class ConsultationResponseViewSet(viewsets.ModelViewSet):
    queryset = ConsultationResponse.objects.all()
    serializer_class = ConsultationResponseSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        doc_id = self.request.query_params.get('document_id')
        if doc_id:
            return ConsultationResponse.objects.filter(document_id=doc_id).order_by('-created_at')
        return super().get_queryset()




