"""
Generator Mẫu Báo cáo V2 — Consolidated Version
Phương pháp: Xây dựng file Word từ code (python-docx) để đảm bảo ổn định và linh hoạt.
Hỗ trợ cả Mẫu 10 và các mẫu tùy chỉnh từ Tab Cấu hình.
"""

import io
import os
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.oxml.ns import qn

FONT_NAME = 'Times New Roman'

def clean_text(text):
    """Giữ lại các ký tự hợp lệ cho XML 1.0, loại bỏ ký tự điều khiển."""
    if text is None:
        return ""
    text = str(text)
    return re.sub(r'[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]', '', text)

def _set_cell_style(cell, text, bold=False, italic=False, size=12, align=None):
    """Định dạng văn bản trong ô bảng"""
    cell.text = ''
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(clean_text(text))
    run.font.name = FONT_NAME
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic

def _get_field_value(field_key, fb_idx, fb, explanation=None):
    """Trích xuất dữ liệu dựa trên field_key (Tối ưu hóa mapping)"""
    # Lấy giải trình đầu tiên nếu không được truyền vào
    if explanation is None:
        explanations = fb.explanations.all() if hasattr(fb, 'explanations') else []
        explanation = explanations[0] if explanations else None
    else:
        explanations = [explanation] if explanation else []
    
    # Lấy tên chuyên viên
    chuyen_vien = ""
    if explanation and explanation.user:
        chuyen_vien = explanation.user.username or explanation.user.first_name
    elif hasattr(fb, 'user') and fb.user:
        chuyen_vien = fb.user.username
        
    # Map status
    status_map = {'pending': 'Chưa xử lý', 'reviewed': 'Đã giải trình', 'approved': 'Đã duyệt'}
    status_tv = status_map.get(fb.status, fb.status)

    # Agency logic
    agency = (fb.agency.name if hasattr(fb, 'agency') and fb.agency else fb.contributing_agency) or 'Ẩn danh'

    # Dieu khoan logic (Full label: Dieu 1, Khoan 2)
    dieu_khoan_val = ""
    if fb.node:
        dieu_khoan_val = fb.node.node_label or ""
        if fb.node.parent:
            dieu_khoan_val = f"{fb.node.parent.node_label}, {dieu_khoan_val}"
    
    # Theo yêu cầu mới: Cột dự thảo chỉ lấy phần nhãn Điều/Khoản
    du_thao_val = dieu_khoan_val or 'Chung'

    mapping = {
        'stt': str(fb_idx),
        'dieu_khoan': dieu_khoan_val or 'Chung',
        'noi_dung_du_thao': du_thao_val,
        'don_vi_gop_y': agency,
        'co_quan': agency,
        'user_name': agency,
        'noi_dung_gop_y': fb.content or '',
        'content': fb.content or '',
        'giai_trinh': "\n".join([f"- {ex.content}" for ex in explanations if ex.content]) if explanations else "Chưa có giải trình",
        'noi_dung_giai_trinh': "\n".join([f"- {ex.content}" for ex in explanations if ex.content]) if explanations else "Chưa có giải trình",
        'explanations': "\n".join([f"- {ex.content}" for ex in explanations if ex.content]) if explanations else "Chưa có giải trình",
        'chuyen_vien': chuyen_vien,
        'trang_thai': status_tv,
        'status': status_tv,
        'need_opinion': fb.need_opinion or '',
    }
    
    # Fallback cho các trường tùy chỉnh khác (nếu có trong model Feedback)
    if field_key not in mapping:
        return str(getattr(fb, field_key, ''))
        
    return mapping.get(field_key, '')

def _build_grouped_data(feedbacks):
    """Nhóm feedbacks theo node để phục vụ Merge Cells (cho Mẫu 10)"""
    nodes_map = {}
    node_order = []

    for fb in feedbacks:
        if fb.node:
            node_key = fb.node.id
            node_label = fb.node.node_label
            if fb.node.parent:
                node_label = f"{fb.node.parent.node_label}, {node_label}"
        else:
            node_key = '__no_node__'
            node_label = 'Chung'

        if node_key not in nodes_map:
            nodes_map[node_key] = {
                'label': node_label,
                'feedbacks': []
            }
            node_order.append(node_key)
        
        nodes_map[node_key]['feedbacks'].append(fb)

    return [nodes_map[k] for k in node_order]

def generate_from_v2_template(document, feedbacks, template_config=None, template_type='mau_10'):
    """
    Sinh báo cáo linh hoạt (Mẫu 10 hoặc Tùy chỉnh) sử dụng python-docx
    """
    doc = Document()
    cfg = template_config or {}
    
    # 1. Page Setup: Mặc định Landscape cho báo cáo bảng
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height
    
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)

    # 2. Header (Cơ quan + Quốc hiệu)
    drafting_agency = getattr(document, 'drafting_agency', '') or cfg.get('header_org_name') or 'BỘ / CƠ QUAN CHỦ TRÌ'
    agency_location = getattr(document, 'agency_location', '') or cfg.get('header_org_location') or 'Hà Nội'
    now = datetime.now()
    export_date = f"ngày {now.day:02d} tháng {now.month:02d} năm {now.year}"

    header_table = doc.add_table(rows=2, cols=2)
    header_table.width = Cm(25)
    
    _set_cell_style(header_table.cell(0, 0), drafting_agency.upper(), bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_style(header_table.cell(0, 1), "CỘNG HÀ XÃ HỘI CHỦ NGHĨA VIỆT NAM", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_style(header_table.cell(1, 0), "-------", align=WD_ALIGN_PARAGRAPH.CENTER)
    _set_cell_style(header_table.cell(1, 1), "Độc lập - Tự do - Hạnh phúc", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_date = header_table.cell(1, 1).add_paragraph()
    p_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_date = p_date.add_run(f"{agency_location}, {export_date}")
    run_date.font.size = Pt(14); run_date.font.name = FONT_NAME; run_date.italic = True

    doc.add_paragraph()

    # 3. Title & Intro
    p_name = (document.project_name or 'Dự thảo văn bản').upper()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run(f"BẢN TỔNG HỢP Ý KIẾN, TIẾP THU, GIẢI TRÌNH Ý KIẾN GÓP Ý, PHẢN BIỆN XÃ HỘI ĐỐI VỚI {p_name}")
    run.bold = True; run.font.size = Pt(14)

    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.first_line_indent = Cm(1.27)
    p_intro.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_intro.add_run(f"Căn cứ Luật Ban hành văn bản quy pháp luật, cơ quan lập đề xuất chính sách/cơ quan chủ trì soạn thảo đã tổ chức lấy ý kiến, tham vấn/phản biện xã hội đối với hồ sơ chính sách {document.project_name}.").font.size = Pt(12)

    # Thống kê chi tiết
    import re
    if hasattr(feedbacks, 'prefetch_related'):
        fbs_list = list(feedbacks.prefetch_related('explanations'))
    else:
        fbs_list = list(feedbacks)

    agencies = set([f.contributing_agency or (f.agency.name if hasattr(f, 'agency') and f.agency else '') for f in fbs_list])
    total_consulted = len([a for a in agencies if a]) or len(fbs_list)
    total_fbs = len(fbs_list)

    count_agreed_agencies = 0
    agreed_agencies = set()
    AGREED_PHRASE = "thống nhất với nội dung dự thảo"
    
    # Lọc danh sách ý kiến "không thống nhất" để tính các chỉ số khác
    active_fbs = []
    for f in fbs_list:
        content_text = (f.content or "").lower()
        if AGREED_PHRASE.lower() in content_text:
            agency_name = f.contributing_agency or (f.agency.name if hasattr(f, 'agency') and f.agency else 'Ẩn danh')
            agreed_agencies.add(agency_name)
        else:
            active_fbs.append(f)
    
    count_agreed_agencies = len(agreed_agencies)
    
    count_accepted = 0
    count_partial = 0
    count_explained_only = 0
    count_pending = 0
    count_need_opinion = 0

    for f in active_fbs:
        if getattr(f, 'need_opinion', None) and str(f.need_opinion).strip():
            count_need_opinion += 1
        exps = f.explanations.all()
        if not exps:
            count_pending += 1
        else:
            exp_text = " ".join([e.content for e in exps if e.content]).lower()
            if 'tiếp thu một phần' in exp_text:
                count_partial += 1
            if 'tiếp thu' in exp_text:
                count_accepted += 1
            else:
                count_explained_only += 1

    total_active = len(active_fbs)
    p_stats = doc.add_paragraph()
    p_stats.paragraph_format.first_line_indent = Cm(1.27)
    p_stats.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_stats.add_run(f"1. Tổng số {total_consulted} cơ quan, tổ chức, cá nhân đã gửi xin ý kiến, tham vấn/góp ý, phản biện xã hội và tổng số {total_active} ý kiến nhận được (không bao gồm các ý kiến thống nhất hoàn toàn).").font.size = Pt(14)
    p_stats.add_run(f"\n- Số lượng cơ quan thống nhất với nội dung dự thảo Nghị định: {count_agreed_agencies} đơn vị.").font.size = Pt(14)
    p_stats.add_run(f"\n- Số ý kiến đã tiếp thu: {count_accepted} ý kiến.").font.size = Pt(14)
    p_stats.add_run(f"\n- Số ý kiến tiếp thu một phần: {count_partial} ý kiến.").font.size = Pt(14)
    p_stats.add_run(f"\n- Số ý kiến đã giải trình (không bao gồm tiếp thu): {count_explained_only} ý kiến.").font.size = Pt(14)
    p_stats.add_run(f"\n- Số ý kiến chưa có giải trình: {count_pending} ý kiến.").font.size = Pt(14)
    if count_need_opinion > 0:
        p_stats.add_run(f"\n- Số ý kiến cần xin ý kiến lãnh đạo: {count_need_opinion} ý kiến.").font.size = Pt(14)

    p_res = doc.add_paragraph()
    p_res.paragraph_format.first_line_indent = Cm(1.27)
    p_res.add_run("2. Kết quả cụ thể như sau:").font.size = Pt(14)
    doc.add_paragraph()

    # 4. Main Data Table
    fields = cfg.get('fields')
    
    # Mẫu 10 chuẩn 4 cột: Nhóm/Điều, Agency, Content, Explanation
    if not fields:
        fields = [
            {'field_key': 'noi_dung_du_thao', 'field_label': 'NHÓM VẤN ĐỀ / ĐIỀU / KHOẢN'},
            {'field_key': 'don_vi_gop_y', 'field_label': 'CHỦ THỂ GÓP Ý'},
            {'field_key': 'noi_dung_gop_y', 'field_label': 'NỘI DUNG GÓP Ý'},
            {'field_key': 'giai_trinh', 'field_label': 'Ý KIẾN TIẾP THU, GIẢI TRÌNH'},
        ]

    table = doc.add_table(rows=1, cols=len(fields))
    table.style = 'Table Grid'
    
    for i, f in enumerate(fields):
        _set_cell_style(table.cell(0, i), f['field_label'].upper(), bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Đổ dữ liệu và gộp dòng
    grouped_data = _build_grouped_data(feedbacks)
    stt_global = 1
    
    for group in grouped_data:
        first_row = None; last_row = None
        is_chung = (group['label'] == 'Chung')
        
        for fb_idx, fb in enumerate(group['feedbacks']):
            row = table.add_row()
            for col_idx, field in enumerate(fields):
                f_key = field['field_key']
                val = _get_field_value(f_key, stt_global, fb)
                
                # Logic gộp dòng cho các cột định danh (Điều/Khoản hoặc Nội dung dự thảo nếu nó đóng vai trò này)
                if f_key in ['dieu_khoan', 'noi_dung_du_thao'] and not is_chung:
                    if fb_idx == 0:
                        first_row = row
                        _set_cell_style(row.cells[col_idx], val, bold=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
                    last_row = row
                else:
                    _set_cell_style(row.cells[col_idx], val, size=10, align=WD_ALIGN_PARAGRAPH.CENTER if f_key=='stt' else None)
            
            stt_global += 1

        if first_row and last_row and first_row != last_row:
            for col_idx, field in enumerate(fields):
                if field['field_key'] in ['dieu_khoan', 'noi_dung_du_thao']:
                    first_row.cells[col_idx].merge(last_row.cells[col_idx])

    # 5. Footer (Chữ ký)
    doc.add_paragraph()
    signer_name = cfg.get('footer_signer_name', '')
    signer_title = cfg.get('footer_signer_title', 'ĐẠI DIỆN CƠ QUAN CHỦ TRÌ')
    
    f_table = doc.add_table(rows=1, cols=2)
    f_table.width = Cm(25)
    # Cột 1 trống, Cột 2 chứa chữ ký
    sign_cell = f_table.cell(0, 1)
    _set_cell_style(sign_cell, signer_title.upper(), bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
    
    p_sign = sign_cell.add_paragraph()
    p_sign.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sign.add_run("(Ký tên, đóng dấu)").font.size = Pt(14)
    
    if signer_name:
        p_name = sign_cell.add_paragraph()
        p_name.paragraph_format.space_before = Pt(60)
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run_name = p_name.add_run(signer_name.upper())
        run_name.bold = True; run_name.font.size = Pt(14)

    return _save_to_stream(doc)

def _save_to_stream(doc):
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

def _get_template_path(template_type):
    """Legacy helper cho Reports/views.py - Trả về đường dẫn template mặc định"""
    base_dir = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
    if template_type == 'custom':
        return os.path.join(base_dir, 'utils', 'template_truong_tu_chinh_v3.docx')
    return os.path.join(base_dir, 'utils', 'template_bao_cao_V2_fixed.docx')
