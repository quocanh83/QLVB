"""
Generator Mẫu 10 — Bản Tổng hợp, Giải trình, Tiếp thu Ý kiến
Tuân thủ Nghị định 30/2020/NĐ-CP về Công tác Văn thư:
  - Phông: Times New Roman
  - Cỡ chữ tiêu đề: 14pt, in hoa, đậm
  - Cỡ chữ nội dung: 13pt, bảng: 12pt
  - Lề: trái 3cm, phải 2cm, trên 2cm, dưới 2cm
"""

import docx
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io


FONT_NAME = 'Times New Roman'

# Map field_key -> hàm trích xuất dữ liệu từ feedback object
def _get_field_value(field_key, index, fb, explanation):
    """Trả về giá trị tương ứng với field_key từ 1 feedback"""
    mapping = {
        'stt': str(index),
        'noi_dung_du_thao': _get_node_text(fb),
        'noi_dung_gop_y': fb.content or '',
        'don_vi_gop_y': fb.contributing_agency or (fb.agency.name if hasattr(fb, 'agency') and fb.agency else 'Ẩn danh'),
        'giai_trinh': explanation.content if explanation else 'Chưa có nội dung giải trình.',
        'chuyen_vien': _get_assignee(fb, explanation),
        'dieu_khoan': _get_dieu_khoan(fb),
        'ghi_chu': '',
    }
    return mapping.get(field_key, '')


def _get_node_text(fb):
    if fb.node:
        content = fb.node.content or ''
        return f"{fb.node.node_label}: {content[:200]}{'...' if len(content) > 200 else ''}"
    return ''


def _get_dieu_khoan(fb):
    if fb.node:
        label = fb.node.node_label or ''
        if fb.node.parent:
            return f"{fb.node.parent.node_label}, {label}"
        return label
    return ''


def _get_assignee(fb, explanation):
    if explanation and hasattr(explanation, 'user') and explanation.user:
        return explanation.user.username
    if hasattr(fb, 'user') and fb.user:
        return fb.user.username
    return ''


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in edge_data:
                element.set(qn('w:{}'.format(key)), str(edge_data[key]))


def _set_cell_text(cell, text, bold=False, size=12, align=None):
    """Ghi text vào cell với định dạng chuẩn"""
    cell.text = ''
    p = cell.paragraphs[0]
    run = p.add_run(str(text) if text else '')
    run.font.name = FONT_NAME
    run.font.size = Pt(size)
    run.bold = bold
    # Fix font fallback cho ký tự Việt
    run._element.rPr.rFonts.set(qn('w:eastAsia'), FONT_NAME)
    if align is not None:
        p.alignment = align


def generate_mau_10(document, feedbacks, template_config=None):
    """
    Sinh file Word Mẫu 10.
    
    Args:
        document: Document model instance
        feedbacks: QuerySet feedback đã filter
        template_config: dict chứa cấu hình mẫu (tuỳ chọn)
            {
                'header_org_name': str,
                'header_org_location': str,
                'footer_signer_name': str,
                'footer_signer_title': str,
                'fields': [{'field_key': str, 'field_label': str, 'column_width_cm': float}, ...]
            }
    """
    doc = docx.Document()

    # ===== THIẾT LẬP TRANG (NĐ 30/2020) =====
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

    # Font mặc định
    style = doc.styles['Normal']
    style.font.name = FONT_NAME
    style.font.size = Pt(13)
    style.paragraph_format.line_spacing = 1.5

    # ===== CẤU HÌNH =====
    if template_config:
        org_name = template_config.get('header_org_name', 'BỘ/CƠ QUAN CHỦ TRÌ')
        org_location = template_config.get('header_org_location', 'Hà Nội')
        signer_name = template_config.get('footer_signer_name', '')
        signer_title = template_config.get('footer_signer_title', 'ĐẠI DIỆN CƠ QUAN CHỦ TRÌ')
        fields = template_config.get('fields', None)
    else:
        org_name = 'BỘ/CƠ QUAN CHỦ TRÌ'
        org_location = 'Hà Nội'
        signer_name = ''
        signer_title = 'ĐẠI DIỆN CƠ QUAN CHỦ TRÌ'
        fields = None

    # Nếu không có cấu hình fields, dùng mặc định 5 cột
    if not fields:
        fields = [
            {'field_key': 'stt', 'field_label': 'TT', 'column_width_cm': 1.0},
            {'field_key': 'noi_dung_du_thao', 'field_label': 'Nội dung dự thảo', 'column_width_cm': 4.0},
            {'field_key': 'noi_dung_gop_y', 'field_label': 'Nội dung góp ý', 'column_width_cm': 4.0},
            {'field_key': 'don_vi_gop_y', 'field_label': 'Đơn vị góp ý', 'column_width_cm': 3.0},
            {'field_key': 'giai_trinh', 'field_label': 'Ý kiến giải trình, tiếp thu', 'column_width_cm': 4.5},
        ]

    # ===== HEADER 2 CỘT =====
    header_table = doc.add_table(rows=1, cols=2)
    header_table.autofit = False

    left_cell = header_table.cell(0, 0)
    left_para = left_cell.paragraphs[0]
    left_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = left_para.add_run(f"{org_name}\n----------")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = FONT_NAME

    right_cell = header_table.cell(0, 1)
    right_para = right_cell.paragraphs[0]
    right_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = right_para.add_run("CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\nĐộc lập - Tự do - Hạnh phúc\n-----------------------")
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = FONT_NAME

    doc.add_paragraph()

    # ===== TIÊU ĐỀ (14pt, in hoa, đậm, giữa) =====
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    run = title.add_run("BẢN TỔNG HỢP, GIẢI TRÌNH, TIẾP THU Ý KIẾN\n")
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = FONT_NAME

    run2 = title.add_run(f"Đối với dự thảo: {document.project_name}")
    run2.bold = True
    run2.font.size = Pt(14)
    run2.font.name = FONT_NAME

    doc.add_paragraph()

    # ===== BẢNG NỘI DUNG (12pt, cột động) =====
    num_cols = len(fields)
    table = doc.add_table(rows=1, cols=num_cols)
    table.style = 'Table Grid'
    table.autofit = False

    # Header row
    hdr_cells = table.rows[0].cells
    for i, field in enumerate(fields):
        _set_cell_text(hdr_cells[i], field['field_label'], bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
        hdr_cells[i].width = Cm(field['column_width_cm'])

    # Data rows
    for idx, fb in enumerate(feedbacks, 1):
        explanation = fb.explanations.first() if hasattr(fb, 'explanations') else None
        row_cells = table.add_row().cells
        for col_idx, field in enumerate(fields):
            value = _get_field_value(field['field_key'], idx, fb, explanation)
            align = WD_ALIGN_PARAGRAPH.CENTER if field['field_key'] == 'stt' else None
            _set_cell_text(row_cells[col_idx], value, bold=False, size=12, align=align)
            row_cells[col_idx].width = Cm(field['column_width_cm'])

    # ===== FOOTER =====
    doc.add_paragraph()

    # Nơi ký + Ngày tháng
    from datetime import datetime
    date_str = datetime.now().strftime('%d/%m/%Y')

    date_para = doc.add_paragraph()
    date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    date_run = date_para.add_run(f"{org_location}, ngày {date_str}")
    date_run.italic = True
    date_run.font.size = Pt(13)
    date_run.font.name = FONT_NAME

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run(f"{signer_title}\n(Ký tên, đóng dấu)")
    run.bold = True
    run.font.size = Pt(13)
    run.font.name = FONT_NAME

    if signer_name:
        footer2 = doc.add_paragraph()
        footer2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        footer2.paragraph_format.space_before = Pt(48)
        run = footer2.add_run(signer_name)
        run.bold = True
        run.font.size = Pt(13)
        run.font.name = FONT_NAME

    # ===== SAVE =====
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream
