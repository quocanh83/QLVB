from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from django.db import models
from rest_framework.decorators import action
from .models import Feedback, Explanation, ActionLog, ConsultationResponse
from documents.models import Document, DocumentNode, DocumentAppendix
from django.contrib.contenttypes.models import ContentType
from .serializers import FeedbackSerializer, ConsultationResponseSerializer
import docx
import re
import os
from openai import OpenAI
from django.http import FileResponse
from .utils.v2_template_generator import generate_from_v2_template, _get_field_value
from .utils.v2_template_generator import generate_from_v2_template
from django.core.files.storage import FileSystemStorage
from django.conf import settings
import shutil
from django.utils.text import get_valid_filename
import traceback



class FeedbackViewSet(viewsets.ModelViewSet):
    queryset = Feedback.objects.all()
    serializer_class = FeedbackSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_permissions(self):
        if self.action == 'export_mau_10':
            return [permissions.AllowAny()]
        return super().get_permissions()

    def create(self, request, *args, **kwargs):
        node_id = request.data.get('node')
        appendix_id = request.data.get('appendix')
        
        document_id = None
        if node_id:
            try:
                node = DocumentNode.objects.get(id=node_id)
                document_id = node.document_id
            except DocumentNode.DoesNotExist:
                return Response({"error": "Điều/Khoản không tồn tại."}, status=404)
        elif appendix_id:
            from documents.models import DocumentAppendix
            try:
                app = DocumentAppendix.objects.get(id=appendix_id)
                document_id = app.document_id
            except DocumentAppendix.DoesNotExist:
                return Response({"error": "Phụ lục không tồn tại."}, status=404)
        
        if document_id and not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Nhập góp ý cho dự thảo này."}, status=403)
            
        return super().create(request, *args, **kwargs)

    def perform_create(self, serializer):
        serializer.save(user=self.request.user)

    def _check_permission(self, request, document_id, required_role=None):
        user = request.user
        # 1. Admin?
        is_admin = user.is_staff or user.is_superuser or (hasattr(user, 'roles') and user.roles.filter(role_name='Admin').exists())
        if is_admin: return True
        
        # 2. Lead?
        try:
            if not document_id: return False
            doc = Document.objects.get(id=document_id)
            if doc.lead == user: return True
        except (Document.DoesNotExist, ValueError, TypeError):
            pass

        # 3. Chuyên viên Role?
        return user.roles.filter(role_name__in=['Contributor', 'Explainer']).exists()
        
    def _find_best_node_match(self, document_id, label, parent_context=None):
        """
        Tìm node phù hợp nhất dựa trên label (Điều 1, Khoản 2, Điểm a...)
        Hỗ trợ phân cấp: nếu label chỉ là 'Khoản 2', sẽ ưu tiên tìm Khoản 2 thuộc parent_context (Điều x).
        """
        if not label or label == "Vấn đề khác": return None
        label = label.strip()
        
        # Thử tìm khớp hoàn toàn trước (case-insensitive)
        exact_match = DocumentNode.objects.filter(document_id=document_id, node_label__iexact=label).first()
        if exact_match: return exact_match

        # Regex linh hoạt hơn để tách Điều/Khoản/Điểm
        dieu_match = re.search(r'(?:Điều|D|Art|Art\.|Chương|Mục)\s*(\d+)', label, re.IGNORECASE)
        khoan_match = re.search(r'(?:Khoản|K|Clause|K\.)\s*(\d+)', label, re.IGNORECASE)
        diem_match = re.search(r'(?:Điểm|Diem|Item|Point|P\.|p)\s*([a-zđ0-9]+)', label, re.IGNORECASE)
        
        # Nếu nhãn chỉ là số
        if not dieu_match and not khoan_match and label.replace('.', '').isdigit():
            clean_num = label.replace('.', '').strip()
            if not parent_context:
                dieu_match = re.search(r'(\d+)', clean_num)
            else:
                khoan_match = re.search(r'(\d+)', clean_num)

        # 1. Xử lý có Điều hoặc Chương/Mục
        if dieu_match:
            d_num_int = int(dieu_match.group(1))
            possible_nodes = DocumentNode.objects.filter(document_id=document_id, node_type='Điều', node_label__icontains=str(d_num_int))
            
            node_dieu = None
            for n in possible_nodes:
                actual_nums = re.findall(r'\d+', n.node_label)
                if any(int(x) == d_num_int for x in actual_nums):
                    node_dieu = n
                    break
            
            if node_dieu:
                if khoan_match:
                    k_num_int = int(khoan_match.group(1))
                    possible_khoans = DocumentNode.objects.filter(document_id=document_id, parent_id=node_dieu.id)
                    for k in possible_khoans:
                        k_actual_nums = re.findall(r'\d+', k.node_label)
                        if any(int(x) == k_num_int for x in k_actual_nums):
                            if diem_match:
                                d_label = diem_match.group(1).lower()
                                possible_diems = DocumentNode.objects.filter(document_id=document_id, parent_id=k.id, node_label__icontains=d_label)
                                if possible_diems.exists(): return possible_diems.first()
                            return k
                return node_dieu

        # 2. Xử lý chỉ có Khoản (trong context Điều)
        if khoan_match and parent_context:
            k_num_int = int(khoan_match.group(1))
            search_parent = parent_context.id
            if parent_context.node_type != 'Điều':
                p = parent_context
                while p and p.node_type != 'Điều':
                    p = DocumentNode.objects.filter(id=p.parent_id).first()
                if p: search_parent = p.id

            possible_khoans = DocumentNode.objects.filter(document_id=document_id, parent_id=search_parent)
            for k in possible_khoans:
                k_actual_nums = re.findall(r'\d+', k.node_label)
                if any(int(x) == k_num_int for x in k_actual_nums):
                    return k
            return parent_context

        # Fallback: icontains đơn thuần
        return DocumentNode.objects.filter(document_id=document_id, node_label__icontains=label).first()

    def _find_similar_agencies(self, name, limit=5):
        from core.models import Agency
        from difflib import get_close_matches
        import unicodedata
        
        def normalize(s):
            if not s: return ""
            return unicodedata.normalize('NFC', str(s)).strip()
            
        norm_input = normalize(name)
        if not norm_input or norm_input.lower() == 'nan': return []
        
        agencies = Agency.objects.all()
        agency_map = {normalize(a.name): a for a in agencies}
        
        if norm_input in agency_map:
            a = agency_map[norm_input]
            return [{"id": a.id, "name": a.name, "is_exact": True}]
        
        # Case-insensitive exact match
        agency_map_lower = {k.lower(): v for k, v in agency_map.items()}
        if norm_input.lower() in agency_map_lower:
            a = agency_map_lower[norm_input.lower()]
            return [{"id": a.id, "name": a.name, "is_exact": True}]

        matches = get_close_matches(norm_input, agency_map.keys(), n=limit, cutoff=0.5)
        return [{"id": agency_map[m].id, "name": agency_map[m].name, "is_exact": False} for m in matches]

    @action(detail=False, methods=['post'])
    def analyze_import(self, request):
        document_id = request.data.get('document_id')
        file_obj = request.FILES.get('file')
        gs_url = request.data.get('google_sheets_url')
        
        if not file_obj and not gs_url:
            return Response({"error": "Không tìm thấy tệp tải lên hoặc đường dẫn Google Sheets."}, status=400)
            
        try:
            import pandas as pd
            import io
            import requests
            
            if gs_url:
                # Regex trích xuất sheet_id và gid
                sheet_id_match = re.search(r'/d/([a-zA-Z0-9-_]+)', gs_url)
                gid_match = re.search(r'gid=([0-9]+)', gs_url)
                if not sheet_id_match:
                    return Response({"error": "Đường dẫn Google Sheets không hợp lệ."}, status=400)
                
                sheet_id = sheet_id_match.group(1)
                gid = gid_match.group(1) if gid_match else '0'
                export_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/export?format=xlsx&gid={gid}"
                
                response = requests.get(export_url)
                if response.status_code != 200:
                    return Response({"error": f"Không thể tải dữ liệu từ Google Sheets (Status: {response.status_code}). Hãy đảm bảo sheet ở chế độ công khai."}, status=400)
                
                df = pd.read_excel(io.BytesIO(response.content))
            else:
                if file_obj.name.endswith('.xlsx') or file_obj.name.endswith('.xls'):
                    df = pd.read_excel(file_obj)
                else:
                    df = pd.read_csv(file_obj, encoding='utf-8-sig')
            
            # Map columns based on keywords (9 columns support)
            col_map = {}
            # Thứ tự ưu tiên quan trọng
            for col in df.columns:
                c = str(col).lower()
                if any(kw in c for kw in ['điều', 'khoản', 'mục', 'phạm vi', 'vị trí']): col_map['node'] = col
                if any(kw in c for kw in ['cơ quan', 'chủ thể', 'đơn vị']): col_map['agency'] = col
                if any(kw in c for kw in ['số văn bản', 'số hiệu']): col_map['doc_number'] = col
                if any(kw in c for kw in ['ngày', 'thời gian']): col_map['doc_date'] = col
                if any(kw in c for kw in ['giải trình', 'tiếp thu']): col_map['explanation'] = col
                if any(kw in c for kw in ['lý do', 'cơ sở']): col_map['reason'] = col
                if any(kw in c for kw in ['ghi chú', 'note']): col_map['note'] = col
            
            # Map cột Nội dung riêng để tránh nhầm với 'Số văn bản góp ý'
            for col in df.columns:
                c = str(col).lower()
                # Ưu tiên có chữ 'nội dung'
                if 'nội dung' in c and any(kw in c for kw in ['góp ý', 'tham vấn', 'phản biện']):
                    col_map['content'] = col
                    break
            
            # Nếu chưa tìm thấy nội dung thì mới dùng từ khóa 'góp ý' chung, 
            # nhưng loại trừ cột đã map là 'số văn bản'
            if 'content' not in col_map:
                for col in df.columns:
                    c = str(col).lower()
                    if 'góp ý' in c and col != col_map.get('doc_number'):
                        col_map['content'] = col
                        break
            
            if 'content' not in col_map:
                return Response({"error": "Không tìm thấy cột 'Nội dung góp ý' trong bảng dữ liệu."}, status=400)

            results = []
            
            # Lấy trước dữ liệu để kiểm tra trùng lặp và giải trình
            feedbacks_qs = Feedback.objects.filter(document_id=document_id).prefetch_related('explanations')
            
            for index, row in df.iterrows():
                raw_node = str(row.get(col_map.get('node'), '')).strip()
                raw_agency = str(row.get(col_map.get('agency'), '')).strip()
                raw_content = str(row.get(col_map['content'], '')).strip()
                raw_reason = str(row.get(col_map.get('reason'), '')).strip()
                raw_note = str(row.get(col_map.get('note'), '')).strip()
                raw_explanation = str(row.get(col_map.get('explanation'), '')).strip()
                raw_doc_num = str(row.get(col_map.get('doc_number'), '')).strip()
                
                # Format date string (only Date)
                val_doc_date = row.get(col_map.get('doc_date'), '')
                raw_doc_date = ""
                if pd.notnull(val_doc_date):
                    if hasattr(val_doc_date, 'strftime'):
                        raw_doc_date = val_doc_date.strftime('%d/%m/%Y')
                    else:
                        raw_doc_date = str(val_doc_date).split(' ')[0]
                
                if not raw_content or raw_content.lower() in ['nan', 'none']: continue
                
                # MỚI: Nếu ND và GT đều là OK (hoặc ND là OK và GT trống) thì bỏ qua
                def is_ok_or_empty(val):
                    v = str(val).strip().upper()
                    return v == "OK" or v in ["", "NAN", "NONE"]

                content_val = str(row.get(col_map['content'], '')).strip().upper()
                exp_val = str(row.get(col_map.get('explanation'), '')).strip().upper()

                if content_val == "OK":
                    # Nếu nội dung đã OK, và (giải trình cũng OK hoặc không có giải trình) thì skip
                    if exp_val in ["OK", "", "NAN", "NONE"]:
                        continue
                
                # Logic trùng lặp (vẫn dùng full_content để so khớp với dữ liệu cũ nếu cần, 
                # hoặc chuyển sang so khớp từng trường)
                full_content_for_dup_check = raw_content
                if raw_reason and raw_reason.lower() not in ['nan', 'none']:
                    full_content_for_dup_check += f" - Lý do: {raw_reason}"
                if raw_note and raw_note.lower() not in ['nan', 'none']:
                    full_content_for_dup_check += f" - Ghi chú: {raw_note}"
                
                # Khớp cơ quan
                agency_matches = self._find_similar_agencies(raw_agency)
                matched_agency = agency_matches[0] if agency_matches and agency_matches[0]['is_exact'] else None
                agency_id = matched_agency['id'] if matched_agency else None
                
                # Khớp Node/Appendix
                suggested_node = self._find_best_node_match(document_id, raw_node)
                appendix_id = None
                if not suggested_node and 'phụ lục' in raw_node.lower():
                    from documents.models import DocumentAppendix
                    app_match = re.search(r'Phụ\s+lục\s*([IVXLCDM\d]+)', raw_node, re.IGNORECASE)
                    if app_match:
                        app_name_part = app_match.group(1).upper()
                        app_obj = DocumentAppendix.objects.filter(document_id=document_id, name__icontains=app_name_part).first()
                        if app_obj: appendix_id = app_obj.id
                # 1. Kiểm tra trùng lặp Góp ý
                is_duplicate = False
                duplicate_id = None
                existing_fb = None
                
                target_node_id = suggested_node.id if suggested_node else None
                
                # Query feedbacks with SAME content, agency, and document
                # (Lưu ý: Feedback model đã import global ở trên)
                f_qs = Feedback.objects.filter(
                    document_id=document_id,
                    content=raw_content,
                    agency_id=agency_id
                )
                
                for fb in f_qs:
                    is_duplicate = True
                    duplicate_id = fb.id
                    existing_fb = fb
                    break
                
                # 2. Logic xử lý Giải trình (Explanation)
                explanation_status = "new"
                existing_exp_content = ""
                
                if raw_explanation and raw_explanation.lower() not in ['nan', 'none']:
                    if existing_fb:
                        # Tìm giải trình hiện có
                        exp = existing_fb.explanations.first()
                        if exp:
                            existing_exp_content = exp.content
                            if exp.content.strip() == raw_explanation.strip():
                                explanation_status = "identical"
                            else:
                                explanation_status = "conflict"
                        else:
                            explanation_status = "new"
                else:
                    explanation_status = "none"

                results.append({
                    "key": f"import-{index}",
                    "original_node": raw_node,
                    "node_id": target_node_id,
                    "node_label": suggested_node.node_label if suggested_node else raw_node,
                    "appendix_id": appendix_id,
                    "original_agency": raw_agency,
                    "agency_id": agency_id,
                    "agency_name": matched_agency['name'] if matched_agency else raw_agency,
                    "agency_suggestions": agency_matches,
                    "content": raw_content,
                    "reason": raw_reason if raw_reason.lower() not in ['nan', 'none'] else "",
                    "note": raw_note if raw_note.lower() not in ['nan', 'none'] else "",
                    # MỚI: Logic thông minh cho Chế độ Lưu (import_mode)
                    # Nếu trùng nội dung nhưng giải trình mới hoặc khác -> "Chỉ nhập giải trình"
                    # Ngược lại nếu trùng -> "skip" (Bỏ qua)
                    "is_duplicate": is_duplicate,
                    "duplicate_id": duplicate_id,
                    "import_mode": "explanation_only" if (is_duplicate and explanation_status in ["new", "conflict"]) else ("skip" if is_duplicate else "add_new"),
                    
                    # Thông tin giải trình
                    "explanation_content": raw_explanation if explanation_status != "none" else "",
                    "explanation_status": explanation_status,
                    "existing_explanation": existing_exp_content,
                    "explanation_import_mode": "overwrite" if explanation_status in ["new", "conflict"] else "skip",
                    
                    # Thông tin công văn (ConsultationResponse)
                    "official_number": raw_doc_num if raw_doc_num.lower() not in ['nan', 'none'] else "",
                    "official_date": raw_doc_date if raw_doc_date.lower() not in ['nan', 'none'] else "",
                    "row_index": index + 2, # GSheet row starts at 2 (1 is header)
                })
                
            return Response({"rows": results})
            
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def confirm_import(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Nhập góp ý cho dự thảo này."}, status=403)
            
        rows = request.data.get('rows', [])
        created_count = 0
        updated_count = 0
        skipped_count = 0
        exp_count = 0
        consultation_count = 0
        
        try:
            doc = Document.objects.get(id=document_id)
        except Document.DoesNotExist:
            return Response({"error": "Dự thảo không tồn tại."}, status=404)
        
        for row in rows:
            mode = row.get('import_mode', 'add_new')
            node_id = row.get('node_id')
            appendix_id = row.get('appendix_id')
            agency_id = row.get('agency_id')
            content = row.get('content', '')
            reason = row.get('reason', '')
            note = row.get('note', '')
            duplicate_id = row.get('duplicate_id')
            
            # 1. Xử lý Feedback (Góp ý)
            feedback_obj = None
            if mode == 'skip' or (mode == 'add_if_diff' and row.get('is_duplicate')):
                skipped_count += 1
                if duplicate_id:
                    feedback_obj = Feedback.objects.filter(id=duplicate_id).first()
            elif mode == 'explanation_only' and duplicate_id:
                feedback_obj = Feedback.objects.filter(id=duplicate_id).first()
                # Không update content, chỉ để feedback_obj có giá trị cho bước sau
            elif mode == 'overwrite' and duplicate_id:
                Feedback.objects.filter(id=duplicate_id).update(
                    node_id=node_id,
                    appendix_id=appendix_id,
                    agency_id=agency_id,
                    content=content,
                    reason=reason,
                    note=note,
                    user=request.user
                )
                feedback_obj = Feedback.objects.get(id=duplicate_id)
                updated_count += 1
            else:
                # add_new (hoặc add_if_diff mà không trùng)
                feedback_obj = Feedback.objects.create(
                    document_id=document_id,
                    node_id=node_id,
                    appendix_id=appendix_id,
                    user=request.user,
                    agency_id=agency_id,
                    content=content,
                    reason=reason,
                    note=note
                )
                created_count += 1

            # 2. Xử lý Giải trình (Explanation)
            exp_content = row.get('explanation_content', '')
            exp_mode = row.get('explanation_import_mode', 'skip')
            
            if feedback_obj and exp_content and exp_mode == 'overwrite':
                # Xóa giải trình cũ nếu có
                feedback_obj.explanations.all().delete()
                # Tạo giải trình mới
                Explanation.objects.create(
                    target_type='Feedback',
                    content_type=ContentType.objects.get_for_model(Feedback),
                    object_id=feedback_obj.id,
                    content=exp_content,
                    user=request.user
                )
                exp_count += 1
                row['has_imported_explanation'] = True
            elif feedback_obj and exp_content and exp_mode == 'skip':
                # Nếu đã có giải trình thì đánh dấu là đã có (không ghi đè nhưng vẫn coi là đã nhập)
                if feedback_obj.explanations.exists():
                    row['has_imported_explanation'] = True

            if feedback_obj:
                row['has_imported_content'] = True

            # 3. Xử lý Công văn (ConsultationResponse)
            doc_num = row.get('official_number', '')
            doc_date_raw = row.get('official_date', '')
            
            if agency_id and doc_num and str(doc_num).lower() != 'nan':
                # Parse date if possible
                official_date = None
                if doc_date_raw and str(doc_date_raw).lower() != 'nan':
                    try:
                        from datetime import datetime
                        # Clean date string (remove time if exists)
                        date_str = str(doc_date_raw).split(' ')[0]
                        # Thử các định dạng phổ biến
                        for fmt in ('%d/%m/%Y', '%Y-%m-%d', '%d-%m-%Y'):
                            try:
                                official_date = datetime.strptime(date_str, fmt).date()
                                break
                            except: continue
                    except: pass
                
                # Tạo hoặc cập nhật ConsultationResponse (Số hiệu công văn của đơn vị này)
                cr, created = ConsultationResponse.objects.get_or_create(
                    document=doc,
                    agency_id=agency_id,
                    official_number=doc_num,
                    defaults={'official_date': official_date}
                )
                if not created and official_date:
                    cr.official_date = official_date
                    cr.save()
                
                if created: consultation_count += 1

        # MỚI: Cập nhật Google Sheets nếu có link
        gs_url = request.data.get('gs_url')
        
        # MỚI: Lưu lại đường link vào Dự thảo nếu được yêu cầu
        save_gs_flag = request.data.get('save_gs_url')
        print(f"DEBUG: gs_url={gs_url}, save_gs_flag={save_gs_flag}, doc_id={document_id}")
        
        if gs_url and str(save_gs_flag).lower() in ['true', 'on', '1', 'yes'] and document_id:
            try:
                Document.objects.filter(id=document_id).update(google_sheets_url=gs_url)
                print(f"✅ Đã cập nhật link Google Sheets vào Dự thảo: {document_id}")
            except Exception as e:
                print(f"❌ Lỗi khi cập nhật link GS vào Dự thảo: {str(e)}")

        if gs_url and rows:
            self._update_google_sheet_status(gs_url, rows)

        return Response({
            "message": f"Nhập hoàn tất: {created_count} góp ý mới, {updated_count} ghi đè. Đã xử lý {exp_count} giải trình và {consultation_count} công văn.",
            "created": created_count,
            "updated": updated_count,
            "skipped": skipped_count,
            "explanations": exp_count,
            "consultations": consultation_count
        })

    @action(detail=False, methods=['post'])
    def delete_all(self, request):
        document_id = request.data.get('document_id')
        if not document_id:
            return Response({"error": "Chưa chọn dự thảo để xóa."}, status=400)
            
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền xóa toàn bộ góp ý của dự thảo này."}, status=403)
            
        count = Feedback.objects.filter(document_id=document_id).count()
        Feedback.objects.filter(document_id=document_id).delete()
        
        return Response({"message": f"Đã xóa thành công toàn bộ {count} nội dung góp ý của dự thảo."}, status=200)

    def _get_gsheet_worksheet(self, gs_url):
        try:
            import gspread
            from google.oauth2.service_account import Credentials
            import re
            
            # 1. Auth (Dùng credentials file)
            scopes = ['https://www.googleapis.com/auth/spreadsheets']
            key_path = os.path.join(settings.BASE_DIR, 'keys.json')
            if not os.path.exists(key_path):
                key_path = os.path.join(settings.BASE_DIR, 'google_keys.json')
                
            if not os.path.exists(key_path):
                return None, f"File keys.json không tồn tại tại {settings.BASE_DIR}"
                
            try:
                creds = Credentials.from_service_account_file(key_path, scopes=scopes)
                client = gspread.authorize(creds)
            except Exception as auth_e:
                return None, f"Lỗi xác thực Google API: {str(auth_e)}"
            
            # 2. Open Sheet (Lấy ID và GID từ link)
            match = re.search(r'/spreadsheets/d/([a-zA-Z0-9-_]+)', gs_url)
            if not match: 
                return None, "Link Google Sheet không hợp lệ (không tìm thấy Spreadsheet ID)."
            
            sheet_id = match.group(1)
            
            gid_match = re.search(r'gid=([0-9]+)', gs_url)
            gid = gid_match.group(1) if gid_match else "0"
            
            try:
                spreadsheet = client.open_by_key(sheet_id)
            except Exception as open_e:
                return None, f"Không thể mở Spreadsheet (ID: {sheet_id}). Hãy kiểm tra quyền chia sẻ với công tác viên. Chi tiết: {str(open_e)}"

            # Tìm worksheet theo GID
            worksheet = None
            try:
                for ws in spreadsheet.worksheets():
                    if str(ws.id) == gid:
                        worksheet = ws
                        break
                if not worksheet: 
                    worksheet = spreadsheet.get_worksheet(0)
            except Exception as ws_e:
                 return None, f"Lỗi khi truy cập Worksheet: {str(ws_e)}"
            
            return worksheet, None
        except Exception as e:
            return None, f"Lỗi hệ thống khi kết nối Google Sheets: {str(e)}"

    def _update_google_sheet_status(self, gs_url, rows):
        try:
            import gspread
            worksheet, err = self._get_gsheet_worksheet(gs_url)
            if not worksheet:
                print(f"❌ _update_google_sheet_status failed: {err}")
                return

            # 3. Headers check & Identify columns
            headers = worksheet.row_values(1)
            
            def get_or_create_col(name):
                try:
                    return headers.index(name) + 1
                except ValueError:
                    new_col = len(headers) + 1
                    worksheet.update_cell(1, new_col, name)
                    headers.append(name)
                    return new_col

            col_content = get_or_create_col("ND")
            col_exp = get_or_create_col("GT")
            
            # 4. Batch Update (Ghi giá trị 'OK' cho các dòng tương ứng)
            cells_to_update = []
            for row in rows:
                r_idx = row.get('row_index')
                if not r_idx: continue
                
                if row.get('has_imported_content'):
                    cells_to_update.append(gspread.Cell(row=r_idx, col=col_content, value="OK"))
                if row.get('has_imported_explanation'):
                    cells_to_update.append(gspread.Cell(row=r_idx, col=col_exp, value="OK"))
            
            if cells_to_update:
                worksheet.update_cells(cells_to_update)
                print(f"✅ Đã cập nhật xong {len(cells_to_update)} trạng thái vào Google Sheets: {gs_url}")
            else:
                print("⚠️ Không có dòng nào cần cập nhật trạng thái vào Google Sheets.")
                
        except Exception as e:
            print(f"❌ Lỗi khi cập nhật Google Sheets: {str(e)}")
            import traceback
            traceback.print_exc()

    @action(detail=False, methods=['post'])
    def parse_file(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Nhập góp ý cho dự thảo này."}, status=403)
        
        file_obj = request.FILES.get('file')
        agency_default = request.data.get('contributing_agency', '')
        agency_id = request.data.get('agency_id')
        
        if not file_obj:
            return Response({"error": "No file uploaded"}, status=400)
            
        doc = docx.Document(file_obj)
        results = []
        metadata = {
            "drafting_agency": "",
            "agency_location": "",
            "total_consulted_doc": 0,
            "total_feedbacks_doc": 0
        }
        
        # Sơ bộ tìm metadata ở 20 đoạn đầu
        for i, para in enumerate(doc.paragraphs[:20]):
            text = para.text.strip()
            if not text: continue
            
            loc_match = re.search(r'^([\w\s]+),\s*ngày\s+\d+\s+tháng\s+\d+\s+năm\s+\d+', text)
            if loc_match and not metadata["agency_location"]:
                metadata["agency_location"] = loc_match.group(1).strip()
            
            if i < 7 and any(kw in text.upper() for kw in ["BỘ ", "TỔNG CỤC ", "ỦY BAN ", "SỞ "]):
                if not metadata["drafting_agency"]:
                    metadata["drafting_agency"] = text
            
            consult_match = re.search(r'tổng số cơ quan.*?là\s*(\d+)', text.lower()) or re.search(r'lấy ý kiến.*?:?\s*(\d+)', text.lower())
            if consult_match:
                metadata["total_consulted_doc"] = int(consult_match.group(1))
            
            feedback_match = re.search(r'tổng số ý kiến.*?là\s*(\d+)', text.lower()) or re.search(r'số ý kiến nhận được.*?:?\s*(\d+)', text.lower())
            if feedback_match:
                metadata["total_feedbacks_doc"] = int(feedback_match.group(1))

        # Phân rã từ BẢNG (thường có trong phụ lục)
        table_context_node = None
        for table in doc.tables:
            # Kiểm tra xem bảng có phải bảng góp ý không (dựa vào header)
            header_cells = [cell.text.strip().lower() for cell in table.rows[0].cells]
            if not any(kw in " ".join(header_cells) for kw in ["góp ý", "ý kiến", "nội dung", "điều", "khoản"]):
                continue
            
            # Tìm cột chứa nội dung và cột chứa node_label
            content_col_idx = -1
            node_col_idx = -1
            agency_col_idx = -1
            
            for idx, text in enumerate(header_cells):
                if any(kw in text for kw in ["góp ý", "nội dung", "ý kiến"]): content_col_idx = idx
                if any(kw in text for kw in ["điều", "khoản", "mục", "vị trí"]): node_col_idx = idx
                if any(kw in text for kw in ["cơ quan", "đơn vị", "chủ thể"]): agency_col_idx = idx
            
            # Nếu không tìm thấy cột nội dung, lấy cột rộng nhất hoặc mặc định cột 2
            if content_col_idx == -1: continue
            
            for row in table.rows[1:]: # Bỏ qua header
                cells = row.cells
                if len(cells) <= content_col_idx: continue
                
                row_content = cells[content_col_idx].text.strip()
                if not row_content: continue
                
                row_node_label = cells[node_col_idx].text.strip() if node_col_idx != -1 and len(cells) > node_col_idx else ""
                row_agency = cells[agency_col_idx].text.strip() if agency_col_idx != -1 and len(cells) > agency_col_idx else agency_default
                
                # Tìm node_id tự động
                suggested_node = self._find_best_node_match(document_id, row_node_label, parent_context=table_context_node)
                if suggested_node and suggested_node.node_type in ['Điều', 'Khoản']:
                    table_context_node = suggested_node # Cập nhật context cho dòng sau
                
                results.append({
                    "key": f"table-{len(results)}",
                    "node_label": row_node_label if row_node_label else "Vấn đề khác",
                    "node_id": suggested_node.id if suggested_node else None,
                    "contributing_agency": row_agency,
                    "content": row_content
                })

        # Phân rã nội dung từ Văn bản (Paragraphs)
        current_node_label = "Vấn đề khác"
        current_node_obj = None # Context cho phân cấp
        current_content = ""
        
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text: continue
            
            node_match = re.search(r'(Điều\s+\d+|Khoản\s+\d+|Điểm\s+[a-zđ])', text, re.IGNORECASE)
            
            if node_match:
                if current_content.strip():
                    results.append({
                        "key": f"para-{len(results)}",
                        "node_label": current_node_label,
                        "node_id": current_node_obj.id if current_node_obj else None,
                        "contributing_agency": agency_default,
                        "agency_id": agency_id,
                        "content": current_content.strip()
                    })
                
                new_label = node_match.group(1).title()
                # Thử tìm node và cập nhật context
                new_node = self._find_best_node_match(document_id, new_label, parent_context=current_node_obj if 'Điều' not in new_label else None)
                
                current_node_label = new_label
                current_node_obj = new_node
                current_content = text
            else:
                if current_content:
                    current_content += f"\n{text}"
                else:
                    current_content = text
                    
        if current_content.strip():
            results.append({
                "key": f"para-{len(results)}",
                "node_label": current_node_label,
                "node_id": current_node_obj.id if current_node_obj else None,
                "contributing_agency": agency_default,
                "agency_id": agency_id,
                "content": current_content.strip()
            })
            
        return Response({
            "metadata": metadata,
            "feedbacks": results
        })

    @action(detail=False, methods=['post'])
    def ocr_parse(self, request):
        """
        Xử lý OCR cho file Ảnh/PDF, trả về văn bản đã qua AI sửa lỗi và highlight diff.
        """
        try:
            document_id = request.data.get('document_id')
            if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
                return Response({"error": "Bạn không có quyền hoặc ID văn bản không hợp lệ."}, status=403)
            
            file_obj = request.FILES.get('file')
            if not file_obj:
                return Response({"error": "Không tìm thấy tệp tin được tải lên."}, status=400)
                
            # Save temporary file
            fs = FileSystemStorage(location=os.path.join(settings.MEDIA_ROOT, 'uploads_ocr'))
            
            # Làm sạch tên file để tránh lỗi tiếng Việt/ký tự lạ trên Windows
            safe_name = get_valid_filename(file_obj.name)
            print(f"--- [OCR API] Đang lưu tệp tin: {safe_name} ---")
            filename = fs.save(safe_name, file_obj)
            file_path = fs.path(filename)
            
            from .utils.ocr_service import OCRService # Lazy load
            print(f"--- [OCR API] Bắt đầu xử lý AI cho {filename} (đợi nạp bộ não)... ---")
            
            selected_pages = request.data.get('selected_pages') # VD: "2, 4" hoặc "1-3, 5"
            
            service = OCRService()
            ocr_results = service.process_file(file_path, target_pages=selected_pages)
            print(f"--- [OCR API] Bóc tách hoàn tất {len(ocr_results)} trang. ---")
            
            # Clean up raw upload if desired, or keep for audit
            # os.remove(file_path) 
            
            return Response({
                "pages": ocr_results
            })
        except Exception as e:
            import traceback
            with open('error_log_ocr.txt', 'w', encoding='utf-8') as f:
                traceback.print_exc(file=f)
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def get_pdf_info(self, request):
        """
        Lấy thông tin số trang và tạo thumbnails cho file PDF/Ảnh.
        """
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "No file"}, status=400)
            
        fs = FileSystemStorage(location=os.path.join(settings.MEDIA_ROOT, 'temp_previews'))
        filename = fs.save(get_valid_filename(file_obj.name), file_obj)
        file_path = fs.path(filename)
        ext = os.path.splitext(file_path)[1].lower()
        
        previews = []
        total_pages = 0
        
        import fitz
        try:
            if ext == '.pdf':
                doc = fitz.open(file_path)
                total_pages = len(doc)
                # Chỉ tạo thumbnail cho tất cả các trang
                for i in range(total_pages):
                    page = doc.load_page(i)
                    # Tạo ảnh nhỏ (thumbnail)
                    pix = page.get_pixmap(matrix=fitz.Matrix(0.5, 0.5)) 
                    img_name = f"thumb_{os.urandom(4).hex()}_{i}.jpg"
                    img_p = os.path.join(fs.location, img_name)
                    pix.save(img_p)
                    rel_p = os.path.relpath(img_p, settings.MEDIA_ROOT).replace('\\', '/')
                    previews.append(settings.MEDIA_URL + rel_p)
                doc.close()
            else:
                total_pages = 1
                rel_p = os.path.relpath(file_path, settings.MEDIA_ROOT).replace('\\', '/')
                previews.append(settings.MEDIA_URL + rel_p)
                
            return Response({
                "total_pages": total_pages,
                "previews": previews,
                "temp_file_path": filename # Để frontend dùng lại nếu cần (tuy nhiên hiện tại vẫn upload lại ocr_parse cho đơn giản)
            })
        except Exception as e:
            return Response({"error": str(e)}, status=500)
        finally:
            # Lưu ý: Không xóa file ở đây vì preview cần URL. 
            # Hệ thống nên có job dọn dẹp temp định kỳ.
            pass

    @action(detail=False, methods=['post'])
    def ocr_finalize_parse(self, request):
        """
        Giai đoạn cuối: biến văn bản sau OCR thành danh sách góp ý có cấu trúc.
        """
        document_id = request.data.get('document_id')
        text = request.data.get('text', '')
        
        if not text:
            return Response({"error": "Văn bản trống"}, status=400)
            
        from .utils.ai_parser import AIParser
        parser = AIParser()
        raw_feedbacks = parser.parse_text_to_feedbacks(text)
        
        if not raw_feedbacks:
            return Response({"error": "Không thể phân tích văn bản này bằng AI."}, status=500)
            
        # Làm sạch và khớp NodeID
        results = []
        for i, f in enumerate(raw_feedbacks):
            node_label = f.get('node_label', 'Chung')
            content = f.get('content', '')
            agency = f.get('agency', '')
            
            # Sử dụng logic khớp thông minh của hệ thống
            suggested_node = self._find_best_node_match(document_id, node_label)
            
            results.append({
                "key": f"ocr-finalize-{i}-{os.urandom(4).hex()}",
                "node_label": suggested_node.node_label if suggested_node else node_label,
                "node_id": suggested_node.id if suggested_node else None,
                "contributing_agency": agency or "",
                "content": content
            })
            
        return Response({
            "feedbacks": results
        })

    @action(detail=False, methods=['get'])
    def get_document_nodes(self, request):
        """Lấy danh sách node và Phụ lục để phục vụ mapping ở frontend với nhãn đầy đủ"""
        doc_id = request.query_params.get('document_id')
        if not doc_id: return Response([])
        
        nodes = DocumentNode.objects.filter(document_id=doc_id).select_related('parent', 'parent__parent').order_by('order_index')
        
        results = []
        for n in nodes:
            # Tạo nhãn phân tầng: Điều 1 > Khoản 2 > Điểm a
            path_parts = []
            curr = n
            while curr:
                if curr.node_type != 'Văn bản':
                    path_parts.insert(0, curr.node_label)
                curr = curr.parent
            
            results.append({
                "unique_id": f"node-{n.id}",
                "id": n.id, 
                "label": " > ".join(path_parts) if path_parts else n.node_label, 
                "type": n.node_type
            })
            
        # Thêm Phụ lục từ bảng DocumentAppendix
        from documents.models import DocumentAppendix
        appendices = DocumentAppendix.objects.filter(document_id=doc_id).order_by('name')
        for app in appendices:
            # Rút gọn tên: "Phụ lục I: ABC" -> "Phụ lục I"
            display_name = app.name
            if ':' in display_name:
                display_name = display_name.split(':')[0].strip()
            if '\n' in display_name:
                display_name = display_name.split('\n')[0].strip()
            
            # Đảm bảo nhãn hiển thị đẹp (Phụ lục I, Phụ lục II...)
            if not display_name.lower().startswith('phụ lục'):
                display_name = f"Phụ lục {display_name}"
            else:
                # Chuẩn hóa viết hoa chữ đầu
                display_name = "Phụ lục" + display_name[7:]
                
            results.append({
                "unique_id": f"app-{app.id}",
                "id": app.id,
                "label": display_name,
                "type": "Appendix"
            })
            
        return Response(results)

    @action(detail=False, methods=['post'])
    def bulk_create(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Góp ý'):
            return Response({"error": "Bạn không có quyền Nhập góp ý cho dự thảo này."}, status=403)
            
        feedbacks_data = request.data.get('feedbacks', [])
        metadata = request.data.get('metadata', {})
        
        # Cập nhật Document Metadata
        from documents.models import Document
        doc = Document.objects.filter(id=document_id).first()
        if doc and metadata:
            if metadata.get('drafting_agency'): doc.drafting_agency = metadata['drafting_agency']
            if metadata.get('agency_location'): doc.agency_location = metadata['agency_location']
            if metadata.get('total_consulted_doc'): doc.total_consulted_doc = metadata['total_consulted_doc']
            if metadata.get('total_feedbacks_doc'): doc.total_feedbacks_doc = metadata['total_feedbacks_doc']
            doc.save()

        # Tìm node làm fallback (Ưu tiên nhãn "Chung")
        fallback_node = DocumentNode.objects.filter(document_id=document_id, node_label='Chung').first()
        if not fallback_node:
            fallback_node = DocumentNode.objects.filter(document_id=document_id, node_type='Vấn đề khác').first()

        created_count = 0
        for item in feedbacks_data:
            node_id = item.get('node_id') # Ưu tiên id nếu frontend đã map
            node_label = item.get('node_label', '')
            
            node = None
            if node_id:
                node = DocumentNode.objects.filter(id=node_id, document_id=document_id).first()
            
            if not node and node_label:
                # Tìm theo label (icontains)
                node = DocumentNode.objects.filter(document_id=document_id, node_label__icontains=node_label).first()
            
            # Fallback
            target_node = node or fallback_node
            
            if target_node:
                # Ưu tiên lấy cơ quan góp ý và số công văn từ metadata (chung cho cả lần nhập)
                global_agency = metadata.get('contributing_agency')
                global_agency_id = metadata.get('agency_id')
                global_doc_number = metadata.get('official_doc_number')

                Feedback.objects.create(
                    document_id=document_id,
                    node=target_node,
                    user=request.user,
                    contributing_agency=global_agency or item.get('contributing_agency', ''),
                    agency_id=global_agency_id or item.get('agency_id'),
                    official_doc_number=global_doc_number or item.get('official_doc_number', ''),
                    content=item.get('content', ''),
                    reason=item.get('reason', ''),
                    note=item.get('note', '')
                )
                created_count += 1
                
        return Response({"message": f"Đã lưu thành công {created_count} góp ý.", "count": created_count})

    @action(detail=False, methods=['post'])
    def ai_suggest(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Giải trình'):
            return Response({"error": "Bạn không có quyền sử dụng AI Giải trình cho dự thảo này."}, status=403)

        node_content = request.data.get('node_content', '')
        feedback_content = request.data.get('feedback_content', '')
        
        system_prompt = """Bạn là một Chuyên viên Pháp chế cấp cao của cơ quan Nhà nước trực thuộc Chính phủ Việt Nam.
NGUYÊN TẮC LẬP LUẬN:
- Bạn mang tâm thế của cơ quan chủ trì soạn thảo. ƯU TIÊN MẶC ĐỊNH là bảo vệ, giữ nguyên nội dung dự thảo ban đầu, trừ khi ý kiến góp ý chỉ ra sai sót pháp lý nghiêm trọng không thể chối cãi.
- Lập luận từ chối tiếp thu phải cực kỳ khéo léo, khách quan, hợp tình hợp lý. Dựa vào các lý do như: tính khả thi, đánh giá tác động thủ tục hành chính, bối cảnh thực tiễn...
- Văn phong mang ĐẬM TÍNH NGÔN NGỮ HÀNH CHÍNH NHÀ NƯỚC.
FORMAT TRẢ LỜI CỐ ĐỊNH:
- KHÔNG XƯNG HÔ "tôi", "bạn", "chúng tôi", "AI". Chỉ dùng đại từ "Cơ quan chủ trì soạn thảo".
- KHÔNG CÓ lời chào hỏi, dạo đầu. TRẢ VỀ TRỰC TIẾP kết quả giải trình để điền thẳng vào khung báo cáo."""
        
        user_prompt = f"Nội dung gốc:\n{node_content}\n\nNội dung góp ý:\n{feedback_content}\n\nHãy đưa ra nội dung giải trình ngắn gọn, đanh thép để bảo vệ dự thảo."
        
        from core.models import SystemSetting
        db_setting = SystemSetting.objects.filter(key='OPENAI_API_KEY').first()
        api_key = db_setting.value if db_setting and db_setting.value else os.environ.get('OPENAI_API_KEY')
        
        if not api_key:
            return Response({
                "suggestion": "Cơ quan chủ trì soạn thảo xin ghi nhận ý kiến. Tuy nhiên, qua rà soát bối cảnh thực tiễn và tính bộ khung của toàn Hệ thống điều khoản, kính đề nghị giữ nguyên như dự thảo để đảm bảo tính khả thi.\n\n(Hệ thống đang chạy MOCK AI do Môi trường Backend chưa nạp biến OPENAI_API_KEY. Để nạp key, vui lòng Set Environment OPENAI_API_KEY=sk-...)"
            })
            
        try:
            client = OpenAI(api_key=api_key)
            response = client.chat.completions.create(
                model="gpt-3.5-turbo",
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ],
                temperature=0.7,
                max_tokens=600
            )
            return Response({
                "suggestion": response.choices[0].message.content.strip()
            })
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def save_explanation(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Giải trình'):
            return Response({"error": "Bạn không có quyền Lưu giải trình cho dự thảo này."}, status=403)
            
        target_type = request.data.get('target_type')
        object_id = request.data.get('object_id')
        content = request.data.get('content')
        
        if target_type == 'Feedback':
            content_type = ContentType.objects.get_for_model(Feedback)
        else:
            content_type = ContentType.objects.get_for_model(DocumentNode)
            
        explanation, created = Explanation.objects.update_or_create(
            target_type=target_type,
            content_type=content_type,
            object_id=object_id,
            defaults={'content': content, 'user': request.user}
        )
        
        # Log action and update feedback status if target is Feedback
        if target_type == 'Feedback':
            fb = Feedback.objects.filter(id=object_id).first()
            if fb:
                if fb.status == 'pending':
                    fb.status = 'reviewed'
                    fb.save()
                ActionLog.objects.create(
                    user=request.user,
                    feedback=fb,
                    action="Lưu giải trình",
                    details=f"Nội dung: {content[:100]}..."
                )
            
        return Response({"message": "Đã lưu giải trình"})

    @action(detail=False, methods=['post'])
    def delete_explanation(self, request):
        document_id = request.data.get('document_id')
        if not self._check_permission(request, document_id, 'Chuyên viên Giải trình'):
            return Response({"error": "Bạn không có quyền Xóa giải trình cho dự thảo này."}, status=403)
            
        target_type = request.data.get('target_type')
        object_id = request.data.get('object_id')
        
        if target_type == 'Feedback':
            content_type = ContentType.objects.get_for_model(Feedback)
        else:
            content_type = ContentType.objects.get_for_model(DocumentNode)
            
        deleted_count, _ = Explanation.objects.filter(
            target_type=target_type,
            content_type=content_type,
            object_id=object_id
        ).delete()
        
        if target_type == 'Feedback' and deleted_count > 0:
            fb = Feedback.objects.filter(id=object_id).first()
            if fb:
                fb.status = 'pending'
                fb.save()
                ActionLog.objects.create(
                    user=request.user,
                    feedback=fb,
                    action="Xóa giải trình",
                    details="Chuyên viên đã xóa nội dung giải trình."
                )
            
        return Response({"message": "Đã xóa giải trình thành công."})

    @action(detail=True, methods=['post'])
    def submit_for_review(self, request, pk=None):
        fb = self.get_object()
        if not self._check_permission(request, fb.document_id, 'Chuyên viên Giải trình'):
            return Response({"error": "No permission"}, status=403)
        
        fb.status = 'reviewed'
        fb.save()
        ActionLog.objects.create(
            user=request.user,
            feedback=fb,
            action="Gửi duyệt",
            details="Chuyên viên đã gửi duyệt nội dung giải trình."
        )
        return Response({"message": "Đã gửi duyệt thành công"})

    @action(detail=True, methods=['post'])
    def approve_feedback(self, request, pk=None):
        fb = self.get_object()
        # Admin or Lead only
        user = request.user
        is_lead = fb.document.lead == user or user.is_staff or user.is_superuser
        if not is_lead:
            return Response({"error": "Chỉ Lãnh đạo/Admin mới có quyền phê duyệt."}, status=403)
        
        fb.status = 'approved'
        fb.save()
        ActionLog.objects.create(
            user=request.user,
            feedback=fb,
            action="Phê duyệt",
            details="Lãnh đạo đã phê duyệt nội dung giải trình."
        )
        return Response({"message": "Đã phê duyệt thành công"})

    @action(detail=False, methods=['get'])
    def by_node(self, request):
        """Lấy danh sách góp ý phẳng cho một node (và các con của nó)"""
        node_id = request.query_params.get('node_id')
        filter_type = request.query_params.get('filter_type', 'has_feedback')
        if not node_id: return Response([])
        
        from documents.models import DocumentNode
        try:
            root_node = DocumentNode.objects.get(id=node_id)
            
            def get_all_descendant_ids(node):
                ids = [node.id]
                for child in node.children.all():
                    ids.extend(get_all_descendant_ids(child))
                return ids
            
            target_ids = get_all_descendant_ids(root_node)
            
            feedbacks = Feedback.objects.filter(node_id__in=target_ids)\
                .select_related('node', 'agency')\
                .prefetch_related('explanations', 'logs', 'logs__user')\
                .order_by('created_at')
            
            results = []
            for fb in feedbacks:
                path = []
                curr = fb.node
                while curr:
                    path.insert(0, curr.node_label)
                    curr = curr.parent
                
                explanation_obj = fb.explanations.first()
                if filter_type == 'resolved' and not explanation_obj: continue
                if filter_type == 'unresolved' and explanation_obj: continue
                
                results.append({
                    "id": fb.id,
                    "node_id": fb.node_id,
                    "node_path": ", ".join(path),
                    "node_content": fb.node.content if fb.node else "",
                    "contributing_agency": fb.contributing_agency or (fb.agency.name if fb.agency else "Ẩn danh"),
                    "agency_category": fb.agency.category if fb.agency else "other",
                    "content": fb.content,
                    "explanation": explanation_obj.content if explanation_obj else "",
                    "status": fb.status,
                    "created_at": fb.created_at.isoformat(),
                    "logs": [
                        {
                            "username": log.user.username,
                            "action": log.action,
                            "time": log.created_at.isoformat(),
                            "details": log.details
                        } for log in fb.logs.all()
                    ]
                })
            return Response(results)
        except DocumentNode.DoesNotExist:
            return Response({"error": "Node not found"}, status=404)

    @action(detail=False, methods=['get'])
    def subject_stats(self, request):
        document_id = request.query_params.get('document_id')
        query = Feedback.objects.all()
        if document_id:
            query = query.filter(document_id=document_id)
            
        from django.db.models.functions import Coalesce
        agency_counts = query.annotate(
            display_name=Coalesce('agency__name', 'contributing_agency')
        ).values('display_name', 'agency__category', 'agency__agency_category__name').annotate(
            total=models.Count('id'),
            resolved=models.Count('id', filter=models.Q(explanations__isnull=False))
        ).order_by('-total')
        
        DEFAULT_CAT = 'Phân loại khác'
        agency_stats_list = []
        found_categories = set()
        
        for item in agency_counts:
            raw_cat = item['agency__agency_category__name'] or item['agency__category'] or 'other'
            cat_label = raw_cat if raw_cat != 'other' else DEFAULT_CAT
            found_categories.add(cat_label)
            
            total = item['total']
            resolved = item['resolved']
            agency_stats_list.append({
                'agency': item['display_name'] or 'Ẩn danh',
                'total': total,
                'resolved': resolved,
                'resolve_rate': round((resolved / total * 100), 1) if total > 0 else 0,
                'category': cat_label
            })

        category_stats = {}
        for stat in query.values('agency__agency_category__name', 'agency__category').annotate(count=models.Count('agency', distinct=True)):
            raw_cat = stat['agency__agency_category__name'] or stat['agency__category'] or 'other'
            label = raw_cat if raw_cat != 'other' else DEFAULT_CAT
            category_stats[label] = category_stats.get(label, 0) + stat['count']
        
        invited_category_stats = {}
        if document_id:
            from documents.models import Document
            try:
                doc = Document.objects.get(id=document_id)
                for cat_stat in doc.consulted_agencies.values('agency_category__name', 'category').annotate(count=models.Count('id')):
                    raw_cat = cat_stat['agency_category__name'] or cat_stat['category'] or 'other'
                    label = raw_cat if raw_cat != 'other' else DEFAULT_CAT
                    invited_category_stats[label] = invited_category_stats.get(label, 0) + cat_stat['count']
            except Document.DoesNotExist:
                pass

        # 5. Thống kê cấp ý kiến (mới bổ sung)
        AGREED_REGEX = r'thống\s+nhất|nhất\s+trí'
        ACCEPTED_REGEX = r'tiếp\s+thu'
        
        total_fbs = query.count()
        count_agreed = query.filter(models.Q(content__iregex=AGREED_REGEX)).count()
        
        # Feedback có giải trình 'tiếp thu'
        from .models import Explanation
        accepted_fb_ids = Explanation.objects.filter(
            target_type='Feedback',
            content__iregex=ACCEPTED_REGEX
        ).values_list('object_id', flat=True)
        
        count_accepted = query.filter(id__in=accepted_fb_ids).distinct().count()
        
        # Feedback có giải trình nhưng không tiếp thu
        count_explained_no_acc = query.filter(
            explanations__isnull=False
        ).exclude(
            id__in=accepted_fb_ids
        ).distinct().count()
        
        # Feedback chưa giải trình
        count_pending = query.filter(explanations__isnull=True).count()

        return Response({
            'agency_stats': agency_stats_list,
            'category_stats': category_stats,
            'invited_category_stats': invited_category_stats,
            'available_categories': sorted(list(found_categories)),
            'summary': {
                'total_fbs': total_fbs,
                'total_agreed': count_agreed,
                'total_accepted': count_accepted,
                'total_explained_no_acc': count_explained_no_acc,
                'total_pending': count_pending
            }
        })

    @action(detail=False, methods=['get'])
    def by_appendix(self, request):
        """Lấy danh sách góp ý cho một Phụ lục"""
        appendix_id = request.query_params.get('appendix_id')
        if not appendix_id: return Response([])
        try:
            feedbacks = Feedback.objects.filter(appendix_id=appendix_id)\
                .select_related('appendix', 'agency')\
                .prefetch_related('explanations', 'logs', 'logs__user').order_by('created_at')
            results = []
            for fb in feedbacks:
                explanation_obj = fb.explanations.first()
                results.append({
                    "id": fb.id,
                    "appendix_id": fb.appendix_id,
                    "node_path": f"Phụ lục: {fb.appendix.name}",
                    "node_content": fb.appendix.content,
                    "contributing_agency": fb.contributing_agency or (fb.agency.name if fb.agency else "Ẩn danh"),
                    "agency_category": fb.agency.category if fb.agency else "other",
                    "content": fb.content,
                    "explanation": explanation_obj.content if explanation_obj else "",
                    "status": fb.status,
                    "created_at": fb.created_at.isoformat(),
                    "logs": [{"username": log.user.username, "action": log.action, "time": log.created_at.isoformat(), "details": log.details} for log in fb.logs.all()]
                })
            return Response(results)
        except Exception as e:
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['get'])
    def by_document(self, request):
        """Lấy toàn bộ danh sách góp ý của một Dự thảo"""
        doc_id = request.query_params.get('document_id')
        if not doc_id:
            return Response({"error": "Vui lòng cung cấp document_id"}, status=400)
            
        feedbacks = Feedback.objects.filter(document_id=doc_id)\
            .select_related('node', 'agency', 'appendix')\
            .prefetch_related('explanations', 'logs', 'logs__user')\
            .order_by('node__order_index', 'created_at')
            
        results = []
        for fb in feedbacks:
            # Xây dựng path cho node hoặc dùng Phụ lục
            node_label = "Chung"
            node_path = ""
            if fb.node:
                node_label = fb.node.node_label
                path_parts = []
                curr = fb.node
                while curr:
                    path_parts.insert(0, curr.node_label)
                    curr = curr.parent
                node_path = " > ".join(path_parts)
            elif fb.appendix:
                node_label = f"Phụ lục: {fb.appendix.name}"
                node_path = "Phụ lục"

            explanation_obj = fb.explanations.first()
            
            results.append({
                "id": fb.id,
                "document_id": fb.document_id,
                "node_id": fb.node_id,
                "node_label": node_label,
                "node_path": node_path,
                "agency": fb.agency_id,
                "contributing_agency": fb.contributing_agency or (fb.agency.name if fb.agency else "Ẩn danh"),
                "official_doc_number": fb.official_doc_number,
                "content": fb.content,
                "explanation": explanation_obj.content if explanation_obj else "",
                "status": fb.status,
                "created_at": fb.created_at.isoformat(),
            })
            
        return Response(results)

    @action(detail=False, methods=['get'])
    def uncontributed(self, request):
        """Trả về danh sách các cơ quan CHƯA có góp ý cho dự thảo cụ thể"""
        doc_id = request.query_params.get('document_id')
        if not doc_id:
            return Response({"error": "Vui lòng cung cấp document_id"}, status=400)
            
        from core.models import Agency
        # Lấy ID của các Agency đã góp ý cho dự thảo này
        contributed_agency_ids = Feedback.objects.filter(
            document_id=doc_id, agency__isnull=False
        ).values_list('agency_id', flat=True).distinct()
        
        # Các Agency còn lại là chưa góp ý
        uncontributed = Agency.objects.exclude(id__in=contributed_agency_ids).values('id', 'name', 'category')
        
        return Response(list(uncontributed))

    @action(detail=False, methods=['get'])
    def custom_report_preview(self, request):
        doc_id = request.query_params.get('document_id')
        agency = request.query_params.get('agency')
        status_filter = request.query_params.get('status')
        specialist = request.query_params.get('specialist')
        report_type = request.query_params.get('report_type', 'mau10')
        
        if not doc_id: return Response([])
        
        feedbacks = Feedback.objects.filter(document_id=doc_id).select_related('node', 'agency').prefetch_related('explanations', 'user').order_by('node__order_index')
        
        if agency and agency != 'all':
            import unicodedata
            norm_target = unicodedata.normalize('NFC', agency).strip().lower()
            feedbacks = [
                fb for fb in feedbacks 
                if (fb.contributing_agency and unicodedata.normalize('NFC', fb.contributing_agency).strip().lower() == norm_target) or
                (fb.agency and unicodedata.normalize('NFC', fb.agency.name).strip().lower() == norm_target)
            ]
            
        if status_filter == 'resolved':
            if isinstance(feedbacks, list):
                feedbacks = [fb for fb in feedbacks if fb.explanations.exists()]
            else:
                feedbacks = feedbacks.filter(explanations__isnull=False).distinct()
        elif status_filter == 'unresolved':
            if isinstance(feedbacks, list):
                feedbacks = [fb for fb in feedbacks if not fb.explanations.exists()]
            else:
                feedbacks = feedbacks.filter(explanations__isnull=True).distinct()

        if specialist and specialist != 'all':
            if specialist == 'none':
                if isinstance(feedbacks, list):
                    feedbacks = [fb for fb in feedbacks if not any(ex.user for ex in fb.explanations.all())]
                else:
                    feedbacks = feedbacks.filter(explanations__user__isnull=True)
            else:
                if isinstance(feedbacks, list):
                    feedbacks = [fb for fb in feedbacks if any(str(ex.user_id) == str(specialist) for ex in fb.explanations.all())]
                else:
                    feedbacks = feedbacks.filter(explanations__user_id=specialist)
            
        # Lay cau hinh truong tu template
        from reports.models import ReportTemplate
        template = ReportTemplate.objects.filter(template_type=report_type, is_active=True).first()
        fields = []
        if template:
            fields = template.field_configs.filter(is_enabled=True).order_by('column_order')

        from .utils.v2_template_generator import _get_field_value
        
        results = []
        for i, fb in enumerate(feedbacks, 1):
            explanation = fb.explanations.first()
            if fields:
                row = { f.field_key: _get_field_value(f.field_key, i, fb, explanation) for f in fields }
            else:
                # Fallback mac dinh
                dieu_khoan = f"{fb.node.node_label}" if fb.node else ""
                if fb.node and fb.node.parent:
                    dieu_khoan = f"{fb.node.parent.node_label}, {fb.node.node_label}"
                row = {
                    "stt": i,
                    "dieu_khoan": dieu_khoan,
                    "co_quan": (fb.agency.name if fb.agency else fb.contributing_agency) or "Khác",
                    "noi_dung_gop_y": fb.content,
                    "noi_dung_giai_trinh": explanation.content if explanation else "",
                }
            results.append(row)
            
        return Response(results)

    @action(detail=False, methods=['get'], permission_classes=[])
    def export_mau_10(self, request):
        doc_id = request.query_params.get('document_id')
        agency = request.query_params.get('agency')
        status_filter = request.query_params.get('status')
        
        # Xác thực qua token ở URL (Cho phép tải file trực tiếp từ trình duyệt)
        user = request.user
        if not user.is_authenticated:
            token = request.query_params.get('token')
            if not token:
                return Response({"error": "Vui lòng cung cấp token"}, status=401)
            from rest_framework_simplejwt.tokens import AccessToken
            from django.contrib.auth import get_user_model
            try:
                access_token = AccessToken(token)
                user = get_user_model().objects.get(id=access_token['user_id'])
            except Exception:
                return Response({"error": "Token không hợp lệ"}, status=401)
        
        if not doc_id:
            return Response({"error": "Vui lòng cung cấp document_id"}, status=400)
            
        try:
            document = Document.objects.get(id=doc_id)
            feedbacks = Feedback.objects.filter(document_id=doc_id).select_related('node', 'agency').prefetch_related('explanations').order_by('node__order_index')
            
            if agency and agency != 'all':
                import unicodedata
                norm_target = unicodedata.normalize('NFC', agency).strip().lower()
                feedbacks = [
                    fb for fb in feedbacks 
                    if (fb.contributing_agency and unicodedata.normalize('NFC', fb.contributing_agency).strip().lower() == norm_target) or
                    (fb.agency and unicodedata.normalize('NFC', fb.agency.name).strip().lower() == norm_target)
                ]
                
            if status_filter == 'resolved':
                if isinstance(feedbacks, list):
                    feedbacks = [fb for fb in feedbacks if fb.explanations.exists()]
                else:
                    feedbacks = feedbacks.filter(explanations__isnull=False).distinct()
            elif status_filter == 'unresolved':
                if isinstance(feedbacks, list):
                    feedbacks = [fb for fb in feedbacks if not fb.explanations.exists()]
                else:
                    feedbacks = feedbacks.filter(explanations__isnull=True).distinct()
            
            # Kiểm tra an toàn xem feedbacks (list hoặc queryset) có trống không
            has_data = feedbacks.exists() if hasattr(feedbacks, 'exists') else (len(feedbacks) > 0)
            if not has_data:
                return Response({"error": "Không có ý kiến góp ý nào thỏa mãn bộ lọc để xuất báo cáo."}, status=404)
            
            # Phan nhanh theo loai bao cao (report_type param)
            report_type = request.query_params.get('report_type', 'mau10')
            if report_type == 'mau10': report_type = 'mau_10' # Chuẩn hóa về format DB
            
            # Doc cau hinh tu DB cho loai tuong ung
            template_config = {}
            tpl_db = None
            try:
                from reports.models import ReportTemplate, ReportFieldConfig
                tpl_db = ReportTemplate.objects.filter(template_type=report_type, is_active=True).first()
                if tpl_db:
                    fields_qs = tpl_db.field_configs.filter(is_enabled=True).order_by('column_order')
                    fields = [
                        {
                            'field_key': f.field_key,
                            'field_label': f.field_label,
                            'column_width_cm': f.column_width_cm
                        } for f in fields_qs
                    ]
                    template_config = {
                        'header_org_name': tpl_db.header_org_name,
                        'header_org_location': tpl_db.header_org_location,
                        'footer_signer_name': tpl_db.footer_signer_name,
                        'footer_signer_title': tpl_db.footer_signer_title,
                        'fields': fields
                    }
            except Exception as e:
                print(f"Error loading template config: {e}")

            from .utils.v2_template_generator import generate_from_v2_template
            from django.http import HttpResponse

            # HỢP NHẤT TUYỆT ĐỐI: Dùng V2 generator (Landscape) cho mọi chế độ xuất
            file_stream = generate_from_v2_template(
                document, 
                feedbacks, 
                template_config=template_config, 
                template_type=report_type
            )
            
            # Cấu hình tên file
            type_label = 'Mau_10' if report_type == 'mau_10' else 'Tuy_chinh'
            filename = f"Bao_cao_{type_label}_{document.id}.docx"
            
            # Trả về file Word cho người dùng (SỬ DỤNG HTTPRESPONSE ĐỂ ĐẢM BẢO ỔN ĐỊNH DỮ LIỆU)
            response = HttpResponse(
                file_stream.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
            
        except Document.DoesNotExist:
            return Response({"error": "Văn bản không tồn tại"}, status=404)
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": f"Lỗi xuất báo cáo: {str(e)}"}, status=500)

    @action(detail=True, methods=['post'])
    def reassign_node(self, request, pk=None):
        fb = self.get_object()
        new_node_id = request.data.get('node_id')
        if not new_node_id:
            return Response({"error": "node_id is required"}, status=400)
        
        try:
            from documents.models import DocumentNode
            new_node = DocumentNode.objects.get(id=new_node_id, document_id=fb.document_id)
            old_label = fb.node.node_label if fb.node else "N/A"
            fb.node = new_node
            fb.save()
            
            ActionLog.objects.create(
                user=request.user,
                feedback=fb,
                action="Gắn lại điều khoản",
                details=f"Chuyển từ '{old_label}' sang '{new_node.node_label}'"
            )
            
            return Response({"message": "Đã chuyển điều khoản thành công."})
        except DocumentNode.DoesNotExist:
            return Response({"error": "Điều khoản mới không hợp lệ hoặc không thuộc dự thảo này."}, status=404)

    def _get_full_node_path(self, node):
        if not node: return "Chung"
        path = []
        current = node
        while current:
            path.insert(0, current.node_label)
            current = current.parent
        return ", ".join(path)

    @action(detail=False, methods=['post'])
    def gsheet_compare(self, request):
        document_id = request.data.get('document_id')
        gs_url = request.data.get('gs_url')
        
        if not document_id or not gs_url:
            return Response({"error": "Thiếu document_id hoặc gs_url."}, status=400)
            
        try:
            # 1. Fetch local feedbacks
            local_feedbacks = Feedback.objects.filter(document_id=document_id).select_related('node', 'agency', 'appendix').prefetch_related('explanations')
            
            # 2. Fetch GS rows
            worksheet, error_msg = self._get_gsheet_worksheet(gs_url)
            if not worksheet:
                return Response({"error": error_msg}, status=400)
            
            all_values = worksheet.get_all_values()
            if not all_values:
                return Response({"error": "Google Sheet trống hoặc không có dòng tiêu đề."}, status=400)
                
            headers = [h.strip().lower() for h in all_values[0]]
            rows = all_values[1:]
            
            # Identify columns using fuzzy matching keywords
            col_map = {}
            for idx, h in enumerate(headers):
                if 'node' not in col_map and any(kw in h for kw in ['điều', 'khoản', 'mục', 'vị trí']): 
                    col_map['node'] = idx
                elif 'explanation' not in col_map and any(kw in h for kw in ['giải trình', 'tiếp thu']): 
                    col_map['explanation'] = idx
                elif 'agency' not in col_map and any(kw in h for kw in ['cơ quan', 'chủ thể', 'đơn vị', 'người']): 
                    col_map['agency'] = idx
                elif 'content' not in col_map and any(kw in h for kw in ['nội dung', 'ý kiến', 'góp ý']): 
                    col_map['content'] = idx

            if 'content' not in col_map:
                 return Response({"error": "Không tìm thấy cột 'Nội dung góp ý' trong Google Sheet để đối chiếu."}, status=400)

            # 3. Compare logic
            import unicodedata
            import re
            from difflib import SequenceMatcher

            def norm(s):
                if not s: return ""
                # Bỏ HTML tags
                s = re.sub(r'<[^>]+>', ' ', str(s))
                # Xoá các ký tự đặc biệt, dấu câu
                s = re.sub(r'[^\w\s]', ' ', s)
                s = re.sub(r'\s+', ' ', s)
                return unicodedata.normalize('NFC', s).strip().lower()

            def get_key(content, agency):
                return f"{norm(content)}|{norm(agency)}"

            gs_data = {}
            unmatched_gs = []
            for i, row in enumerate(rows):
                raw_content = row[col_map['content']] if len(row) > col_map['content'] else ""
                raw_agency = row[col_map['agency']] if 'agency' in col_map and len(row) > col_map['agency'] else ""
                raw_exp = row[col_map['explanation']] if 'explanation' in col_map and len(row) > col_map['explanation'] else ""

                entry = {
                    "row": i + 2,
                    "content": raw_content,
                    "agency": raw_agency,
                    "explanation": raw_exp,
                    "norm_content": norm(raw_content),
                    "norm_agency": norm(raw_agency)
                }

                key = get_key(raw_content, raw_agency)
                if key and key != "|":
                    gs_data[key] = entry
                    unmatched_gs.append(entry)

            results = []
            for fb in local_feedbacks:
                norm_fb_content = norm(fb.content)
                norm_fb_agency = norm(fb.contributing_agency)
                key = f"{norm_fb_content}|{norm_fb_agency}"
                
                # Check exact match first
                gs_info = gs_data.get(key)
                
                # Try fuzzy matching if no exact match
                if not gs_info and norm_fb_content:
                    best_match = None
                    best_ratio = 0
                    for gs in unmatched_gs:
                        # Thử fuzzy trên nội dung chính. 
                        # Yêu cầu cơ quan đóng góp trùng khớp hoặc gần giống nếu có khai báo.
                        valid_agency = (not gs["norm_agency"] or not norm_fb_agency or 
                                        gs["norm_agency"] == norm_fb_agency or 
                                        SequenceMatcher(None, norm_fb_agency, gs["norm_agency"]).ratio() > 0.8)
                        
                        if valid_agency:
                            ratio = SequenceMatcher(None, norm_fb_content, gs["norm_content"]).ratio()
                            if ratio > best_ratio and ratio >= 0.85:
                                best_ratio = ratio
                                best_match = gs
                    
                    if best_match:
                        gs_info = best_match

                # Nếu đã khớp thì bỏ khỏi list danh sách chờ để không móc trùng
                if gs_info and gs_info in unmatched_gs:
                    unmatched_gs.remove(gs_info)
                
                exp = fb.explanations.first()
                exp_content = exp.content if exp else ""
                
                # Check for specific diffs even if matched
                is_content_diff = False
                is_exp_diff = False
                if gs_info:
                    is_content_diff = norm_fb_content != gs_info["norm_content"]
                    is_exp_diff = norm(exp_content) != norm(gs_info["explanation"])

                results.append({
                    "id": fb.id,
                    "node_label": self._get_full_node_path(fb.node) if fb.node else (f"Phụ lục: {fb.appendix.name}" if fb.appendix else "Chung"),
                    "agency": fb.contributing_agency,
                    "content": fb.content,
                    "explanation": exp_content,
                    "is_in_gs": gs_info is not None,
                    "gs_row": gs_info["row"] if gs_info else None,
                    "gs_content": gs_info["content"] if gs_info else "",
                    "gs_explanation": gs_info["explanation"] if gs_info else "",
                    "is_content_diff": is_content_diff,
                    "is_exp_diff": is_exp_diff,
                    "status": "synced" if (gs_info and not is_exp_diff) else ("diff" if is_exp_diff else "new_in_db")
                })
                
            return Response({"feedbacks": results})
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": str(e)}, status=500)

    @action(detail=False, methods=['post'])
    def gsheet_push(self, request):
        document_id = request.data.get('document_id')
        gs_url = request.data.get('gs_url')
        push_items = request.data.get('push_items', [])
        
        # Hỗ trợ legacy (trường hợp api gửi danh sách id thô)
        feedback_ids = request.data.get('feedback_ids', [])
        
        if not document_id or not gs_url or (not push_items and not feedback_ids):
            return Response({"error": "Thiếu tham số bắt buộc (document_id, gs_url, danh sách đẩy)."}, status=400)
            
        try:
            worksheet, error_msg = self._get_gsheet_worksheet(gs_url)
            if not worksheet:
                return Response({"error": error_msg}, status=400)
            
            if not push_items and feedback_ids:
                push_items = [{"id": fid, "gs_row": None} for fid in feedback_ids]
                
            all_fids = [item['id'] for item in push_items]
            feedbacks = Feedback.objects.filter(id__in=all_fids).select_related('node', 'appendix').prefetch_related('explanations')
            feedback_dict = {fb.id: fb for fb in feedbacks}
            
            # Map headers to know where to write
            raw_headers = worksheet.row_values(1)
            headers = [h.strip().lower() for h in raw_headers]
            
            col_node = col_agency = col_content = col_reason = col_explanation = col_note = col_nd = col_gt = None
            
            for idx, h in enumerate(headers):
                # Ưu tiên các cột trạng thái viết tắt trước để tránh bị hốt vào cột nội dung chính
                if not col_nd and h == 'nd': col_nd = idx + 1
                elif not col_gt and h == 'gt': col_gt = idx + 1
                elif not col_node and any(kw in h for kw in ['điều', 'khoản', 'mục', 'vị trí']): col_node = idx + 1
                elif not col_explanation and any(kw in h for kw in ['giải trình', 'tiếp thu']): col_explanation = idx + 1
                elif not col_agency and any(kw in h for kw in ['cơ quan', 'chủ thể', 'đơn vị', 'người']): col_agency = idx + 1
                elif not col_content and any(kw in h for kw in ['nội dung', 'ý kiến', 'góp ý']): col_content = idx + 1
                elif not col_reason and any(kw in h for kw in ['lý do', 'cơ sở']): col_reason = idx + 1
                elif not col_note and any(kw in h for kw in ['ghi chú', 'note']): col_note = idx + 1

            if not col_content:
                return Response({"error": "Google Sheet không có cột 'Nội dung góp ý' để định danh dữ liệu."}, status=400)
            
            rows_to_append = []
            cells_to_update = []
            from gspread.cell import Cell
            
            for item in push_items:
                fb = feedback_dict.get(item['id'])
                if not fb: continue
                
                node_val = self._get_full_node_path(fb.node) if fb.node else (f"Phụ lục: {fb.appendix.name}" if fb.appendix else "Chung")
                exp = fb.explanations.first()
                exp_val = exp.content if exp else ""

                gs_row = item.get('gs_row')
                if gs_row: # Cập nhật
                    if col_node: cells_to_update.append(Cell(row=gs_row, col=col_node, value=node_val))
                    if col_agency: cells_to_update.append(Cell(row=gs_row, col=col_agency, value=fb.contributing_agency))
                    if col_content: cells_to_update.append(Cell(row=gs_row, col=col_content, value=fb.content))
                    if col_reason: cells_to_update.append(Cell(row=gs_row, col=col_reason, value=fb.reason or ""))
                    if col_explanation: cells_to_update.append(Cell(row=gs_row, col=col_explanation, value=exp_val))
                    if col_note: cells_to_update.append(Cell(row=gs_row, col=col_note, value=fb.note or ""))
                    
                    # Tự động điền trạng thái OK
                    if col_nd and fb.content: cells_to_update.append(Cell(row=gs_row, col=col_nd, value="OK"))
                    if col_gt and exp_val: cells_to_update.append(Cell(row=gs_row, col=col_gt, value="OK"))
                else: # Thêm mới
                    row = [""] * len(headers)
                    if col_node: row[col_node-1] = node_val
                    if col_agency: row[col_agency-1] = fb.contributing_agency
                    if col_content: row[col_content-1] = fb.content
                    if col_reason: row[col_reason-1] = fb.reason or ""
                    if col_explanation: row[col_explanation-1] = exp_val
                    if col_note: row[col_note-1] = fb.note or ""
                    
                    # Trạng thái OK cho dòng mới
                    if col_nd and fb.content: row[col_nd-1] = "OK"
                    if col_gt and exp_val: row[col_gt-1] = "OK"
                    
                    rows_to_append.append(row)
                
            if rows_to_append:
                worksheet.append_rows(rows_to_append)
                
            if cells_to_update:
                worksheet.update_cells(cells_to_update)
                
            msg = []
            if rows_to_append: msg.append(f"Thêm mới {len(rows_to_append)} dòng.")
            if cells_to_update: 
                updated_count = len(push_items) - len(rows_to_append)
                msg.append(f"Cập nhật {updated_count} dòng có sẵn.")
                
            return Response({"message": " ".join(msg) or "Không có thay đổi nào."})
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            return Response({"error": str(e)}, status=500)

class ConsultationResponseViewSet(viewsets.ModelViewSet):
    queryset = ConsultationResponse.objects.all()
    serializer_class = ConsultationResponseSerializer
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        doc_id = self.request.query_params.get('document_id')
        if doc_id:
            return ConsultationResponse.objects.filter(document_id=doc_id).order_by('-created_at')
        return super().get_queryset()




