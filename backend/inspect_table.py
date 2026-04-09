import docx
import sys

def inspect_table(file_path):
    doc = docx.Document(file_path)
    if not doc.tables:
        print("No tables found.")
        return

    table = doc.tables[0]
    print(f"Total rows: {len(table.rows)}")
    
    # Check column count
    col_count = len(table.rows[0].cells)
    print(f"Columns: {col_count}")

    # Inspect rows 20 to 40
    for i in range(max(0, 20), min(len(table.rows), 50)):
        row = table.rows[i]
        c1 = row.cells[0].text.strip().replace('\n', ' ')[:50]
        c2 = row.cells[1].text.strip().replace('\n', ' ')[:50] if len(row.cells) > 1 else "N/A"
        print(f"Row {i:03d} | C1: [{c1}] | C2: [{c2}]")

if __name__ == "__main__":
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    inspect_table('c:/Users/Quoc Anh/Desktop/QLVB/bangthuyetminh.docx')
