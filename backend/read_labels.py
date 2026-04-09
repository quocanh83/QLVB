import docx
import sys

def read_docx_tables(file_path):
    doc = docx.Document(file_path)
    if not doc.tables:
        print("No tables found in the document.")
        return

    table = doc.tables[0]
    print(f"Total rows in Table 0: {len(table.rows)}")
    
    # Print the first 100 rows' first column to see the numbering
    for i, row in enumerate(table.rows):
        label = row.cells[0].text.strip()
        # Only print labels to avoid huge output
        print(f"Row {i:03d} | Label: {label}")

if __name__ == "__main__":
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    read_docx_tables('c:/Users/Quoc Anh/Desktop/QLVB/bangthuyetminh.docx')
