from rest_framework import viewsets, permissions, status
from rest_framework.response import Response
from rest_framework.decorators import action
from django.db.models import Count, Q
from django.db import transaction
from .models import Document, DocumentNode, NodeAssignment, DocumentType
from .serializers import DocumentListSerializer, DocumentUploadSerializer, DocumentNodeSerializer, DocumentTypeSerializer

from accounts.models import User
import docx
import re
import io
from datetime import datetime
from django.http import FileResponse, HttpResponse
from docxtpl import DocxTemplate
import os
from .utils.parser_engine import ParserEngine

def clean_text(text):
    """Giữ lại các ký tự hợp lệ cho XML 1.0, loại bỏ ký tự điều khiển và escape các ký tự đặc biệt."""
    if text is None:
        return ""
    text = str(text)
    # Loại bỏ các ký tự không hợp lệ cho XML 1.0
    text = re.sub(r'[^\x09\x0A\x0D\x20-\uD7FF\uE000-\uFFFD\U00010000-\U0010FFFF]', '', text)
    # Escape các ký tự đặc biệt của XML để tránh làm hỏng cấu trúc file Word
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

class DocumentViewSet(viewsets.ModelViewSet):
    permission_classes = [permissions.IsAuthenticated]

    def get_queryset(self):
        queryset = Document.objects.annotate(
            total_dieu=Count('nodes', filter=Q(nodes__node_type='Điều'), distinct=True),
            total_khoan=Count('nodes', filter=Q(nodes__node_type='Khoản'), distinct=True),
            total_diem=Count('nodes', filter=Q(nodes__node_type='Điểm'), distinct=True),
            total_phu_luc=Count('nodes', filter=Q(nodes__node_type='Phụ lục'), distinct=True),
            total_nodes=Count('nodes', distinct=True),
            total_feedbacks=Count('feedbacks', distinct=True),
            resolved_feedbacks=Count('feedbacks', filter=Q(feedbacks__explanations__isnull=False), distinct=True)
        )
        
        document_type_id = self.request.query_params.get('document_type')
        if document_type_id:
            queryset = queryset.filter(document_type_id=document_type_id)
            
        return queryset.order_by('-id')

    def get_serializer_class(self):
        if self.action == 'create':
            return DocumentUploadSerializer
        return DocumentListSerializer

    @action(detail=True, methods=['post'])
    def set_lead(self, request, pk=None):
        """Admin chỉ định Chủ trì (Lead) cho một Dự thảo"""
        user = request.user
        user = request.user
        is_admin = user.is_staff or user.is_superuser or (hasattr(user, 'roles') and user.roles.filter(role_name='Admin').exists())
        if not is_admin:
            return Response({"error": "Chỉ Admin mới có thể chỉ định Chủ trì."}, status=403)
        document = self.get_object()
        lead_id = request.data.get('lead_id')
        if lead_id:
            try:
                lead_user = User.objects.get(id=lead_id)
                document.lead = lead_user
            except User.DoesNotExist:
                return Response({"error": "Không tìm thấy tài khoản cán bộ."}, status=404)
        else:
            document.lead = None  # Gỡ chủ trì
        document.save()
        return Response({"message": "Đã cập nhật thông tin Chủ trì (Lead) thành công!"})

    def create(self, request, *args, **kwargs):
        serializer = self.get_serializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        try:
            with transaction.atomic():
                document = serializer.save(uploaded_by=request.user)
                file_obj = request.FILES.get('attached_file_path')
                if file_obj:
                    parser = ParserEngine(file_obj)
                    structure = parser.parse()
                    
                    def save_nodes_recursive(node_list, parent=None):
                        nonlocal index
                        for node_data in node_list:
                            children = node_data.pop('children', [])
                            node = DocumentNode.objects.create(
                                document=document,
                                parent=parent,
                                **node_data
                            )
                            if children:
                                save_nodes_recursive(children, parent=node)

                    index = 0
                    save_nodes_recursive(structure)
                
                # Luôn tạo node "Chung" cho dự thảo
                DocumentNode.objects.get_or_create(
                    document=document,
                    node_type='Vấn đề khác',
                    node_label='Chung',
                    defaults={'content': '', 'order_index': -1}
                )

                # Tự động tạo node Vấn đề khác ở cuối cùng (nếu có file structure)
                if file_obj:
                    DocumentNode.objects.get_or_create(
                        document=document,
                        node_type='Vấn đề khác',
                        node_label='Vấn đề khác',
                        defaults={'content': '', 'order_index': 9999}
                    )
        except Exception as e:
            return Response({'error': str(e)}, status=status.HTTP_400_BAD_REQUEST)
        headers = self.get_success_headers(serializer.data)
        return Response(serializer.data, status=status.HTTP_201_CREATED, headers=headers)

    @action(detail=False, methods=['post'])
    def parse_preview(self, request):
        """Trả về cấu trúc cây thư mục của file docx mà không lưu database"""
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "Vui lòng đính kèm file .docx"}, status=400)
        
        try:
            parser = ParserEngine(file_obj)
            structure = parser.parse()
            return Response(structure)
        except Exception as e:
            return Response({"error": f"Lỗi parse: {str(e)}"}, status=400)

    @action(detail=False, methods=['get'])
    def type_stats(self, request):
        """Lấy danh sách các Loại văn bản kèm theo số lượng Dự thảo của từng loại"""
        from .models import DocumentType, Document
        from django.db.models import Count
        
        stats = DocumentType.objects.annotate(
            doc_count=Count('documents')
        ).order_by('-doc_count')
        
        data = [
            {
                "id": s.id,
                "name": s.name,
                "count": s.doc_count
            } for s in stats
        ]
        # Thêm mục 'Tổng số'
        total_count = Document.objects.count()
        data.insert(0, {"id": "All", "name": "Tất cả", "count": total_count})
        
        return Response(data)

    @action(detail=True, methods=['get'], permission_classes=[])
    def nodes(self, request, pk=None):
        document = self.get_object()
        user = request.user
        
        is_admin = user.is_staff or user.is_superuser or (hasattr(user, 'roles') and user.roles.filter(role_name='Admin').exists())
        is_lead = (document.lead == user)
        
        if is_admin or is_lead:
            all_nodes = DocumentNode.objects.filter(document=document).annotate(
                total_feedbacks=Count('feedbacks', distinct=True),
                resolved_feedbacks=Count('feedbacks', filter=Q(feedbacks__explanations__isnull=False), distinct=True)
            ).prefetch_related('assignments', 'assignments__user').order_by('order_index')
            for n in all_nodes:
                n.is_editable = True
        else:
            # Data Isolation ORM Magic
            assigned_node_ids = NodeAssignment.objects.filter(
                user=user, 
                node__document=document
            ).values_list('node_id', flat=True)
            
            all_nodes = DocumentNode.objects.filter(document=document).filter(
                Q(id__in=assigned_node_ids) |
                Q(children__id__in=assigned_node_ids) |
                Q(children__children__id__in=assigned_node_ids)
            ).annotate(
                total_feedbacks=Count('feedbacks', distinct=True),
                resolved_feedbacks=Count('feedbacks', filter=Q(feedbacks__explanations__isnull=False), distinct=True)
            ).prefetch_related('assignments', 'assignments__user').distinct().order_by('order_index')
            
            assigned_set = set(assigned_node_ids)
            for n in all_nodes:
                n.is_editable = (n.id in assigned_set)

        # Build In-Memory Tree (O(N) Complexity)
        node_dict = {n.id: n for n in all_nodes}
        root_nodes = []
        for n in all_nodes:
            n.prefetched_children = []
            
        for n in all_nodes:
            if n.parent_id and n.parent_id in node_dict:
                node_dict[n.parent_id].prefetched_children.append(n)
            else:
                root_nodes.append(n)
                
        serializer = DocumentNodeSerializer(root_nodes, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['post'])
    def assign_nodes(self, request, pk=None):
        document = self.get_object()
        user = request.user
        is_admin = user.is_staff or user.is_superuser or (hasattr(user, 'roles') and user.roles.filter(role_name='Admin').exists())
        if not (is_admin or document.lead == user):
            return Response({"error": "Bạn không có quyền Phân công dự án này."}, status=403)
            
        assignments_data = request.data.get('assignments', [])
        
        with transaction.atomic():
            for item in assignments_data:
                node_id = item.get('node_id')
                user_ids = item.get('user_ids', [])
                
                NodeAssignment.objects.filter(node_id=node_id).delete()
                
                new_assignments = [
                    NodeAssignment(node_id=node_id, user_id=uid, assigned_by=user)
                    for uid in user_ids
                ]
                NodeAssignment.objects.bulk_create(new_assignments)
                
        return Response({"message": "Cập nhật phân công thành công!"})

    @action(detail=True, methods=['get'], permission_classes=[])
    def node_details(self, request, pk=None):
        node_id = request.query_params.get('node_id')
        from feedbacks.models import Feedback, Explanation
        from django.contrib.contenttypes.models import ContentType
        try:
            root_node = DocumentNode.objects.get(id=node_id, document_id=pk)
            
            # Hàm đệ quy lấy con
            def get_all_descendants(node):
                descendants = [node]
                for child in node.children.all():
                    descendants.extend(get_all_descendants(child))
                return descendants
            
            all_target_nodes = get_all_descendants(root_node)
            node_ct = ContentType.objects.get_for_model(DocumentNode)
            
            results = []
            for n in all_target_nodes:
                feedbacks = Feedback.objects.filter(node=n).order_by('-created_at')
                feedback_data = []
                for fb in feedbacks:
                    explanations = fb.explanations.all()
                    feedback_data.append({
                        "id": fb.id,
                        "contributing_agency": fb.contributing_agency,
                        "content": fb.content,
                        "created_at": fb.created_at,
                        "explanations": [{"id": ex.id, "content": ex.content} for ex in explanations]
                    })
                
                node_explanations = Explanation.objects.filter(content_type=node_ct, object_id=n.id)
                results.append({
                    "id": n.id,
                    "node_label": n.node_label,
                    "node_type": n.node_type,
                    "content": n.content,
                    "feedbacks": feedback_data,
                    "node_explanations": [{"id": ex.id, "content": ex.content} for ex in node_explanations]
                })
            
            return Response(results)
        except DocumentNode.DoesNotExist:
            return Response({"error": "Node not found"}, status=404)

    @action(detail=False, methods=['get'])
    def explanation_stats(self, request):
        """Thống kê giải trình cho danh sách dự thảo"""
        queryset = Document.objects.annotate(
            total_feedbacks=Count('feedbacks', distinct=True),
            resolved_feedbacks=Count('feedbacks', filter=Q(feedbacks__explanations__isnull=False), distinct=True)
        ).order_by('-id')
        
        serializer = DocumentListSerializer(queryset, many=True)
        return Response(serializer.data)


    @action(detail=False, methods=['get'])
    def dashboard_stats(self, request):
        """API cung cấp dữ liệu cho trang Tổng quan (Dashboard)"""
        from feedbacks.models import Feedback
        from django.db.models.functions import TruncDate
        from datetime import timedelta
        from django.utils import timezone

        # KPI Cards
        total_docs = Document.objects.count()
        total_feedbacks = Feedback.objects.count()
        
        # resolved_feedbacks: feedbacks that have status 'reviewed' or 'approved'
        resolved_feedbacks = Feedback.objects.filter(status__in=['reviewed', 'approved']).distinct().count()
        
        # Lấy count danh sách agencies distinct không bị none/rỗng
        agencies_qs = Feedback.objects.exclude(contributing_agency__isnull=True).exclude(contributing_agency='').values('contributing_agency').distinct()
        agencies_count = agencies_qs.count()
        
        # Top Documents (Dự thảo nóng nhất - Top 5)
        top_docs_qs = Document.objects.annotate(
            feedback_count=Count('feedbacks', distinct=True)
        ).order_by('-feedback_count', '-id')[:5]
        
        top_docs = [{"id": d.id, "name": d.project_name, "feedbacks": d.feedback_count} for d in top_docs_qs]
        
        # Trend Data (Góp ý theo ngày - 7 ngày gần nhất)
        today = timezone.now()
        seven_days_ago = (today - timedelta(days=6)).replace(hour=0, minute=0, second=0, microsecond=0)
        
        trend_qs = Feedback.objects.filter(created_at__gte=seven_days_ago).annotate(
            date=TruncDate('created_at')
        ).values('date').annotate(count=Count('id', distinct=True)).order_by('date')
        
        trend_dict = {item['date'].strftime('%d/%m'): item['count'] for item in trend_qs if item['date']}
        
        trend_data = []
        for i in range(6, -1, -1):
            day_str = (today - timedelta(days=i)).strftime('%d/%m')
            trend_data.append({
                "date": day_str,
                "count": trend_dict.get(day_str, 0)
            })
        
        # Merged Recent Activity (5 hoạt động mới nhất: Documents + Feedbacks)
        recent_feedbacks = Feedback.objects.select_related('user', 'document').order_by('-created_at')[:5]
        recent_documents = Document.objects.select_related('uploaded_by').order_by('-created_at')[:5]
        
        combined_activity = []
        for fb in recent_feedbacks:
            combined_activity.append({
                "id": f"fb-{fb.id}",
                "type": "feedback",
                "user": fb.user.username if fb.user else "Hệ thống",
                "document": fb.document.project_name,
                "time": fb.created_at,
                "content": "đã gửi góp ý mới"
            })
        for doc in recent_documents:
            combined_activity.append({
                "id": f"doc-{doc.id}",
                "type": "document",
                "user": doc.uploaded_by.username if doc.uploaded_by else "Hệ thống",
                "document": doc.project_name,
                "time": doc.created_at,
                "content": "đã tải lên dự thảo mới"
            })
        
        # Sort combined and take top 5
        combined_activity.sort(key=lambda x: x['time'], reverse=True)
        recent_activity = combined_activity[:5]
            
        return Response({
            "cards": {
                "totalDocs": total_docs,
                "totalFeedbacks": total_feedbacks,
                "resolvedFeedbacks": resolved_feedbacks,
                "agenciesCount": agencies_count
            },
            "topDocs": top_docs,
            "trendData": trend_data,
            "recentActivity": recent_activity
        })

    @action(detail=True, methods=['get'])
    def feedback_nodes(self, request, pk=None):
        """Lấy cây thư mục chỉ chứa các node có góp ý, hỗ trợ lọc theo cơ quan"""
        document = self.get_object()
        agency = request.query_params.get('agency')
        
        # Tìm tất cả các node có góp ý hoặc có con/cháu có góp ý
        # Lọc nhãn số lượng theo cơ quan nếu có
        query_all_fb = Q(feedbacks__isnull=False)
        query_exp_fb = Q(feedbacks__explanations__isnull=False)
        
        if agency:
            query_all_fb &= Q(feedbacks__contributing_agency__icontains=agency)
            query_exp_fb &= Q(feedbacks__contributing_agency__icontains=agency)

        nodes = DocumentNode.objects.filter(document=document).annotate(
            total_fb=Count('feedbacks', filter=query_all_fb, distinct=True),
            explained_fb=Count('feedbacks', filter=query_exp_fb, distinct=True)
        ).order_by('order_index')
        
        # Build tree in-memory
        node_dict = {n.id: n for n in nodes}
        for n in nodes:
            n.prefetched_children = []
            n.total_feedbacks = n.total_fb
            n.resolved_feedbacks = n.explained_fb
            
        root_nodes = []
        for n in nodes:
            if n.parent_id and n.parent_id in node_dict:
                node_dict[n.parent_id].prefetched_children.append(n)
            else:
                root_nodes.append(n)
        
        filter_type = request.query_params.get('filter_type', 'has_feedback')

        # Hàm đệ quy để tính tổng số lượng góp ý từ dưới lên
        def aggregate_counts_recursive(node):
            total = node.total_fb
            resolved = node.explained_fb
            
            valid_children = []
            for child in node.prefetched_children:
                child_total, child_resolved, should_keep = aggregate_counts_recursive(child)
                
                # Luôn cộng dồn số lượng để hiển thị đúng tổng số, bất kể filter
                total += child_total
                resolved += child_resolved
                
                if should_keep:
                    valid_children.append(child)
            
            node.total_feedbacks = total
            node.resolved_feedbacks = resolved
            node.prefetched_children = valid_children
            
            # Quyết định xem branch này có được giữ lại không dựa trên filter_type
            if filter_type == 'all':
                keep_node = (total > 0) if agency else True
            elif filter_type == 'resolved':
                keep_node = (resolved > 0)
            elif filter_type == 'unresolved':
                keep_node = ((total - resolved) > 0)
            else: # has_feedback (default)
                keep_node = (total > 0)
                
            return total, resolved, keep_node

        pruned_roots = []
        for r in root_nodes:
            _, _, should_keep = aggregate_counts_recursive(r)
            if should_keep:
                pruned_roots.append(r)
        
        serializer = DocumentNodeSerializer(pruned_roots, many=True)
        return Response(serializer.data)

    @action(detail=True, methods=['get'], permission_classes=[])
    def unresolved_feedbacks(self, request, pk=None):
        """Liệt kê các góp ý chưa có giải trình của một văn bản"""
        from feedbacks.models import Feedback
        # Feedback chưa có explanation
        unresolved = Feedback.objects.filter(
            document_id=pk,
            explanations__isnull=True
        ).select_related('node').order_by('node__order_index')
        
        data = []
        for fb in unresolved:
            # Tạo nhãn đường dẫn: Điều X, Khoản Y...
            path = []
            curr = fb.node
            while curr:
                path.insert(0, curr.node_label)
                curr = curr.parent
            
            data.append({
                "id": fb.id,
                "node_id": fb.node_id,
                "node_path": " > ".join(path),
                "contributing_agency": fb.contributing_agency or "Ẩn danh",
                "content": fb.content[:200] + ("..." if len(fb.content) > 200 else ""),
                "created_at": fb.created_at
            })
        return Response(data)

    @action(detail=True, methods=['get'], permission_classes=[])
    def export_report(self, request, pk=None):
        try:
            # Xác thực qua token ở URL (Cho phép tải file trực tiếp từ trình duyệt)
            user = request.user
            if not user.is_authenticated:
                token = request.query_params.get('token')
                if not token:
                    return Response({"error": "Vui lòng cung cấp token"}, status=401)
                from rest_framework_simplejwt.tokens import AccessToken
                from django.contrib.auth import get_user_model
                try:
                    access_token = AccessToken(token)
                    user = get_user_model().objects.get(id=access_token['user_id'])
                except Exception:
                    return Response({"error": "Token không hợp lệ"}, status=401)

            doc_obj = self.get_object()
            
            # Hợp nhất: Sử dụng generator chuẩn Mẫu 10 mới
            from feedbacks.utils.mau10_generator import generate_mau_10
            from feedbacks.models import Feedback
            try:
                from reports.models import ReportTemplate
            except ImportError:
                ReportTemplate = None
            
            feedbacks = Feedback.objects.filter(document_id=doc_obj.id).select_related('node').prefetch_related('explanations').order_by('node__order_index')
            
            # Đọc cấu hình mẫu
            template_config = None
            if ReportTemplate:
                tpl = ReportTemplate.objects.filter(template_type='mau_10', is_active=True).first()
                if tpl:
                    enabled_fields = tpl.field_configs.filter(is_enabled=True).order_by('column_order')
                    if enabled_fields.exists():
                        template_config = {
                            'header_org_name': tpl.header_org_name,
                            'header_org_location': tpl.header_org_location,
                            'footer_signer_name': tpl.footer_signer_name,
                            'footer_signer_title': tpl.footer_signer_title,
                            'fields': [
                                {
                                    'field_key': f.field_key,
                                    'field_label': f.field_label,
                                    'column_width_cm': f.column_width_cm,
                                }
                                for f in enabled_fields
                            ]
                        }
            
            file_stream = generate_mau_10(doc_obj, feedbacks, template_config=template_config)
            
            response = FileResponse(
                file_stream, 
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            filename = f"Bao_cao_tong_hop_{doc_obj.id}.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response
        except Exception as e:
            import traceback
            print(traceback.format_exc())
            return Response({"error": f"Lỗi trong quá trình xuất báo cáo: {str(e)}"}, status=500)

    @action(detail=False, methods=['post'])
    def match_agencies_from_file(self, request):
        file_obj = request.FILES.get('file')
        if not file_obj:
            return Response({"error": "Không có tệp tải lên."}, status=400)
        
        try:
            from core.models import Agency
            all_agencies = Agency.objects.all()
            text = ""
            
            if file_obj.name.endswith('.docx'):
                doc = docx.Document(file_obj)
                text = "\n".join([p.text for p in doc.paragraphs])
                for table in doc.tables:
                    for row in table.rows:
                        text += "\n" + " ".join([cell.text for cell in row.cells])
            elif file_obj.name.endswith('.xlsx') or file_obj.name.endswith('.xls'):
                import pandas as pd
                df = pd.read_excel(file_obj)
                text = df.to_string()
            else:
                return Response({"error": "Chỉ hỗ trợ tệp .docx hoặc .xlsx"}, status=400)
            
            matched_ids = []
            for agency in all_agencies:
                # Search for agency name in text (case-insensitive)
                # Using regex to ensure word boundary if possible, but simple icontains is safer for long names
                if agency.name.lower() in text.lower():
                    matched_ids.append(agency.id)
            
            return Response({"matched_ids": matched_ids})
        except Exception as e:
            return Response({"error": f"Lỗi xử lý tệp: {str(e)}"}, status=500)

class NodeViewSet(viewsets.ModelViewSet):
    queryset = DocumentNode.objects.all()
    serializer_class = DocumentNodeSerializer
    permission_classes = [permissions.IsAuthenticated]

    @action(detail=True, methods=['get'])
    def full_context(self, request, pk=None):
        """Lấy trọn bộ 'Điều' chứa node này kèm các con của nó"""
        node = self.get_object()
        
        # Nếu là Khoản/Điểm, tìm ngược lên Điều
        root = node
        while root.parent and root.node_type != 'Điều':
            root = root.parent
            
        # Serialize root kèm children
        def serialize_recursive(n):
            children = n.children.all().order_by('order_index')
            return {
                "id": n.id,
                "node_type": n.node_type,
                "node_label": n.node_label,
                "content": n.content,
                "children": [serialize_recursive(c) for c in children]
            }
            
        return Response(serialize_recursive(root))

class DocumentTypeViewSet(viewsets.ModelViewSet):
    queryset = DocumentType.objects.all().order_by('-created_at')
    serializer_class = DocumentTypeSerializer
    permission_classes = [permissions.IsAuthenticated]
