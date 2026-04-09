import docx
import sys
import re

def check_and_fix_numbering(file_path, output_path):
    doc = docx.Document(file_path)
    if not doc.tables:
        print("No tables found.")
        return

    table = doc.tables[0]
    print(f"Total rows: {len(table.rows)}")
    
    current_num = 0
    fixed_count = 0

    # Vietnamese "Điều" literal. 
    # Let's use a regex that is more flexible with characters.
    # [\u0110\u0111][i\u1ec1u\u00ea\u00f4] (D/d + ieu/eo)
    # Actually just match anything that looks like "Word Number" if Column 1 is dedicated to labels.
    
    for i, row in enumerate(table.rows):
        c1_text = row.cells[0].text.strip()
        if not c1_text:
            continue
            
        print(f"Row {i:03d} | RAW: [{c1_text[:30]}]")
        
        # Match "Điều X" or any word followed by space and digit
        match = re.match(r'^([^\d]+\s+)(\d+)(.*)', c1_text, re.IGNORECASE | re.DOTALL)
        if match:
            current_num += 1
            prefix = match.group(1)
            original_num_str = match.group(2)
            original_num = int(original_num_str)
            suffix = match.group(3)
            
            if original_num != current_num:
                new_text = f"{prefix}{current_num}{suffix}"
                row.cells[0].text = new_text
                print(f"  -> FIXED: {original_num} TO {current_num}")
                fixed_count += 1
            else:
                print(f"  -> OK: {original_num}")
                pass
        else:
            print(f"  -> NO MATCH")

    doc.save(output_path)
    print(f"\nSUCCESS: Fixed {fixed_count} articles. Saved to {output_path}")

if __name__ == "__main__":
    if hasattr(sys.stdout, 'reconfigure'):
        sys.stdout.reconfigure(encoding='utf-8')
    check_and_fix_numbering('c:/Users/Quoc Anh/Desktop/QLVB/bangthuyetminh.docx', 'c:/Users/Quoc Anh/Desktop/QLVB/bangthuyetminh_fixed.docx')
