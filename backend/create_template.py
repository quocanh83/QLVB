from docx import Document

def create_template():
    doc = Document()
    doc.add_heading('BẢN TỔNG HỢP Ý KIẾN VÀ GIẢI TRÌNH', 0)
    doc.add_paragraph('Tên cơ quan, tổ chức chủ trì: {{ agency_name }}')
    doc.add_paragraph('Địa danh: {{ headquarters_location }}')
    doc.add_paragraph('Tên dự thảo văn bản: {{ document_title }}')
    doc.add_paragraph('Ngày xuất báo cáo: {{ export_date }}')
    doc.add_paragraph('---')

    # Bắt đầu vòng lặp Điều
    doc.add_paragraph('{% for dieu in dieu_list %}')
    doc.add_heading('{{ dieu.node_label }}', level=1)
    doc.add_paragraph('Nội dung gốc: {{ dieu.content }}')

    # Lặp Feedbacks của Điều
    doc.add_paragraph('{% for fb in dieu.feedbacks %}')
    doc.add_paragraph('Góp ý của [{{ fb.user_name }}]: {{ fb.content }}', style='List Bullet')
    doc.add_paragraph('{% for exp in fb.explanations %}')
    doc.add_paragraph('Giải trình: {{ exp.content }}', style='List Number')
    doc.add_paragraph('{% endfor %}') # End Explanations
    doc.add_paragraph('{% endfor %}') # End Feedbacks

    # --- Bắt đầu vòng lặp Khoản ---
    doc.add_paragraph('{% for khoan in dieu.khoan_list %}')
    doc.add_heading('{{ khoan.node_label }}', level=2)
    doc.add_paragraph('Nội dung gốc: {{ khoan.content }}')

    doc.add_paragraph('{% for fb in khoan.feedbacks %}')
    doc.add_paragraph('Góp ý của [{{ fb.user_name }}]: {{ fb.content }}', style='List Bullet')
    doc.add_paragraph('{% for exp in fb.explanations %}')
    doc.add_paragraph('Giải trình: {{ exp.content }}', style='List Number')
    doc.add_paragraph('{% endfor %}') # End Explanations
    doc.add_paragraph('{% endfor %}') # End Feedbacks

    # --- Bắt đầu vòng lặp Điểm ---
    doc.add_paragraph('{% for diem in khoan.diem_list %}')
    doc.add_heading('{{ diem.node_label }}', level=3)
    doc.add_paragraph('Nội dung gốc: {{ diem.content }}')

    doc.add_paragraph('{% for fb in diem.feedbacks %}')
    doc.add_paragraph('Góp ý của [{{ fb.user_name }}]: {{ fb.content }}', style='List Bullet')
    doc.add_paragraph('{% for exp in fb.explanations %}')
    doc.add_paragraph('Giải trình: {{ exp.content }}', style='List Number')
    doc.add_paragraph('{% endfor %}') # End Explanations
    doc.add_paragraph('{% endfor %}') # End Feedbacks

    doc.add_paragraph('{% endfor %}') # End Loop Điểm
    doc.add_paragraph('{% endfor %}') # End Loop Khoản
    doc.add_paragraph('{% endfor %}') # End Loop Điều

    save_path = r'C:\Users\Quoc Anh\Desktop\QLVB\template_bao_cao_chuan_v2.docx'
    doc.save(save_path)
    print(f"Đã tạo file template chuẩn tại: {save_path}")

if __name__ == "__main__":
    create_template()
