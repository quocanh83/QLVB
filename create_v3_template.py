from docx import Document
from docx.shared import Cm, Pt
from docx.enum.section import WD_ORIENT
import os

def create_template():
    doc = Document()
    
    # Set orientation to Landscape
    section = doc.sections[0]
    new_width, new_height = section.page_height, section.page_width
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = new_width
    section.page_height = new_height
    
    # Administrative Header
    table_header = doc.add_table(rows=1, cols=2)
    table_header.autofit = True
    cells = table_header.rows[0].cells
    cells[0].text = "{{ drafting_agency }}\\nSố: ....../BC-......"
    cells[1].text = "CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM\\nĐộc lập - Tự do - Hạnh phúc\\n-------------------\\n{{ agency_location }}, {{ export_date }}"
    
    doc.add_paragraph()
    title = doc.add_heading("BÁO CÁO TỔNG HỢP Ý KIẾN GÓP Ý\\n{{ document_title }}", 0)
    
    doc.add_paragraph("Tổng số cơ quan gửi ý kiến: {{ total_consulted }}")
    doc.add_paragraph("Tổng số ý kiến nhận được: {{ total_feedbacks }}")

    # Main Data Table
    # Rows: 1 (Header), 1 (Loop start), 1 (Loop end)
    # Actually we just need rows that docxtpl can duplicate
    table = doc.add_table(rows=3, cols=7)
    table.style = 'Table Grid'
    
    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_names = ["STT", "Điều/Khoản", "Cơ quan góp ý", "Nội dung góp ý", "Nội dung giải trình", "Chuyên viên", "Trạng thái"]
    for i, name in enumerate(hdr_names):
        hdr_cells[i].text = name
        
    # Data Loop Row (for d in dieu_list)
    # This row will show the Node Label across all columns
    row_d = table.rows[1]
    row_d.cells[0].text = "{% for d in dieu_list %}"
    row_d.cells[1].text = "{{ d.node_label }}"
    # Merge cells 1 to 6 for node label
    row_d.cells[1].merge(row_d.cells[6])
    
    # Data Loop Row (for fb in d.feedbacks)
    row_fb = table.rows[2]
    # We need to insert a row for the feedback loop
    # Actually docxtpl can handle this if we put tags in cells
    row_fb.cells[0].text = "{% for fb in d.feedbacks %}{{ fb.stt }}"
    row_fb.cells[1].text = "{{ d.node_label }}"
    row_fb.cells[2].text = "{{ fb.user_name }}"
    row_fb.cells[3].text = "{{ fb.content }}"
    # Explanation loop inside cell 4
    row_fb.cells[4].text = "{% for ex in fb.explanations %}{{ ex.content }} {% endfor %}"
    row_fb.cells[5].text = "{{ fb.chuyen_vien }}"
    row_fb.cells[6].text = "{{ fb.status }}{% endfor %}{% endfor %}"

    save_path = 'backend/feedbacks/utils/template_truong_tu_chinh_v3.docx'
    doc.save(save_path)
    print(f"Template created at: {save_path}")

if __name__ == "__main__":
    create_template()
