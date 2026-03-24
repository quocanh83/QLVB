import docx
import os

def create_sample_docx(path):
    doc = docx.Document()
    doc.add_heading('DỰ THẢO NGHỊ ĐỊNH', 0)
    
    doc.add_paragraph('Chương I: QUY ĐỊNH CHUNG')
    doc.add_paragraph('Điều 1. Phạm vi điều chỉnh')
    doc.add_paragraph('Nghị định này quy định về việc quản lý và vận hành hệ thống QLVB.')
    
    doc.add_paragraph('Điều 2. Đối tượng áp dụng')
    doc.add_paragraph('1. Cơ quan nhà nước.')
    doc.add_paragraph('2. Tổ chức, cá nhân có liên quan.')
    
    doc.add_paragraph('Chương II: ĐIỀU KIỆN HOẠT ĐỘNG')
    doc.add_paragraph('Điều 15. Điều kiện cấp phép')
    doc.add_paragraph('1. Có tư cách pháp nhân.')
    doc.add_paragraph('2. Có phương án kinh doanh.')
    doc.add_paragraph('a) Phương án tài chính.')
    doc.add_paragraph('b) Phương án nhân sự.')
    
    doc.add_paragraph('Phụ lục I: DANH MỤC BIỂU MẪU')
    doc.add_paragraph('Mẫu số 10: Báo cáo kết quả tiếp thu.')
    
    doc.save(path)
    print(f"Created sample docx at {path}")

if __name__ == "__main__":
    path = os.path.join(os.getcwd(), 'sample_draft.docx')
    create_sample_docx(path)
